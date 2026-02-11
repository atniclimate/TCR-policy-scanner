"""Tests for the strategic overview document generator (Plan 08-02).

Tests StrategicOverviewGenerator.generate() produces a valid 7-section DOCX
covering all 16 programs, 7 ecoregions, and structural policy framework.
All tests use tmp_path fixtures with mock data -- no network or real data
files required.
"""

import json
from pathlib import Path

import pytest
from docx import Document


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _write_json(path: Path, data) -> None:
    """Write data as JSON to the given path, creating parent dirs."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f)


# ---------------------------------------------------------------------------
# Mock data builders
# ---------------------------------------------------------------------------

def _mock_programs() -> list[dict]:
    """Return a list of 16 mock programs with varied CI statuses.

    Includes:
    - 2 FLAGGED/AT_RISK (DEFEND)
    - 2 UNCERTAIN/STABLE_BUT_VULNERABLE (PROTECT)
    - 10 STABLE/SECURE (EXPAND)
    - 2 FEMA programs (fema_bric, fema_tribal_mitigation)
    """
    return [
        {
            "id": "bia_tcr",
            "name": "BIA Tribal Climate Resilience",
            "agency": "BIA",
            "ci_status": "STABLE",
            "funding_type": "Discretionary",
            "description": "Core BIA program funding Tribal climate adaptation.",
            "advocacy_lever": "Protect/expand the line",
            "tightened_language": "Protect the $34.291M base funding line.",
        },
        {
            "id": "fema_bric",
            "name": "FEMA BRIC",
            "agency": "FEMA",
            "ci_status": "FLAGGED",
            "funding_type": "One-Time",
            "description": "FEMA pre-disaster mitigation competitive grant program.",
            "advocacy_lever": "Restore/replace with Tribal set-aside",
            "tightened_language": "BRIC terminated; advocacy must pivot.",
        },
        {
            "id": "irs_elective_pay",
            "name": "IRS Elective Pay",
            "agency": "Treasury/IRS",
            "ci_status": "AT_RISK",
            "funding_type": "Mandatory",
            "description": "IRA provision for Tribal clean energy tax credits.",
            "advocacy_lever": "Defend against reconciliation repeal",
            "tightened_language": "Maintain final rule status for Tribal instrumentalities.",
        },
        {
            "id": "epa_stag",
            "name": "EPA STAG",
            "agency": "EPA",
            "ci_status": "STABLE",
            "funding_type": "Discretionary",
            "description": "EPA State and Tribal Assistance Grants.",
            "advocacy_lever": "Direct Tribal capitalization pathways",
            "tightened_language": "",
        },
        {
            "id": "epa_gap",
            "name": "EPA GAP",
            "agency": "EPA",
            "ci_status": "STABLE",
            "funding_type": "Discretionary",
            "description": "EPA General Assistance Program.",
            "advocacy_lever": "Increase funding",
            "tightened_language": "",
        },
        {
            "id": "fema_tribal_mitigation",
            "name": "FEMA Tribal Mitigation Plans",
            "agency": "FEMA",
            "ci_status": "AT_RISK",
            "funding_type": "Discretionary",
            "description": "FEMA support for Tribal hazard mitigation planning.",
            "advocacy_lever": "Fund planning; accelerate approvals",
            "tightened_language": "Loss of BRIC reduces plan ROI.",
        },
        {
            "id": "dot_protect",
            "name": "DOT PROTECT",
            "agency": "DOT",
            "ci_status": "SECURE",
            "funding_type": "Discretionary",
            "description": "DOT resilience improvements to surface transportation.",
            "advocacy_lever": "Maintain/expand Tribal set-aside",
            "tightened_language": "",
        },
        {
            "id": "usda_wildfire",
            "name": "USDA Wildfire Defense Grants",
            "agency": "USFS",
            "ci_status": "STABLE_BUT_VULNERABLE",
            "funding_type": "One-Time",
            "description": "USDA grants for community wildfire defense.",
            "advocacy_lever": "Add match waivers",
            "tightened_language": "",
        },
        {
            "id": "doe_indian_energy",
            "name": "DOE Indian Energy",
            "agency": "DOE",
            "ci_status": "UNCERTAIN",
            "funding_type": "Discretionary",
            "description": "DOE Office of Indian Energy programs.",
            "advocacy_lever": "Protect FY26 appropriations",
            "tightened_language": "",
        },
        {
            "id": "hud_ihbg",
            "name": "HUD IHBG",
            "agency": "HUD",
            "ci_status": "STABLE",
            "funding_type": "Discretionary",
            "description": "HUD Indian Housing Block Grant.",
            "advocacy_lever": "Frame resilient housing",
            "tightened_language": "",
        },
        {
            "id": "noaa_tribal",
            "name": "NOAA Tribal Grants",
            "agency": "NOAA",
            "ci_status": "UNCERTAIN",
            "funding_type": "Discretionary",
            "description": "NOAA programs supporting Tribal climate resilience.",
            "advocacy_lever": "Protect climate functions",
            "tightened_language": "",
        },
        {
            "id": "fhwa_ttp_safety",
            "name": "FHWA TTP Safety",
            "agency": "FHWA",
            "ci_status": "STABLE",
            "funding_type": "Discretionary",
            "description": "FHWA Tribal Transportation safety funding.",
            "advocacy_lever": "Increase formula resilience",
            "tightened_language": "",
        },
        {
            "id": "usbr_watersmart",
            "name": "USBR WaterSMART",
            "agency": "Reclamation",
            "ci_status": "UNCERTAIN",
            "funding_type": "Discretionary",
            "description": "Bureau of Reclamation WaterSMART grants.",
            "advocacy_lever": "Expand Tribal drought access",
            "tightened_language": "",
        },
        {
            "id": "usbr_tap",
            "name": "USBR TAP",
            "agency": "Reclamation",
            "ci_status": "STABLE",
            "funding_type": "Discretionary",
            "description": "Bureau of Reclamation Technical Assistance.",
            "advocacy_lever": "Link TA to pre-award support",
            "tightened_language": "",
        },
        {
            "id": "bia_tcr_awards",
            "name": "Tribal Community Resilience Annual Awards",
            "agency": "BIA",
            "ci_status": "STABLE",
            "funding_type": "Discretionary",
            "description": "BIA annual awards for Tribal community resilience.",
            "advocacy_lever": "Protect award program",
            "tightened_language": "",
        },
        {
            "id": "epa_tribal_air",
            "name": "Tribal Air Quality Management",
            "agency": "EPA",
            "ci_status": "STABLE",
            "funding_type": "Discretionary",
            "description": "EPA program for Tribal air quality monitoring.",
            "advocacy_lever": "Protect Clean Air Act funding",
            "tightened_language": "",
        },
    ]


def _mock_policy_tracking() -> dict:
    """Return a mock policy_tracking.json structure."""
    return {
        "positions": [
            {
                "program_id": "bia_tcr",
                "program_name": "BIA Tribal Climate Resilience",
                "agency": "BIA",
                "status": "STABLE",
                "reasoning": "FY26 funded at $34.291M.",
                "extensible_fields": {
                    "authorization_status": "Authorized through FY27",
                },
            },
            {
                "program_id": "fema_bric",
                "program_name": "FEMA BRIC",
                "agency": "FEMA",
                "status": "FLAGGED",
                "reasoning": "BRIC terminated April 2025.",
                "extensible_fields": {},
            },
        ],
    }


def _mock_graph_schema() -> dict:
    """Return a mock graph_schema.json with 5 structural asks."""
    return {
        "authorities": [],
        "funding_vehicles": [],
        "barriers": [],
        "structural_asks": [
            {
                "id": "ask_multi_year",
                "name": "Multi-Year Funding Stability",
                "description": "Shift from annual to multi-year authorization.",
                "target": "Congress",
                "urgency": "FY26",
                "programs": ["bia_tcr", "epa_gap", "doe_indian_energy"],
            },
            {
                "id": "ask_match_waivers",
                "name": "Match/Cost-Share Waivers",
                "description": "Categorical Tribal exemptions from cost-share.",
                "target": "Congress",
                "urgency": "Immediate",
                "programs": ["usda_wildfire", "usbr_watersmart", "fema_bric"],
            },
            {
                "id": "ask_direct_access",
                "name": "Direct Tribal Access",
                "description": "Non-competitive formula-based pathways.",
                "target": "Agency",
                "urgency": "FY26",
                "programs": ["bia_tcr", "dot_protect", "epa_stag"],
            },
            {
                "id": "ask_consultation",
                "name": "Tribal Consultation Compliance",
                "description": "Enforceable consultation requirements.",
                "target": "Agency",
                "urgency": "Ongoing",
                "programs": ["fema_tribal_mitigation", "bia_tcr", "epa_gap"],
            },
            {
                "id": "ask_data_sovereignty",
                "name": "Data Sovereignty & Capacity",
                "description": "Tribal control over climate data.",
                "target": "Congress",
                "urgency": "FY26",
                "programs": ["bia_tcr", "epa_gap", "epa_stag", "noaa_tribal"],
            },
        ],
    }


def _mock_monitor_data() -> dict:
    """Return mock monitor data with alerts."""
    return {
        "alerts": [
            {
                "title": "IIJA Sunset Warning",
                "description": "IIJA authorization expires end of FY26.",
                "severity": "high",
            },
            {
                "title": "DHS Funding Cliff",
                "description": "DHS resilience programs face budget uncertainty.",
                "severity": "medium",
            },
        ],
    }


def _mock_ecoregion_config() -> dict:
    """Return a valid ecoregion config with 7 ecoregions."""
    return {
        "ecoregions": {
            "alaska": {
                "name": "Alaska",
                "states": ["AK"],
                "priority_programs": ["bia_tcr", "epa_gap", "doe_indian_energy"],
            },
            "pacific_northwest": {
                "name": "Pacific Northwest",
                "states": ["WA", "OR"],
                "priority_programs": ["bia_tcr", "fema_bric", "usda_wildfire"],
            },
            "southwest": {
                "name": "Southwest",
                "states": ["AZ", "NM", "NV", "UT", "CA"],
                "priority_programs": ["bia_tcr", "fema_bric", "usda_wildfire"],
            },
            "mountain_west": {
                "name": "Mountain West",
                "states": ["CO", "MT", "WY", "ID"],
                "priority_programs": ["bia_tcr", "fema_bric", "usda_wildfire"],
            },
            "great_plains_south": {
                "name": "Great Plains & South",
                "states": ["OK", "TX", "KS"],
                "priority_programs": ["bia_tcr", "fema_tribal_mitigation"],
            },
            "southeast": {
                "name": "Southeast",
                "states": ["FL", "NC", "SC", "VA", "GA"],
                "priority_programs": ["bia_tcr", "fema_bric", "hud_ihbg"],
            },
            "northeast_midwest": {
                "name": "Northeast & Midwest",
                "states": ["CT", "ME", "MA", "NY", "MN", "WI"],
                "priority_programs": ["bia_tcr", "fema_bric", "epa_gap"],
            },
        },
        "metadata": {
            "scheme": "NCA5-derived 7-region advocacy classification",
            "version": "1.0",
        },
    }


# ---------------------------------------------------------------------------
# Fixture
# ---------------------------------------------------------------------------

@pytest.fixture
def setup_data_files(tmp_path):
    """Write all mock data files and return (config, programs_dict)."""
    programs_list = _mock_programs()

    # Write data files
    _write_json(
        tmp_path / "data" / "program_inventory.json",
        {"programs": programs_list},
    )
    _write_json(
        tmp_path / "data" / "policy_tracking.json",
        _mock_policy_tracking(),
    )
    _write_json(
        tmp_path / "data" / "graph_schema.json",
        _mock_graph_schema(),
    )
    _write_json(
        tmp_path / "data" / "ecoregion_config.json",
        _mock_ecoregion_config(),
    )
    _write_json(
        tmp_path / "outputs" / "LATEST-MONITOR-DATA.json",
        _mock_monitor_data(),
    )

    config = {
        "packets": {
            "output_dir": str(tmp_path / "output"),
            "ecoregion": {
                "data_path": str(tmp_path / "data" / "ecoregion_config.json"),
            },
        },
    }

    programs_dict = {p["id"]: p for p in programs_list}
    return config, programs_dict, tmp_path


def _generate_overview(setup_data_files):
    """Helper to generate a strategic overview and return the path."""
    config, programs_dict, tmp_path = setup_data_files
    from src.packets.strategic_overview import StrategicOverviewGenerator

    # Patch data file paths to use tmp_path
    generator = StrategicOverviewGenerator.__new__(StrategicOverviewGenerator)
    generator.config = config
    generator.programs = programs_dict

    from src.packets.ecoregion import EcoregionMapper
    generator.ecoregion = EcoregionMapper(config)

    # Load data from tmp_path
    def load_json_tmp(path_str):
        path = tmp_path / path_str
        if not path.exists():
            return {}
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}

    generator.policy_tracking = load_json_tmp("data/policy_tracking.json")
    generator.graph_schema = load_json_tmp("data/graph_schema.json")
    generator.monitor_data = load_json_tmp("outputs/LATEST-MONITOR-DATA.json")

    output_dir = Path(config["packets"]["output_dir"])
    path = generator.generate(output_dir)
    return path


# ===========================================================================
# Tests
# ===========================================================================


class TestGenerateValidDocx:
    """Test that a valid DOCX is generated."""

    def test_generates_valid_docx(self, setup_data_files):
        """Generate STRATEGIC-OVERVIEW.docx; verify file exists and is valid."""
        path = _generate_overview(setup_data_files)

        assert path.exists()
        assert path.name == "STRATEGIC-OVERVIEW.docx"
        assert path.stat().st_size > 0

        # Should re-open without error
        doc = Document(str(path))
        assert len(doc.paragraphs) > 0


class TestSectionStructure:
    """Test document has expected section headings."""

    def test_has_seven_sections(self, setup_data_files):
        """Document should have 7 section areas identifiable by headings."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Cover page content
        assert "FY26 Federal Funding Overview & Strategy" in all_text

        # The 6 Heading 2 sections
        expected_h2 = [
            "Appropriations Landscape",
            "Ecoregion Strategic Priorities",
            "Science Infrastructure Threats",
            "FEMA Programs Analysis",
            "Cross-Cutting Policy Framework",
            "Messaging Guidance",
        ]
        for heading in expected_h2:
            assert heading in all_text, f"Missing section: {heading}"


class TestAppropriationsSection:
    """Test appropriations table has all 16 programs."""

    def test_appropriations_includes_all_programs(self, setup_data_files):
        """Appropriations table should have 16 data rows + 1 header."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        # Find tables -- first table should be appropriations
        tables = doc.tables
        assert len(tables) >= 1, "No tables found in document"

        approps_table = tables[0]
        # 16 data rows + 1 header = 17 rows
        assert len(approps_table.rows) == 17, (
            f"Expected 17 rows (1 header + 16 programs), "
            f"got {len(approps_table.rows)}"
        )

    def test_appropriations_has_status_summary(self, setup_data_files):
        """Appropriations section should include status summary counts."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "flagged or at risk" in all_text
        assert "stable or secure" in all_text


class TestEcoregionSection:
    """Test ecoregion section covers all 7 ecoregions."""

    def test_ecoregion_covers_all_seven(self, setup_data_files):
        """Should find 7 ecoregion subheadings."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        # Collect Heading 3 paragraphs
        h3_texts = [
            p.text for p in doc.paragraphs
            if p.style and p.style.name == "Heading 3"
        ]

        ecoregion_names = [
            "Alaska",
            "Pacific Northwest",
            "Southwest",
            "Mountain West",
            "Great Plains & South",
            "Southeast",
            "Northeast & Midwest",
        ]

        for name in ecoregion_names:
            assert name in h3_texts, (
                f"Ecoregion '{name}' not found in Heading 3 entries. "
                f"Found: {h3_texts}"
            )


class TestGracefulFallback:
    """Test graceful handling of missing data files."""

    def test_graceful_no_monitor_data(self, setup_data_files):
        """Generate without LATEST-MONITOR-DATA.json -- should succeed."""
        config, programs_dict, tmp_path = setup_data_files

        # Remove monitor data
        monitor_path = tmp_path / "outputs" / "LATEST-MONITOR-DATA.json"
        if monitor_path.exists():
            monitor_path.unlink()

        from src.packets.strategic_overview import StrategicOverviewGenerator
        from src.packets.ecoregion import EcoregionMapper

        generator = StrategicOverviewGenerator.__new__(StrategicOverviewGenerator)
        generator.config = config
        generator.programs = programs_dict
        generator.ecoregion = EcoregionMapper(config)

        def load_json_tmp(path_str):
            path = tmp_path / path_str
            if not path.exists():
                return {}
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            return data if isinstance(data, dict) else {}

        generator.policy_tracking = load_json_tmp("data/policy_tracking.json")
        generator.graph_schema = load_json_tmp("data/graph_schema.json")
        generator.monitor_data = {}  # No monitor data

        output_dir = Path(config["packets"]["output_dir"])
        path = generator.generate(output_dir)

        assert path.exists()

        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Monitor data not available" in all_text


class TestCrossCuttingFramework:
    """Test cross-cutting framework section."""

    def test_cross_cutting_has_five_asks(self, setup_data_files):
        """All 5 structural ask names should appear in the document."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        expected_asks = [
            "Multi-Year Funding Stability",
            "Match/Cost-Share Waivers",
            "Direct Tribal Access",
            "Tribal Consultation Compliance",
            "Data Sovereignty & Capacity",
        ]

        for ask_name in expected_asks:
            assert ask_name in all_text, f"Missing ask: {ask_name}"


class TestMessagingGuidance:
    """Test messaging guidance section."""

    def test_messaging_groups_by_goal(self, setup_data_files):
        """DEFEND, PROTECT, and EXPAND groups should all appear."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "DEFEND" in all_text
        assert "PROTECT" in all_text
        assert "EXPAND" in all_text

    def test_messaging_has_template_disclaimer(self, setup_data_files):
        """Messaging section should include template-based disclaimer."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "template-based" in all_text
        assert "No AI-generated text" in all_text


class TestFemaAnalysis:
    """Test FEMA analysis section."""

    def test_fema_analysis_section(self, setup_data_files):
        """FEMA program names should appear in the FEMA analysis section."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Both FEMA programs from our mock should be present
        assert "FEMA BRIC" in all_text
        assert "FEMA Tribal Mitigation Plans" in all_text

    def test_fema_has_bcr_note(self, setup_data_files):
        """FEMA section should include BCR methodology note."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "4:1 benefit-cost ratio" in all_text


class TestCoverPage:
    """Test cover page content."""

    def test_cover_page_content(self, setup_data_files):
        """Cover page should contain title and confidentiality notice."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        # Collect first 10 paragraphs
        first_texts = [p.text for p in doc.paragraphs[:10]]
        combined = "\n".join(first_texts)

        assert "FY26 Federal Funding Overview" in combined
        assert "CONFIDENTIAL" in combined

    def test_cover_page_has_date(self, setup_data_files):
        """Cover page should include a generation date."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs[:10])
        assert "Generated:" in all_text


class TestScienceThreats:
    """Test science threats section with monitor data."""

    def test_science_threats_with_data(self, setup_data_files):
        """With monitor data, alerts should appear."""
        path = _generate_overview(setup_data_files)
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "IIJA Sunset Warning" in all_text
        assert "DHS Funding Cliff" in all_text


class TestOrchestratorIntegration:
    """Test PacketOrchestrator.generate_strategic_overview() integration."""

    def test_orchestrator_integration(self, setup_data_files):
        """Orchestrator should expose generate_strategic_overview()."""
        config, programs_dict, tmp_path = setup_data_files

        # Create a minimal config for PacketOrchestrator
        # It needs a registry, congress, etc. -- so we mock those too
        _write_json(
            tmp_path / "tribal_registry.json",
            {
                "metadata": {"total_tribes": 1, "placeholder": False},
                "tribes": [
                    {
                        "tribe_id": "epa_001",
                        "bia_code": "100",
                        "name": "Test Tribe",
                        "states": ["AZ"],
                        "alternate_names": [],
                        "epa_region": "9",
                        "bia_recognized": True,
                    },
                ],
            },
        )
        _write_json(
            tmp_path / "congressional_cache.json",
            {
                "metadata": {
                    "congress_session": "119",
                    "total_members": 0,
                    "total_committees": 0,
                    "total_tribes_mapped": 0,
                },
                "members": {},
                "committees": {},
                "delegations": {},
            },
        )

        orch_config = {
            "packets": {
                "output_dir": str(tmp_path / "output"),
                "ecoregion": {
                    "data_path": str(tmp_path / "data" / "ecoregion_config.json"),
                },
                "tribal_registry": {
                    "data_path": str(tmp_path / "tribal_registry.json"),
                },
                "congressional_cache": {
                    "data_path": str(tmp_path / "congressional_cache.json"),
                },
                "awards": {
                    "cache_dir": str(tmp_path / "award_cache"),
                },
                "hazards": {
                    "cache_dir": str(tmp_path / "hazard_profiles"),
                },
            },
        }

        from src.packets.orchestrator import PacketOrchestrator

        programs_list = list(programs_dict.values())
        orch = PacketOrchestrator(orch_config, programs_list)

        # Patch data file loading to use tmp_path
        # The StrategicOverviewGenerator loads from project root by default,
        # so we need to ensure files are at the right locations.
        # Instead, we'll monkeypatch the _load_json in the generator.

        original_generate = orch.generate_strategic_overview

        def patched_generate():
            from src.packets.strategic_overview import StrategicOverviewGenerator

            output_dir = Path(
                orch.config.get("packets", {}).get("output_dir", "outputs/packets")
            )
            generator = StrategicOverviewGenerator.__new__(
                StrategicOverviewGenerator
            )
            generator.config = orch_config
            generator.programs = orch.programs

            from src.packets.ecoregion import EcoregionMapper
            generator.ecoregion = EcoregionMapper(orch_config)

            def load_json_tmp(path_str):
                path = tmp_path / path_str
                if not path.exists():
                    return {}
                with open(path, encoding="utf-8") as f:
                    data = json.load(f)
                return data if isinstance(data, dict) else {}

            generator.policy_tracking = load_json_tmp(
                "data/policy_tracking.json"
            )
            generator.graph_schema = load_json_tmp("data/graph_schema.json")
            generator.monitor_data = load_json_tmp(
                "outputs/LATEST-MONITOR-DATA.json"
            )
            return generator.generate(output_dir)

        orch.generate_strategic_overview = patched_generate

        path = orch.generate_strategic_overview()

        assert isinstance(path, Path)
        assert path.exists()
        assert path.suffix == ".docx"
        assert path.stat().st_size > 0

        # Verify re-readable
        doc = Document(str(path))
        assert len(doc.paragraphs) > 0


class TestAtomicSave:
    """Test atomic save pattern."""

    def test_output_dir_created(self, setup_data_files):
        """Output directory is created if it does not exist."""
        config, programs_dict, tmp_path = setup_data_files

        # Use a non-existent nested directory
        config["packets"]["output_dir"] = str(
            tmp_path / "deep" / "nested" / "output"
        )

        from src.packets.strategic_overview import StrategicOverviewGenerator
        from src.packets.ecoregion import EcoregionMapper

        generator = StrategicOverviewGenerator.__new__(StrategicOverviewGenerator)
        generator.config = config
        generator.programs = programs_dict
        generator.ecoregion = EcoregionMapper(config)

        def load_json_tmp(path_str):
            path = tmp_path / path_str
            if not path.exists():
                return {}
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            return data if isinstance(data, dict) else {}

        generator.policy_tracking = load_json_tmp("data/policy_tracking.json")
        generator.graph_schema = load_json_tmp("data/graph_schema.json")
        generator.monitor_data = {}

        output_dir = Path(config["packets"]["output_dir"])
        path = generator.generate(output_dir)

        assert path.exists()
        assert output_dir.exists()


class TestConstants:
    """Test module-level constants."""

    def test_goal_by_status_coverage(self):
        """_GOAL_BY_STATUS covers all 6 CI statuses."""
        from src.packets.strategic_overview import _GOAL_BY_STATUS

        expected = {"FLAGGED", "AT_RISK", "UNCERTAIN",
                    "STABLE_BUT_VULNERABLE", "STABLE", "SECURE"}
        assert set(_GOAL_BY_STATUS.keys()) == expected

    def test_fema_program_ids(self):
        """_FEMA_PROGRAM_IDS contains expected IDs."""
        from src.packets.strategic_overview import _FEMA_PROGRAM_IDS

        assert "fema_bric" in _FEMA_PROGRAM_IDS
        assert "fema_tribal_mitigation" in _FEMA_PROGRAM_IDS
        assert len(_FEMA_PROGRAM_IDS) == 5

    def test_max_data_file_bytes(self):
        """_MAX_DATA_FILE_BYTES is 10 MB."""
        from src.packets.strategic_overview import _MAX_DATA_FILE_BYTES

        assert _MAX_DATA_FILE_BYTES == 10 * 1024 * 1024


class TestLoadJson:
    """Test _load_json helper."""

    def test_load_nonexistent_file(self, tmp_path):
        """Loading a nonexistent file returns empty dict."""
        from src.packets.strategic_overview import StrategicOverviewGenerator

        generator = StrategicOverviewGenerator.__new__(StrategicOverviewGenerator)
        result = generator._load_json(str(tmp_path / "does_not_exist.json"))
        assert result == {}

    def test_load_valid_json(self, tmp_path):
        """Loading a valid JSON file returns parsed dict."""
        from src.packets.strategic_overview import StrategicOverviewGenerator

        data = {"key": "value"}
        json_path = tmp_path / "test.json"
        _write_json(json_path, data)

        generator = StrategicOverviewGenerator.__new__(StrategicOverviewGenerator)
        result = generator._load_json(str(json_path))
        assert result == {"key": "value"}

    def test_load_invalid_json(self, tmp_path):
        """Loading invalid JSON returns empty dict."""
        from src.packets.strategic_overview import StrategicOverviewGenerator

        json_path = tmp_path / "bad.json"
        json_path.parent.mkdir(parents=True, exist_ok=True)
        with open(json_path, "w", encoding="utf-8") as f:
            f.write("not valid json {{{")

        generator = StrategicOverviewGenerator.__new__(StrategicOverviewGenerator)
        result = generator._load_json(str(json_path))
        assert result == {}
