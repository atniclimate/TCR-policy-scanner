"""Tests for audience-differentiated document generation (Plan 14-03).

Verifies that Doc A (internal strategy) contains leverage, approach,
and messaging framework content while Doc B (congressional overview)
excludes all strategy content and provides evidence-only Key Ask boxes.

Also tests orchestrator multi-doc generation: complete-data Tribes
produce both Doc A and Doc B; partial-data Tribes produce Doc B only.

All tests use tmp_path fixtures with mock data -- no network or real
data files required.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from src.packets.context import TribePacketContext
from src.packets.doc_types import DOC_A, DOC_B
from src.packets.docx_engine import DocxEngine
from src.packets.economic import (
    EconomicImpactCalculator,
    TribeEconomicSummary,
)


# ---------------------------------------------------------------------------
# Mock data builders
# ---------------------------------------------------------------------------


def _mock_programs() -> list[dict]:
    """Return a list of mock programs."""
    return [
        {
            "id": "bia_tcr",
            "name": "BIA Tribal Climate Resilience",
            "agency": "Bureau of Indian Affairs",
            "federal_home": "DOI / BIA",
            "priority": "critical",
            "access_type": "direct_service",
            "funding_type": "Discretionary",
            "cfda": "15.156",
            "ci_status": "AT_RISK",
            "ci_determination": "FY26 budget uncertainty.",
            "description": "Annual awards to Tribes for climate adaptation planning.",
            "advocacy_lever": "Increase TCR funding to $50M annually",
            "tightened_language": "TCR faces flat funding despite rising demand.",
            "keywords": ["climate", "resilience"],
            "search_queries": ["tribal climate resilience"],
            "confidence_index": 0.85,
        },
        {
            "id": "fema_bric",
            "name": "FEMA BRIC",
            "agency": "FEMA",
            "federal_home": "DHS / FEMA",
            "priority": "critical",
            "access_type": "competitive",
            "funding_type": "Discretionary",
            "cfda": "97.047",
            "ci_status": "FLAGGED",
            "ci_determination": "BRIC terminated April 2025.",
            "description": "Building Resilient Infrastructure and Communities.",
            "advocacy_lever": "Restore BRIC funding and Tribal set-aside",
            "tightened_language": "BRIC program terminated; Tribal projects stranded.",
            "keywords": ["mitigation", "resilience"],
            "search_queries": ["BRIC tribal"],
            "confidence_index": 0.30,
        },
    ]


def _mock_context() -> TribePacketContext:
    """Return a TribePacketContext with full test data."""
    return TribePacketContext(
        tribe_id="epa_001",
        tribe_name="Alpha Tribe",
        states=["WA", "OR"],
        ecoregions=["pacific_northwest"],
        bia_code="100",
        epa_id="epa_001",
        districts=[
            {"district": "WA-06", "state": "WA", "overlap_pct": 80.0},
            {"district": "OR-01", "state": "OR", "overlap_pct": 20.0},
        ],
        senators=[
            {
                "bioguide_id": "S001",
                "formatted_name": "Sen. Jane Smith (D-WA)",
                "state": "WA",
                "committees": [
                    {
                        "committee_id": "SSAP",
                        "committee_name": "Senate Appropriations Committee",
                        "role": "member",
                    },
                ],
            },
        ],
        representatives=[
            {
                "bioguide_id": "R001",
                "formatted_name": "Rep. Carol Lee (D-WA-06)",
                "state": "WA",
                "district": "WA-06",
                "committees": [
                    {
                        "committee_id": "SLIA",
                        "committee_name": "Committee on Indian Affairs",
                        "role": "member",
                    },
                ],
            },
        ],
        awards=[
            {
                "award_id": "A001",
                "cfda": "15.156",
                "program_id": "bia_tcr",
                "obligation": 100000.0,
                "start_date": "2024-01-01",
                "end_date": "2025-01-01",
            },
        ],
        hazard_profile={
            "sources": {
                "fema_nri": {
                    "composite": {
                        "risk_rating": "Relatively High",
                        "risk_score": 28.5,
                        "sovi_rating": "Relatively Moderate",
                        "sovi_score": 15.2,
                    },
                    "top_hazards": [
                        {
                            "type": "Wildfire",
                            "code": "WFIR",
                            "risk_score": 45.0,
                            "risk_rating": "Very High",
                            "eal_total": 1200000,
                        },
                        {
                            "type": "Flooding",
                            "code": "RFLD",
                            "risk_score": 35.0,
                            "risk_rating": "Relatively High",
                            "eal_total": 800000,
                        },
                    ],
                },
                "usfs_wildfire": {
                    "risk_to_homes": 0.85,
                    "likelihood": 0.45,
                },
            },
        },
        generated_at="2026-02-10T12:00:00+00:00",
    )


def _mock_economic_summary() -> TribeEconomicSummary:
    """Return a TribeEconomicSummary with mock data."""
    calc = EconomicImpactCalculator()
    return calc.compute(
        tribe_id="epa_001",
        tribe_name="Alpha Tribe",
        awards=[
            {"obligation": 100000.0, "program_id": "bia_tcr", "cfda": "15.156"},
        ],
        districts=[
            {"district": "WA-06", "overlap_pct": 80.0},
            {"district": "OR-01", "overlap_pct": 20.0},
        ],
        programs={p["id"]: p for p in _mock_programs()},
    )


def _mock_structural_asks() -> list[dict]:
    """Return structural ask dicts for testing."""
    return [
        {
            "id": "ask_multi_year",
            "name": "Multi-Year Funding Stability",
            "description": "Shift from annual to multi-year authorization.",
            "target": "Congress",
            "urgency": "FY26",
            "programs": ["bia_tcr", "epa_gap"],
        },
        {
            "id": "ask_match_waivers",
            "name": "Match/Cost-Share Waivers",
            "description": "Categorical Tribal exemptions from cost-share.",
            "target": "Congress",
            "urgency": "Immediate",
            "programs": ["fema_bric"],
        },
    ]


def _make_engine(tmp_path: Path, doc_type_config=None) -> DocxEngine:
    """Create a DocxEngine with tmp_path output directory."""
    output_dir = tmp_path / "outputs" / "packets"
    config = {
        "packets": {
            "output_dir": str(output_dir),
        },
    }
    programs = {p["id"]: p for p in _mock_programs()}
    return DocxEngine(config, programs, doc_type_config=doc_type_config)


def _generate_doc(tmp_path: Path, doc_type_config=None) -> Path:
    """Generate a full document with the specified doc type config."""
    engine = _make_engine(tmp_path, doc_type_config=doc_type_config)
    context = _mock_context()
    programs = _mock_programs()
    economic = _mock_economic_summary()
    asks = _mock_structural_asks()
    return engine.generate(
        context, programs, economic, asks,
        doc_type_config=doc_type_config,
    )


def _get_all_text(path: Path) -> str:
    """Read a DOCX and return all paragraph text joined."""
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


# ===========================================================================
# Doc A tests (internal strategy)
# ===========================================================================


class TestDocAContainsStrategyContent:
    """Verify Doc A includes leverage, approach, and messaging."""

    def test_doc_a_contains_strategy_content(self, tmp_path):
        """Doc A contains strategic framing in executive summary."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_A)
        text = _get_all_text(path)

        # Strategic framing paragraph from executive summary
        assert "strategy document" in text.lower()
        assert "leverage" in text.lower() or "strategic" in text.lower()

    def test_doc_a_contains_advocacy_language(self, tmp_path):
        """Doc A Hot Sheets contain advocacy position sections."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_A)
        text = _get_all_text(path)

        # Advocacy language is included in internal docs
        assert "Advocacy Position" in text

    def test_doc_a_contains_messaging_framework(self, tmp_path):
        """Doc A includes messaging framework section."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_A)
        text = _get_all_text(path)

        assert "Messaging Framework" in text
        assert "Core Message" in text
        assert "Congressional Office Framing" in text

    def test_doc_a_contains_strategic_notes(self, tmp_path):
        """Doc A delegation section includes Strategic Notes."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_A)
        text = _get_all_text(path)

        assert "Strategic Notes" in text

    def test_doc_a_contains_hazard_alignment(self, tmp_path):
        """Doc A hazard summary includes hazard-to-program alignment."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_A)
        text = _get_all_text(path)

        assert "Hazard-to-Program Alignment" in text
        assert "strategic advocacy" in text.lower()

    def test_doc_a_structural_asks_strategic_framing(self, tmp_path):
        """Doc A structural asks use strategic framing."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_A)
        text = _get_all_text(path)

        assert "priority advocacy targets" in text.lower()


# ===========================================================================
# Doc B tests (congressional overview)
# ===========================================================================


class TestDocBExcludesStrategyContent:
    """Verify Doc B excludes leverage, approach, and messaging."""

    def test_doc_b_excludes_strategy_content(self, tmp_path):
        """Doc B does not contain strategic framing paragraph."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_B)
        text = _get_all_text(path)

        # No strategic framing paragraph
        assert "strategy document" not in text.lower()
        # No messaging framework
        assert "Messaging Framework" not in text

    def test_doc_b_excludes_advocacy_language(self, tmp_path):
        """Doc B Hot Sheets do NOT contain advocacy position sections."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_B)
        text = _get_all_text(path)

        # Advocacy language is excluded from congressional docs
        assert "Advocacy Position" not in text

    def test_doc_b_excludes_strategic_notes(self, tmp_path):
        """Doc B delegation section has NO Strategic Notes."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_B)
        text = _get_all_text(path)

        assert "Strategic Notes" not in text

    def test_doc_b_excludes_hazard_alignment_narrative(self, tmp_path):
        """Doc B hazard summary has NO hazard-to-program alignment narrative."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_B)
        text = _get_all_text(path)

        assert "Hazard-to-Program Alignment" not in text
        assert "strategic advocacy" not in text.lower()

    def test_doc_b_structural_asks_policy_framing(self, tmp_path):
        """Doc B structural asks use policy recommendation framing."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_B)
        text = _get_all_text(path)

        # Policy recommendation framing (default intro, not strategic)
        assert "would advance climate resilience funding" in text.lower()
        assert "priority advocacy targets" not in text.lower()


# ===========================================================================
# Cover page / confidentiality tests
# ===========================================================================


class TestCoverPageDifferentiation:
    """Verify cover page audience markers."""

    def test_doc_a_confidential_header(self, tmp_path):
        """Doc A cover page has CONFIDENTIAL banner."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_A)
        text = _get_all_text(path)

        assert "CONFIDENTIAL" in text

    def test_doc_b_no_confidential(self, tmp_path):
        """Doc B cover page has NO CONFIDENTIAL marking."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_B)
        text = _get_all_text(path)

        # Doc B should not have CONFIDENTIAL anywhere
        # (it appears in footers too, so check cover area)
        assert "CONFIDENTIAL -- FOR INTERNAL TRIBAL USE ONLY" not in text


# ===========================================================================
# Key Ask differentiation tests
# ===========================================================================


class TestKeyAskDifferentiation:
    """Verify Key Ask box content varies by audience."""

    def test_doc_b_has_key_ask_box(self, tmp_path):
        """Doc B Hot Sheets have Key Ask with ASK/WHY/IMPACT."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_B)
        text = _get_all_text(path)

        # Key Ask section should still exist
        assert "Key Ask" in text
        assert "ASK:" in text
        assert "WHY:" in text
        assert "IMPACT:" in text

    def test_doc_b_key_ask_uses_description_not_lever(self, tmp_path):
        """Doc B Key Ask ASK line uses program description, not advocacy lever."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_B)
        doc = DocxDocument(str(path))

        # Find the ASK: lines
        ask_lines = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.startswith("ASK:"):
                    ask_lines.append(run.text)

        # At least one ASK line should exist
        assert len(ask_lines) > 0

        # The ASK lines should use descriptions, not lever language
        # (advocacy_lever contains "Increase TCR funding" and "Restore BRIC")
        for ask_line in ask_lines:
            # Congressional docs use program description instead of advocacy lever
            assert "Increase TCR funding" not in ask_line
            assert "Restore BRIC" not in ask_line

    def test_doc_a_key_ask_uses_advocacy_lever(self, tmp_path):
        """Doc A Key Ask ASK line uses full advocacy lever text."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_A)
        text = _get_all_text(path)

        # Internal docs get full advocacy lever
        assert "Increase TCR funding" in text or "Restore BRIC" in text


# ===========================================================================
# Economic-hazard integration tests (INTG-01)
# ===========================================================================


class TestEconomicHazardIntegration:
    """Verify hazard context weaving into economic narrative."""

    def test_economic_impact_has_hazard_alignment(self, tmp_path):
        """Hot sheet economic impact includes hazard alignment paragraph."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_A)
        text = _get_all_text(path)

        # INTG-01: hazard alignment in economic section
        assert "Hazard Alignment:" in text
        assert "wildfire risk profile" in text.lower()

    def test_no_hazard_alignment_without_data(self, tmp_path):
        """No hazard alignment paragraph when hazard data is empty."""
        engine = _make_engine(tmp_path, doc_type_config=DOC_A)
        context = TribePacketContext(
            tribe_id="epa_empty",
            tribe_name="Empty Hazard Tribe",
            states=["WA"],
            districts=[
                {"district": "WA-01", "overlap_pct": 100.0},
            ],
            senators=[],
            representatives=[],
            awards=[],
            hazard_profile={},
            generated_at="2026-02-10T12:00:00+00:00",
        )
        programs = _mock_programs()
        economic = TribeEconomicSummary(
            tribe_id="epa_empty", tribe_name="Empty Hazard Tribe"
        )
        asks = _mock_structural_asks()

        path = engine.generate(context, programs, economic, asks)
        text = _get_all_text(path)

        assert "Hazard Alignment:" not in text


# ===========================================================================
# Orchestrator multi-doc tests
# ===========================================================================


class TestOrchestratorMultiDoc:
    """Tests for orchestrator generating Doc A + Doc B per Tribe."""

    def test_orchestrator_generates_two_docs_for_complete_tribe(self, tmp_path):
        """Complete-data Tribe produces 2 documents (Doc A + Doc B)."""
        from src.packets.orchestrator import PacketOrchestrator

        output_dir = tmp_path / "packets"
        config = _make_orchestrator_config(tmp_path, output_dir)

        programs = _mock_programs()
        orch = PacketOrchestrator(config, programs)

        context = _mock_context()
        tribe = {
            "tribe_id": "epa_001",
            "name": "Alpha Tribe",
            "states": ["WA", "OR"],
        }

        paths = orch.generate_tribal_docs(context, tribe)

        assert len(paths) == 2

        # Check that internal/ and congressional/ directories exist
        internal_dir = output_dir / "internal"
        congressional_dir = output_dir / "congressional"
        assert internal_dir.exists()
        assert congressional_dir.exists()

        # Check files in each directory
        internal_files = list(internal_dir.glob("*.docx"))
        congressional_files = list(congressional_dir.glob("*.docx"))
        assert len(internal_files) == 1
        assert len(congressional_files) == 1

    def test_orchestrator_generates_one_doc_for_partial_tribe(self, tmp_path):
        """Partial-data Tribe (no awards) produces Doc B only."""
        from src.packets.orchestrator import PacketOrchestrator

        output_dir = tmp_path / "packets"
        config = _make_orchestrator_config(tmp_path, output_dir)

        programs = _mock_programs()
        orch = PacketOrchestrator(config, programs)

        # Context with no awards = partial data
        context = TribePacketContext(
            tribe_id="epa_partial",
            tribe_name="Partial Tribe",
            states=["MT"],
            ecoregions=["northern_rockies"],
            districts=[
                {"district": "MT-02", "state": "MT", "overlap_pct": 80.0},
            ],
            senators=[
                {
                    "bioguide_id": "S001",
                    "formatted_name": "Sen. Test (R-MT)",
                    "state": "MT",
                    "committees": [],
                },
            ],
            representatives=[],
            awards=[],  # No awards = partial
            hazard_profile={"fema_nri": {"composite": {}, "top_hazards": []}},
            generated_at="2026-02-10T12:00:00+00:00",
        )
        tribe = {
            "tribe_id": "epa_partial",
            "name": "Partial Tribe",
            "states": ["MT"],
        }

        paths = orch.generate_tribal_docs(context, tribe)

        assert len(paths) == 1

        # Only congressional/ should have a file
        congressional_dir = output_dir / "congressional"
        assert congressional_dir.exists()
        congressional_files = list(congressional_dir.glob("*.docx"))
        assert len(congressional_files) == 1


# ===========================================================================
# Air gap compliance test
# ===========================================================================


class TestAirGapCompliance:
    """Ensure no strategy content leaks into congressional documents."""

    def test_no_air_gap_violations(self, tmp_path):
        """Doc B contains zero internal-only strategy content."""
        path = _generate_doc(tmp_path, doc_type_config=DOC_B)
        text = _get_all_text(path)
        text_lower = text.lower()

        # List of terms that must NEVER appear in Doc B
        forbidden_terms = [
            "Messaging Framework",
            "Core Message",
            "Congressional Office Framing",
            "Committee-Specific Guidance",
            "Strategic Notes",
            "Hazard-to-Program Alignment",
            "Advocacy Position",
        ]

        violations = []
        for term in forbidden_terms:
            if term in text:
                violations.append(term)

        assert not violations, (
            f"Air gap violation: Doc B contains internal-only terms: "
            f"{violations}"
        )

        # Additional check: no strategic advocacy framing
        assert "strategic advocacy" not in text_lower
        assert "priority advocacy targets" not in text_lower
        assert "strategy document" not in text_lower


# ===========================================================================
# Helpers for orchestrator tests
# ===========================================================================


def _make_orchestrator_config(tmp_path: Path, output_dir: Path) -> dict:
    """Create a config dict suitable for PacketOrchestrator.

    Sets up minimal tribal_registry.json, ecoregion_config.json,
    and congressional data directory for the orchestrator.
    """
    import json

    # Minimal tribal registry
    registry_path = tmp_path / "tribal_registry.json"
    registry_data = {
        "metadata": {"total_tribes": 1, "placeholder": False},
        "tribes": [
            {
                "tribe_id": "epa_001",
                "bia_code": "100",
                "name": "Alpha Tribe",
                "states": ["WA", "OR"],
                "alternate_names": [],
                "epa_region": "10",
                "bia_recognized": True,
            },
        ],
    }
    registry_path.parent.mkdir(parents=True, exist_ok=True)
    with open(registry_path, "w", encoding="utf-8") as f:
        json.dump(registry_data, f)

    # Minimal ecoregion config
    ecoregion_path = tmp_path / "ecoregion_config.json"
    ecoregion_data = {
        "ecoregions": {
            "pacific_northwest": {
                "name": "Pacific Northwest",
                "states": ["WA", "OR"],
                "priority_programs": ["bia_tcr"],
                "hazards": ["wildfire", "flooding"],
            },
        },
    }
    with open(ecoregion_path, "w", encoding="utf-8") as f:
        json.dump(ecoregion_data, f)

    # Congressional data dir (empty is fine for this test)
    congress_dir = tmp_path / "congressional"
    congress_dir.mkdir(parents=True, exist_ok=True)

    # Graph schema (for structural asks)
    graph_path = tmp_path / "graph_schema.json"
    graph_data = {
        "structural_asks": [
            {
                "id": "ask_1",
                "name": "Test Ask",
                "description": "Test structural ask.",
                "target": "Congress",
                "urgency": "FY26",
                "programs": ["bia_tcr"],
            },
        ],
    }
    with open(graph_path, "w", encoding="utf-8") as f:
        json.dump(graph_data, f)

    # State dir for change tracking
    state_dir = tmp_path / "packet_state"
    state_dir.mkdir(parents=True, exist_ok=True)

    return {
        "packets": {
            "output_dir": str(output_dir),
            "state_dir": str(state_dir),
            "tribal_registry": {"data_path": str(registry_path)},
            "ecoregion": {"config_path": str(ecoregion_path)},
            "congressional": {"data_dir": str(congress_dir)},
            "awards": {"cache_dir": str(tmp_path / "award_cache")},
            "hazards": {"cache_dir": str(tmp_path / "hazard_profiles")},
            "graph_schema_path": str(graph_path),
        },
    }
