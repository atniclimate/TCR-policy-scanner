"""Tests for DOCX styling foundation and engine skeleton.

Covers: StyleManager, CI_STATUS_COLORS, helper functions (set_cell_shading,
apply_zebra_stripe, add_status_badge, format_header_row), and DocxEngine
(create_document, save, header/footer, output directory creation).
"""

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.packets.docx_styles import (
    CI_STATUS_COLORS,
    CI_STATUS_LABELS,
    COLORS,
    StyleManager,
    add_status_badge,
    apply_zebra_stripe,
    format_header_row,
    set_cell_shading,
)
from src.packets.docx_engine import DocxEngine


# ---------------------------------------------------------------------------
# StyleManager tests
# ---------------------------------------------------------------------------


class TestStyleManager:
    """Tests for StyleManager style creation and idempotency."""

    def test_style_manager_creates_styles(self):
        """All expected custom style names exist after StyleManager init."""
        doc = Document()
        StyleManager(doc)
        names = {s.name for s in doc.styles}

        expected = [
            "HS Title",
            "HS Subtitle",
            "HS Section",
            "HS Body",
            "HS Small",
            "HS Callout",
            "HS Callout Detail",
        ]
        for name in expected:
            assert name in names, f"Missing style: {name}"

    def test_style_manager_idempotent(self):
        """Calling StyleManager twice on same Document does not raise ValueError."""
        doc = Document()
        StyleManager(doc)
        # Second call must not raise
        StyleManager(doc)

        # Styles still exist
        names = {s.name for s in doc.styles}
        assert "HS Title" in names

    def test_heading_styles_arial(self):
        """Built-in Heading 1-3 styles have font.name == 'Arial' after StyleManager."""
        doc = Document()
        StyleManager(doc)

        for level in range(1, 4):
            style = doc.styles[f"Heading {level}"]
            assert style.font.name == "Arial", f"Heading {level} font is not Arial"

    def test_custom_style_properties(self):
        """HS Title is 18pt bold, HS Body is 10pt, HS Small is 8pt."""
        doc = Document()
        StyleManager(doc)

        title = doc.styles["HS Title"]
        assert title.font.size == Pt(18)
        assert title.font.bold is True

        body = doc.styles["HS Body"]
        assert body.font.size == Pt(10)

        small = doc.styles["HS Small"]
        assert small.font.size == Pt(8)

    def test_hs_section_style(self):
        """HS Section is 12pt bold with space_before and space_after."""
        doc = Document()
        StyleManager(doc)

        section = doc.styles["HS Section"]
        assert section.font.size == Pt(12)
        assert section.font.bold is True
        assert section.paragraph_format.space_before == Pt(12)
        assert section.paragraph_format.space_after == Pt(4)

    def test_hs_callout_indent(self):
        """HS Callout has left_indent of 0.3 inches."""
        doc = Document()
        StyleManager(doc)

        from docx.shared import Inches

        callout = doc.styles["HS Callout"]
        assert callout.paragraph_format.left_indent == Inches(0.3)
        assert callout.font.bold is True


# ---------------------------------------------------------------------------
# Helper function tests
# ---------------------------------------------------------------------------


class TestCellShading:
    """Tests for set_cell_shading helper."""

    def test_set_cell_shading(self):
        """Cell background has correct fill color in XML."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        set_cell_shading(cell, "F2F2F2")

        tc = cell._tc
        shading_elements = tc.findall(f".//{qn('w:shd')}")
        assert len(shading_elements) >= 1
        assert shading_elements[-1].get(qn("w:fill")) == "F2F2F2"

    def test_set_cell_shading_fresh_element(self):
        """Shading two cells both get correct colors (not shared element)."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        cell_a = table.cell(0, 0)
        cell_b = table.cell(0, 1)

        set_cell_shading(cell_a, "FF0000")
        set_cell_shading(cell_b, "00FF00")

        # Both cells should have their own shading
        shading_a = cell_a._tc.findall(f".//{qn('w:shd')}")
        shading_b = cell_b._tc.findall(f".//{qn('w:shd')}")

        assert shading_a[-1].get(qn("w:fill")) == "FF0000"
        assert shading_b[-1].get(qn("w:fill")) == "00FF00"


class TestZebraStripe:
    """Tests for apply_zebra_stripe helper."""

    def test_apply_zebra_stripe(self):
        """Even/odd data rows get different shading colors."""
        doc = Document()
        table = doc.add_table(rows=4, cols=2)
        apply_zebra_stripe(table, header_rows=1)

        # Row 1 (first data row, index 0) should be even color
        row1_cell = table.cell(1, 0)
        shading = row1_cell._tc.findall(f".//{qn('w:shd')}")
        assert shading[-1].get(qn("w:fill")) == COLORS.table_stripe_even

        # Row 2 (second data row, index 1) should be odd color
        row2_cell = table.cell(2, 0)
        shading = row2_cell._tc.findall(f".//{qn('w:shd')}")
        assert shading[-1].get(qn("w:fill")) == COLORS.table_stripe_odd

        # Row 3 (third data row, index 2) should be even color again
        row3_cell = table.cell(3, 0)
        shading = row3_cell._tc.findall(f".//{qn('w:shd')}")
        assert shading[-1].get(qn("w:fill")) == COLORS.table_stripe_even

    def test_apply_zebra_stripe_skips_header(self):
        """First row (header) is not striped."""
        doc = Document()
        table = doc.add_table(rows=3, cols=1)
        apply_zebra_stripe(table, header_rows=1)

        # Header row should have NO shading elements added by zebra stripe
        header_cell = table.cell(0, 0)
        shading = header_cell._tc.findall(f".//{qn('w:shd')}")
        assert len(shading) == 0


class TestStatusBadge:
    """Tests for add_status_badge helper."""

    def test_add_status_badge_colors(self):
        """Each CI status category produces correct color."""
        doc = Document()

        test_cases = {
            "FLAGGED": RGBColor(0xDC, 0x26, 0x26),
            "AT_RISK": RGBColor(0xDC, 0x26, 0x26),
            "UNCERTAIN": RGBColor(0xB4, 0x53, 0x09),
            "STABLE_BUT_VULNERABLE": RGBColor(0xB4, 0x53, 0x09),
            "STABLE": RGBColor(0x16, 0xA3, 0x4A),
            "SECURE": RGBColor(0x16, 0xA3, 0x4A),
        }

        for status, expected_color in test_cases.items():
            para = doc.add_paragraph()
            add_status_badge(para, status)
            # Circle run is the first run added
            circle_run = para.runs[0]
            assert circle_run.font.color.rgb == expected_color, (
                f"Status {status}: expected {expected_color}, got {circle_run.font.color.rgb}"
            )

    def test_add_status_badge_all_statuses(self):
        """All 6 CI status values render without error."""
        doc = Document()
        all_statuses = list(CI_STATUS_COLORS.keys())
        assert len(all_statuses) == 6

        for status in all_statuses:
            para = doc.add_paragraph()
            add_status_badge(para, status)
            # Should have 2 runs (circle + label)
            assert len(para.runs) >= 2, f"Status {status}: expected at least 2 runs"

    def test_add_status_badge_label_text(self):
        """Status badge label text matches CI_STATUS_LABELS."""
        doc = Document()
        para = doc.add_paragraph()
        add_status_badge(para, "FLAGGED")
        label_run = para.runs[1]
        assert label_run.text == "Flagged -- Requires Specialist Attention"

    def test_add_status_badge_unknown_status(self):
        """Unknown status uses muted color fallback and raw status text."""
        doc = Document()
        para = doc.add_paragraph()
        add_status_badge(para, "UNKNOWN_STATUS")
        circle_run = para.runs[0]
        # Should use muted color fallback
        assert circle_run.font.color.rgb == RGBColor(0x6B, 0x72, 0x80)
        label_run = para.runs[1]
        assert label_run.text == "UNKNOWN_STATUS"


class TestFormatHeaderRow:
    """Tests for format_header_row helper."""

    def test_format_header_row(self):
        """Header cells have dark background and white bold text."""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        format_header_row(table, ["Col A", "Col B", "Col C"])

        for i, expected in enumerate(["Col A", "Col B", "Col C"]):
            cell = table.cell(0, i)
            # Check shading
            shading = cell._tc.findall(f".//{qn('w:shd')}")
            assert len(shading) >= 1
            assert shading[-1].get(qn("w:fill")) == COLORS.table_header_bg

            # Check text content
            text = cell.paragraphs[0].runs[0].text
            assert text == expected

            # Check font color is white
            run = cell.paragraphs[0].runs[0]
            assert run.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF)
            assert run.font.bold is True
            assert run.font.size == Pt(9)
            assert run.font.name == "Arial"


# ---------------------------------------------------------------------------
# CI_STATUS constants tests
# ---------------------------------------------------------------------------


class TestCIStatusConstants:
    """Tests for CI status color and label mappings."""

    def test_ci_status_colors_count(self):
        """CI_STATUS_COLORS has exactly 6 entries."""
        assert len(CI_STATUS_COLORS) == 6

    def test_ci_status_labels_count(self):
        """CI_STATUS_LABELS has exactly 6 entries."""
        assert len(CI_STATUS_LABELS) == 6

    def test_ci_status_keys_match(self):
        """CI_STATUS_COLORS and CI_STATUS_LABELS have the same keys."""
        assert set(CI_STATUS_COLORS.keys()) == set(CI_STATUS_LABELS.keys())


# ---------------------------------------------------------------------------
# DocxEngine tests
# ---------------------------------------------------------------------------


class TestDocxEngine:
    """Tests for DocxEngine document creation and save."""

    def _make_engine(self, tmp_path):
        """Create a DocxEngine with output_dir pointing to tmp_path."""
        config = {
            "packets": {
                "output_dir": str(tmp_path / "packets"),
            }
        }
        programs = {"PROG_01": {"id": "PROG_01", "name": "Test Program"}}
        return DocxEngine(config, programs)

    def test_engine_create_document(self, tmp_path):
        """create_document() returns (Document, StyleManager) tuple."""
        import docx.document

        engine = self._make_engine(tmp_path)
        result = engine.create_document()

        assert isinstance(result, tuple)
        assert len(result) == 2
        doc, sm = result
        assert isinstance(doc, docx.document.Document)
        assert isinstance(sm, StyleManager)

    def test_engine_save_creates_file(self, tmp_path):
        """Saving produces a .docx file on disk."""
        engine = self._make_engine(tmp_path)
        doc, _ = engine.create_document()

        path = engine.save(doc, "test_tribe")
        assert path.exists()
        assert path.suffix == ".docx"
        assert path.name == "test_tribe.docx"

    def test_engine_save_atomic(self, tmp_path):
        """File exists after save (atomic pattern works)."""
        engine = self._make_engine(tmp_path)
        doc, _ = engine.create_document()

        path = engine.save(doc, "atomic_test")

        # File should exist and be readable
        assert path.exists()
        assert path.stat().st_size > 0

        # Should be a valid DOCX (can be opened)
        reopened = Document(str(path))
        assert reopened is not None

    def test_engine_header_footer(self, tmp_path):
        """Document sections have air-gap-compliant header and footer text."""
        engine = self._make_engine(tmp_path)
        doc, _ = engine.create_document()

        section = doc.sections[0]

        # Header -- air gap clean (no tool name or org name)
        header_text = section.header.paragraphs[0].text
        from src.config import FISCAL_YEAR_SHORT
        assert FISCAL_YEAR_SHORT in header_text
        assert "Climate Resilience" in header_text
        assert "TCR Policy Scanner" not in header_text

        # Footer -- air gap clean (no tool name or org name)
        footer_text = section.footer.paragraphs[0].text
        assert "Congressional Office Use" in footer_text
        assert "TCR Policy Scanner" not in footer_text

    def test_engine_output_dir_created(self, tmp_path):
        """Output directory is created if missing."""
        nested = tmp_path / "deep" / "nested" / "packets"
        config = {
            "packets": {
                "output_dir": str(nested),
            }
        }
        programs = {}
        DocxEngine(config, programs)
        assert nested.exists()

    def test_engine_page_margins(self, tmp_path):
        """Document has 1-inch margins on all sides."""
        from docx.shared import Inches

        engine = self._make_engine(tmp_path)
        doc, _ = engine.create_document()

        section = doc.sections[0]
        assert section.top_margin == Inches(1)
        assert section.bottom_margin == Inches(1)
        assert section.left_margin == Inches(1)
        assert section.right_margin == Inches(1)

    def test_engine_generate_full(self, tmp_path):
        """generate() creates a complete document and returns a valid path."""
        from src.packets.context import TribePacketContext
        from src.packets.economic import TribeEconomicSummary

        engine = self._make_engine(tmp_path)

        context = TribePacketContext(
            tribe_id="test_123",
            tribe_name="Test Tribe",
            states=["AZ"],
            ecoregions=["southwest"],
            congress_session="119",
            generated_at="2026-01-15T00:00:00Z",
        )
        economic_summary = TribeEconomicSummary(
            tribe_id="test_123", tribe_name="Test Tribe"
        )

        path = engine.generate(context, [], economic_summary, [])
        assert path.exists()
        assert path.name == "test_123.docx"

    def test_engine_styles_available_in_document(self, tmp_path):
        """Created document has all custom styles available."""
        engine = self._make_engine(tmp_path)
        doc, _ = engine.create_document()

        names = {s.name for s in doc.styles}
        expected = ["HS Title", "HS Body", "HS Section", "HS Small"]
        for name in expected:
            assert name in names, f"Missing style in engine document: {name}"
