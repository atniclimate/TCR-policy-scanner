"""Tests for full document assembly with all 8 sections (Plan 08-01).

Tests the complete DocxEngine.generate() output with new section renderers:
executive summary, congressional delegation, hazard summary, and structural
asks. Verifies section ordering, content correctness, zero-data edge cases,
multi-state Tribes, and regression against existing tests.

All tests use tmp_path fixtures with mock data -- no network or real
data files required.
"""

import json
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from src.packets.context import TribePacketContext
from src.packets.docx_engine import DocxEngine
from src.packets.economic import (
    EconomicImpactCalculator,
    ProgramEconomicImpact,
    TribeEconomicSummary,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _write_json(path: Path, data) -> None:
    """Write data as JSON to the given path, creating parent dirs."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f)


# ---------------------------------------------------------------------------
# Mock data builders
# ---------------------------------------------------------------------------

def _mock_programs() -> list[dict]:
    """Return a list of 4 mock programs (2 critical, 2 high)."""
    return [
        {
            "id": "bia_tcr",
            "name": "BIA Tribal Climate Resilience",
            "agency": "Bureau of Indian Affairs",
            "federal_home": "DOI / BIA",
            "priority": "critical",
            "access_type": "direct_service",
            "funding_type": "Discretionary",
            "cfda": "15.156",
            "ci_status": "AT_RISK",
            "ci_determination": "FY26 budget uncertainty.",
            "description": "Annual awards to Tribes for climate adaptation planning.",
            "advocacy_lever": "Increase TCR funding to $50M annually",
            "tightened_language": "TCR faces flat funding despite rising demand.",
            "keywords": ["climate", "resilience"],
            "search_queries": ["tribal climate resilience"],
            "confidence_index": 0.85,
        },
        {
            "id": "fema_bric",
            "name": "FEMA BRIC",
            "agency": "FEMA",
            "federal_home": "DHS / FEMA",
            "priority": "critical",
            "access_type": "competitive",
            "funding_type": "Discretionary",
            "cfda": "97.047",
            "ci_status": "FLAGGED",
            "ci_determination": "BRIC terminated April 2025.",
            "description": "Building Resilient Infrastructure and Communities.",
            "advocacy_lever": "Restore BRIC funding and Tribal set-aside",
            "tightened_language": "BRIC program terminated; Tribal projects stranded.",
            "keywords": ["mitigation", "resilience"],
            "search_queries": ["BRIC tribal"],
            "confidence_index": 0.30,
        },
        {
            "id": "epa_gap",
            "name": "EPA Indian Environmental GAP",
            "agency": "EPA",
            "federal_home": "EPA / OEJ",
            "priority": "high",
            "access_type": "direct_service",
            "funding_type": "Discretionary",
            "cfda": "66.926",
            "ci_status": "STABLE",
            "ci_determination": "Permanent authorization.",
            "description": "General Assistance Program for Tribal environmental capacity.",
            "advocacy_lever": "Expand GAP funding for climate monitoring",
            "keywords": ["environmental", "capacity"],
            "search_queries": ["EPA GAP tribal"],
            "confidence_index": 0.70,
        },
        {
            "id": "doe_indian_energy",
            "name": "DOE Indian Energy",
            "agency": "Department of Energy",
            "federal_home": "DOE / OIE",
            "priority": "high",
            "access_type": "competitive",
            "funding_type": "Discretionary",
            "cfda": "81.104",
            "ci_status": "UNCERTAIN",
            "ci_determination": "IIJA supplement timing unclear.",
            "description": "Energy development and efficiency for Tribal communities.",
            "advocacy_lever": "Streamline Indian energy project approvals",
            "keywords": ["energy", "tribal"],
            "search_queries": ["DOE indian energy"],
            "confidence_index": 0.55,
        },
    ]


def _mock_context() -> TribePacketContext:
    """Return a TribePacketContext with full test data."""
    return TribePacketContext(
        tribe_id="epa_001",
        tribe_name="Alpha Tribe",
        states=["WA", "OR"],
        ecoregions=["pacific_northwest"],
        bia_code="100",
        epa_id="epa_001",
        districts=[
            {"district": "WA-06", "state": "WA", "overlap_pct": 80.0},
            {"district": "OR-01", "state": "OR", "overlap_pct": 20.0},
        ],
        senators=[
            {
                "bioguide_id": "S001",
                "formatted_name": "Sen. Jane Smith (D-WA)",
                "state": "WA",
                "committees": [
                    {
                        "committee_id": "SSAP",
                        "committee_name": "Senate Appropriations Committee",
                        "role": "member",
                    },
                    {
                        "committee_id": "SSEG",
                        "committee_name": "Senate Energy and Natural Resources Committee",
                        "role": "member",
                    },
                ],
            },
            {
                "bioguide_id": "S002",
                "formatted_name": "Sen. Bob Jones (R-OR)",
                "state": "OR",
                "committees": [
                    {
                        "committee_id": "SSFI",
                        "committee_name": "Senate Finance Committee",
                        "role": "member",
                    },
                ],
            },
        ],
        representatives=[
            {
                "bioguide_id": "R001",
                "formatted_name": "Rep. Carol Lee (D-WA-06)",
                "state": "WA",
                "district": "WA-06",
                "committees": [
                    {
                        "committee_id": "SLIA",
                        "committee_name": "Committee on Indian Affairs",
                        "role": "member",
                    },
                ],
            },
            {
                "bioguide_id": "R002",
                "formatted_name": "Rep. Dave Kim (R-OR-01)",
                "state": "OR",
                "district": "OR-01",
                "committees": [
                    {
                        "committee_id": "HSNR",
                        "committee_name": "House Natural Resources Committee",
                        "role": "member",
                    },
                ],
            },
        ],
        awards=[
            {
                "award_id": "A001",
                "cfda": "15.156",
                "program_id": "bia_tcr",
                "obligation": 100000.0,
                "start_date": "2024-01-01",
                "end_date": "2025-01-01",
            },
            {
                "award_id": "A002",
                "cfda": "15.156",
                "program_id": "bia_tcr",
                "obligation": 200000.0,
                "start_date": "2023-01-01",
                "end_date": "2024-01-01",
            },
            {
                "award_id": "A003",
                "cfda": "15.156",
                "program_id": "bia_tcr",
                "obligation": 50000.0,
                "start_date": "2022-01-01",
                "end_date": "2023-01-01",
            },
        ],
        hazard_profile={
            "sources": {
                "fema_nri": {
                    "composite": {
                        "risk_rating": "Relatively High",
                        "risk_score": 28.5,
                        "sovi_rating": "Relatively Moderate",
                        "sovi_score": 15.2,
                    },
                    "top_hazards": [
                        {
                            "type": "Wildfire",
                            "code": "WFIR",
                            "risk_score": 45.0,
                            "risk_rating": "Very High",
                            "eal_total": 1200000,
                        },
                        {
                            "type": "Flooding",
                            "code": "RFLD",
                            "risk_score": 35.0,
                            "risk_rating": "Relatively High",
                            "eal_total": 800000,
                        },
                        {
                            "type": "Earthquake",
                            "code": "ERQK",
                            "risk_score": 25.0,
                            "risk_rating": "Relatively Moderate",
                            "eal_total": 500000,
                        },
                        {
                            "type": "Drought",
                            "code": "DRGT",
                            "risk_score": 20.0,
                            "risk_rating": "Relatively Moderate",
                            "eal_total": 300000,
                        },
                        {
                            "type": "Hurricane",
                            "code": "HRCN",
                            "risk_score": 10.0,
                            "risk_rating": "Relatively Low",
                            "eal_total": 100000,
                        },
                    ],
                },
                "usfs_wildfire": {
                    "risk_to_homes": 0.85,
                    "likelihood": 0.45,
                },
            },
        },
        generated_at="2026-02-10T12:00:00+00:00",
    )


def _mock_economic_summary() -> TribeEconomicSummary:
    """Return a TribeEconomicSummary with mock data."""
    calc = EconomicImpactCalculator()
    return calc.compute(
        tribe_id="epa_001",
        tribe_name="Alpha Tribe",
        awards=[
            {"obligation": 100000.0, "program_id": "bia_tcr", "cfda": "15.156"},
            {"obligation": 200000.0, "program_id": "bia_tcr", "cfda": "15.156"},
            {"obligation": 50000.0, "program_id": "bia_tcr", "cfda": "15.156"},
        ],
        districts=[
            {"district": "WA-06", "overlap_pct": 80.0},
            {"district": "OR-01", "overlap_pct": 20.0},
        ],
        programs={p["id"]: p for p in _mock_programs()},
    )


def _mock_structural_asks() -> list[dict]:
    """Return 5 structural ask dicts for testing."""
    return [
        {
            "id": "ask_multi_year",
            "name": "Multi-Year Funding Stability",
            "description": "Shift from annual to multi-year authorization.",
            "target": "Congress",
            "urgency": "FY26",
            "programs": ["bia_tcr", "epa_gap", "doe_indian_energy"],
        },
        {
            "id": "ask_match_waivers",
            "name": "Match/Cost-Share Waivers",
            "description": "Categorical Tribal exemptions from cost-share.",
            "target": "Congress",
            "urgency": "Immediate",
            "programs": ["fema_bric"],
        },
        {
            "id": "ask_set_asides",
            "name": "Tribal Set-Aside Expansion",
            "description": "Increase Tribal set-aside percentages in competitive programs.",
            "target": "Agency rulemaking",
            "urgency": "FY26",
            "programs": ["fema_bric", "doe_indian_energy"],
        },
        {
            "id": "ask_capacity_building",
            "name": "Capacity Building Grants",
            "description": "Dedicated pre-application technical assistance funding.",
            "target": "Congress",
            "urgency": "HIGH",
            "programs": ["bia_tcr", "epa_gap"],
        },
        {
            "id": "ask_data_sovereignty",
            "name": "Tribal Data Sovereignty",
            "description": "Protect Tribal rights over environmental and climate data.",
            "target": "Executive Order",
            "urgency": "FY26",
            "programs": ["epa_gap"],
        },
    ]


def _zero_data_context() -> TribePacketContext:
    """Return a TribePacketContext with no data (empty Tribe)."""
    return TribePacketContext(
        tribe_id="epa_999",
        tribe_name="Empty Tribe",
        states=[],
        ecoregions=[],
        districts=[],
        senators=[],
        representatives=[],
        awards=[],
        hazard_profile={},
        generated_at="2026-02-10T12:00:00+00:00",
    )


def _zero_economic_summary() -> TribeEconomicSummary:
    """Return a zero-data TribeEconomicSummary."""
    return TribeEconomicSummary(
        tribe_id="epa_999",
        tribe_name="Empty Tribe",
    )


def _make_engine(tmp_path: Path) -> DocxEngine:
    """Create a DocxEngine with tmp_path output directory."""
    output_dir = tmp_path / "outputs" / "packets"
    config = {
        "packets": {
            "output_dir": str(output_dir),
        },
    }
    programs = {p["id"]: p for p in _mock_programs()}
    return DocxEngine(config, programs)


def _generate_full_doc(tmp_path: Path) -> Path:
    """Generate a full document with all sections and return the path."""
    engine = _make_engine(tmp_path)
    context = _mock_context()
    programs = _mock_programs()
    economic = _mock_economic_summary()
    asks = _mock_structural_asks()
    return engine.generate(context, programs, economic, asks)


# ===========================================================================
# Test classes
# ===========================================================================


class TestFullAssembly:
    """Tests for full 8-section document assembly."""

    def test_full_assembly_produces_all_sections(self, tmp_path):
        """Generate full document and verify all section headings present."""
        path = _generate_full_doc(tmp_path)

        doc = DocxDocument(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Executive Summary" in all_text
        assert "Congressional Delegation" in all_text
        assert "Climate Hazard Profile" in all_text
        assert "Structural Policy Asks" in all_text
        assert "Appendix" in all_text

    def test_document_opens_without_error(self, tmp_path):
        """Generated document opens correctly and has non-zero size."""
        path = _generate_full_doc(tmp_path)

        assert path.exists()
        assert path.stat().st_size > 0

        # Must not raise
        doc = DocxDocument(str(path))
        assert len(doc.paragraphs) > 0


class TestExecutiveSummary:
    """Tests for executive summary section content."""

    def test_executive_summary_contains_tribe_name(self, tmp_path):
        """Executive summary area contains the Tribe name."""
        path = _generate_full_doc(tmp_path)
        doc = DocxDocument(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Alpha Tribe" in all_text

    def test_executive_summary_references_programs(self, tmp_path):
        """Executive summary references at least 1 program name."""
        path = _generate_full_doc(tmp_path)
        doc = DocxDocument(str(path))

        # Find paragraphs between "Executive Summary" heading and next heading
        paras = [p.text for p in doc.paragraphs]
        in_exec = False
        exec_text = []
        for text in paras:
            if text == "Executive Summary":
                in_exec = True
                continue
            if in_exec:
                # Stop at next Heading 2 level section
                if text in (
                    "Congressional Delegation",
                    "Climate Hazard Profile",
                    "Structural Policy Asks",
                    "Appendix: Additional Programs",
                    "Program Hot Sheets",
                ):
                    break
                exec_text.append(text)

        exec_combined = "\n".join(exec_text)
        # At least one of the first 3 programs should be referenced
        programs = _mock_programs()
        found = any(
            p.get("name", "") in exec_combined for p in programs[:3]
        )
        assert found, (
            f"Expected at least one program name in executive summary. "
            f"Got: {exec_combined[:300]}"
        )


class TestDelegationSection:
    """Tests for congressional delegation section."""

    def test_delegation_table_has_senators_and_reps(self, tmp_path):
        """Delegation table contains both Senator and Representative rows."""
        path = _generate_full_doc(tmp_path)
        doc = DocxDocument(str(path))

        # Find tables with Name/Role headers
        found_delegation_table = False
        has_senator = False
        has_rep = False

        for table in doc.tables:
            header_cells = [cell.text for cell in table.rows[0].cells]
            if "Name" in header_cells and "Role" in header_cells:
                found_delegation_table = True
                role_idx = header_cells.index("Role")
                for row in table.rows[1:]:
                    role = row.cells[role_idx].text
                    if role == "Senator":
                        has_senator = True
                    elif role == "Representative":
                        has_rep = True
                break

        assert found_delegation_table, "No delegation table found with Name/Role headers"
        assert has_senator, "No Senator found in delegation table"
        assert has_rep, "No Representative found in delegation table"

    def test_delegation_committee_filtering(self, tmp_path):
        """Delegation section reports members with relevant committee assignments."""
        path = _generate_full_doc(tmp_path)
        doc = DocxDocument(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Mock data has matching committees (Appropriations, Energy, Indian Affairs, Natural Resources)
        assert "relevant committee assignments" in all_text.lower()

        # Extract the count
        for para in doc.paragraphs:
            if "members with relevant committee assignments" in para.text:
                # Should be > 0
                count_str = para.text.split()[0]
                assert count_str.isdigit()
                assert int(count_str) > 0
                break
        else:
            pytest.fail("No paragraph with committee assignment count found")

    def test_multi_state_tribe_delegation(self, tmp_path):
        """Multi-state Tribe with 4 senators and 3 reps shows all 7 members."""
        engine = _make_engine(tmp_path)

        context = TribePacketContext(
            tribe_id="epa_multi",
            tribe_name="Multi State Tribe",
            states=["WA", "OR", "ID"],
            ecoregions=["pacific_northwest"],
            districts=[
                {"district": "WA-06", "state": "WA", "overlap_pct": 50.0},
                {"district": "OR-01", "state": "OR", "overlap_pct": 30.0},
                {"district": "ID-02", "state": "ID", "overlap_pct": 20.0},
            ],
            senators=[
                {"bioguide_id": "S001", "formatted_name": "Sen. A (D-WA)", "state": "WA", "committees": []},
                {"bioguide_id": "S002", "formatted_name": "Sen. B (R-WA)", "state": "WA", "committees": []},
                {"bioguide_id": "S003", "formatted_name": "Sen. C (D-OR)", "state": "OR", "committees": []},
                {"bioguide_id": "S004", "formatted_name": "Sen. D (R-ID)", "state": "ID", "committees": []},
            ],
            representatives=[
                {"bioguide_id": "R001", "formatted_name": "Rep. X (D-WA-06)", "state": "WA", "district": "WA-06", "committees": []},
                {"bioguide_id": "R002", "formatted_name": "Rep. Y (R-OR-01)", "state": "OR", "district": "OR-01", "committees": []},
                {"bioguide_id": "R003", "formatted_name": "Rep. Z (R-ID-02)", "state": "ID", "district": "ID-02", "committees": []},
            ],
            awards=[],
            hazard_profile={},
            generated_at="2026-02-10T12:00:00+00:00",
        )

        programs = _mock_programs()
        economic = _zero_economic_summary()
        economic.tribe_id = "epa_multi"
        economic.tribe_name = "Multi State Tribe"
        asks = _mock_structural_asks()

        path = engine.generate(context, programs, economic, asks)
        doc = DocxDocument(str(path))

        # Find delegation table and count data rows
        for table in doc.tables:
            header_cells = [cell.text for cell in table.rows[0].cells]
            if "Name" in header_cells and "Role" in header_cells:
                # 7 data rows (4 senators + 3 reps)
                data_rows = len(table.rows) - 1  # subtract header row
                assert data_rows == 7, (
                    f"Expected 7 members in delegation table, got {data_rows}"
                )
                break
        else:
            pytest.fail("No delegation table found")


class TestHazardSummary:
    """Tests for hazard summary section content."""

    def test_hazard_summary_renders_top_hazards(self, tmp_path):
        """Hazard summary contains top hazard type names."""
        path = _generate_full_doc(tmp_path)
        doc = DocxDocument(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Wildfire is the top hazard in mock data
        assert "Wildfire" in all_text

    def test_hazard_summary_has_fema_ratings(self, tmp_path):
        """Hazard summary contains FEMA NRI composite ratings."""
        path = _generate_full_doc(tmp_path)
        doc = DocxDocument(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Check for rating text
        assert "Risk Rating" in all_text or "Social Vulnerability" in all_text


class TestStructuralAsks:
    """Tests for structural asks section content."""

    def test_structural_asks_includes_all_five(self, tmp_path):
        """Structural asks section contains all 5 ask names."""
        path = _generate_full_doc(tmp_path)
        doc = DocxDocument(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        asks = _mock_structural_asks()
        for ask in asks:
            assert ask["name"] in all_text, (
                f"Ask '{ask['name']}' not found in document"
            )


class TestZeroDataTribe:
    """Tests for zero-data edge case (no awards, no hazards, no delegation)."""

    def test_zero_data_tribe_produces_valid_docx(self, tmp_path):
        """Zero-data Tribe generates a valid DOCX without crashing."""
        engine = _make_engine(tmp_path)
        context = _zero_data_context()
        programs = _mock_programs()
        economic = _zero_economic_summary()
        asks = _mock_structural_asks()

        path = engine.generate(context, programs, economic, asks)

        # Must open without error
        doc = DocxDocument(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Graceful fallback messages
        assert (
            "Hazard profile data not available" in all_text
            or "Hazard profile data pending" in all_text
        )
        assert "Congressional delegation data not available" in all_text


class TestSectionOrder:
    """Tests for correct section ordering in generated document."""

    def test_section_order_correct(self, tmp_path):
        """Heading 2 paragraphs appear in expected order."""
        path = _generate_full_doc(tmp_path)
        doc = DocxDocument(str(path))

        # Collect all Heading 2 paragraphs in order
        h2_texts = []
        for para in doc.paragraphs:
            if para.style and para.style.name == "Heading 2":
                h2_texts.append(para.text)

        # Expected sequence (Hot Sheet titles between delegation and hazard)
        expected_markers = [
            "Program Hot Sheets",
            "Executive Summary",
            "Congressional Delegation",
            # Hot Sheet program titles would be here (HS Title style, not Heading 2)
            "Climate Hazard Profile",
            "Structural Policy Asks",
            "Appendix: Additional Programs",
        ]

        # Verify each expected marker appears in order
        last_idx = -1
        for marker in expected_markers:
            found = False
            for i, h2 in enumerate(h2_texts):
                if marker in h2 and i > last_idx:
                    last_idx = i
                    found = True
                    break
            assert found, (
                f"Section '{marker}' not found in correct order. "
                f"H2 headings: {h2_texts}"
            )


class TestRegression:
    """Regression tests to ensure existing tests are not broken."""

    def test_existing_tests_unaffected(self):
        """All existing test modules pass without regression."""
        from src.paths import PROJECT_ROOT
        result = subprocess.run(
            [
                sys.executable, "-m", "pytest",
                "tests/test_docx_integration.py",
                "tests/test_docx_hotsheet.py",
                "tests/test_docx_styles.py",
                "tests/test_economic.py",
                "-v", "--tb=short",
            ],
            capture_output=True,
            text=True,
            cwd=str(PROJECT_ROOT),
            timeout=120,
        )

        assert result.returncode == 0, (
            f"Existing tests failed:\n{result.stdout}\n{result.stderr}"
        )
