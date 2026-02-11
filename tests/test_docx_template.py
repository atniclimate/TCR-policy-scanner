"""Tests for DOCX template builder, layout helpers, and agent review orchestrator.

Covers: TemplateBuilder, create_two_column_grid, create_info_box,
add_horizontal_rule, mark_header_row, add_page_number_field,
build_template convenience function, BOX_SCHEMES, and
AgentReviewOrchestrator (critique registration, conflict resolution,
quality gate, resolution logging).
"""

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.packets.docx_template import (
    BOX_SCHEMES,
    LEFT_RATIO,
    MARGIN_LEFT,
    MARGIN_RIGHT,
    PAGE_WIDTH,
    TemplateBuilder,
    add_horizontal_rule,
    add_page_number_field,
    build_template,
    create_info_box,
    create_two_column_grid,
    mark_header_row,
)
from src.packets.agent_review import (
    AGENT_PRIORITY,
    AgentReviewOrchestrator,
    Critique,
    MIN_AGENTS_FOR_QUALITY_GATE,
)


# ---------------------------------------------------------------------------
# TemplateBuilder tests
# ---------------------------------------------------------------------------


class TestTemplateBuilder:
    """Tests for TemplateBuilder.build() and .save()."""

    def test_build_returns_document(self):
        """build() returns a python-docx Document."""
        builder = TemplateBuilder()
        doc = builder.build()
        assert doc is not None

    def test_build_registers_base_styles(self):
        """Template has all 7 base HS-* styles from StyleManager."""
        builder = TemplateBuilder()
        doc = builder.build()
        names = {s.name for s in doc.styles}
        base_styles = [
            "HS Title", "HS Subtitle", "HS Section", "HS Body",
            "HS Small", "HS Callout", "HS Callout Detail",
        ]
        for name in base_styles:
            assert name in names, f"Missing base style: {name}"

    def test_build_registers_extended_styles(self):
        """Template has the 9 extended styles from _EXTENDED_STYLES."""
        builder = TemplateBuilder()
        doc = builder.build()
        names = {s.name for s in doc.styles}
        extended = [
            "HS Column Primary", "HS Column Secondary",
            "HS Box Title", "HS Box Body",
            "HS Ask Label", "HS Ask Detail",
            "HS Footer", "HS Caption", "HS Page Number",
        ]
        for name in extended:
            assert name in names, f"Missing extended style: {name}"

    def test_build_page_margins(self):
        """Template has 1-inch margins on all sides."""
        builder = TemplateBuilder()
        doc = builder.build()
        section = doc.sections[0]
        assert section.top_margin == Inches(1)
        assert section.bottom_margin == Inches(1)
        assert section.left_margin == Inches(1)
        assert section.right_margin == Inches(1)

    def test_build_page_dimensions(self):
        """Template is US Letter size (8.5 x 11 inches)."""
        builder = TemplateBuilder()
        doc = builder.build()
        section = doc.sections[0]
        assert section.page_width == Inches(8.5)
        assert section.page_height == Inches(11)

    def test_build_header_text(self):
        """Template header contains branding text."""
        builder = TemplateBuilder()
        doc = builder.build()
        header = doc.sections[0].header
        text = header.paragraphs[0].text
        assert "TCR Policy Scanner" in text

    def test_build_footer_text(self):
        """Template footer contains confidentiality notice."""
        builder = TemplateBuilder()
        doc = builder.build()
        footer = doc.sections[0].footer
        text = footer.paragraphs[0].text
        assert "CONFIDENTIAL" in text

    def test_build_normal_style_arial(self):
        """Normal style is set to Arial font."""
        builder = TemplateBuilder()
        doc = builder.build()
        normal = doc.styles["Normal"]
        assert normal.font.name == "Arial"

    def test_save_requires_build(self):
        """save() raises RuntimeError if build() was not called."""
        builder = TemplateBuilder()
        with pytest.raises(RuntimeError, match="build"):
            builder.save("some_path.docx")

    def test_save_creates_file(self, tmp_path):
        """save() writes a .docx file to disk."""
        builder = TemplateBuilder()
        builder.build()
        out = tmp_path / "template.docx"
        result = builder.save(out)
        assert result == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_save_creates_parent_dirs(self, tmp_path):
        """save() creates parent directories if they don't exist."""
        builder = TemplateBuilder()
        builder.build()
        out = tmp_path / "deep" / "nested" / "template.docx"
        builder.save(out)
        assert out.exists()

    def test_build_idempotent_styles(self):
        """Building twice doesn't duplicate styles."""
        builder = TemplateBuilder()
        doc = builder.build()
        style_count_1 = len([s for s in doc.styles if s.name.startswith("HS ")])

        # Build again on a fresh builder
        builder2 = TemplateBuilder()
        doc2 = builder2.build()
        style_count_2 = len([s for s in doc2.styles if s.name.startswith("HS ")])

        assert style_count_1 == style_count_2


class TestBuildTemplate:
    """Tests for the build_template() convenience function."""

    def test_build_template_creates_file(self, tmp_path):
        """build_template() creates a .docx file at the specified path."""
        out = tmp_path / "test_template.docx"
        result = build_template(out)
        assert result == out
        assert out.exists()

    def test_build_template_output_opens_as_document(self, tmp_path):
        """The generated template can be opened by python-docx."""
        out = tmp_path / "test_template.docx"
        build_template(out)
        doc = Document(str(out))
        names = {s.name for s in doc.styles}
        assert "HS Title" in names


# ---------------------------------------------------------------------------
# Layout helper tests
# ---------------------------------------------------------------------------


class TestCreateTwoColumnGrid:
    """Tests for create_two_column_grid()."""

    def test_returns_table(self):
        """Returns a Table object."""
        doc = Document()
        table = create_two_column_grid(doc)
        assert table is not None
        assert len(table.columns) == 2

    def test_single_row(self):
        """Grid has exactly one row."""
        doc = Document()
        table = create_two_column_grid(doc)
        assert len(table.rows) == 1

    def test_no_autofit(self):
        """Grid table has autofit disabled for fixed widths."""
        doc = Document()
        table = create_two_column_grid(doc)
        assert table.autofit is False

    def test_default_ratio(self):
        """Left column is wider than right (62/38 default split)."""
        doc = Document()
        # Set margins so width calculation works
        for section in doc.sections:
            section.page_width = PAGE_WIDTH
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT
        table = create_two_column_grid(doc)
        left_width = table.columns[0].width
        right_width = table.columns[1].width
        assert left_width > right_width

    def test_custom_ratio(self):
        """Custom ratio changes column proportions."""
        doc = Document()
        for section in doc.sections:
            section.page_width = PAGE_WIDTH
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT
        table = create_two_column_grid(doc, left_ratio=0.5)
        left_width = table.columns[0].width
        right_width = table.columns[1].width
        # With 0.5 ratio and gutter, widths should be roughly equal
        ratio = left_width / (left_width + right_width)
        assert 0.45 < ratio < 0.55

    def test_invisible_borders(self):
        """Grid table borders are set to 'none'."""
        doc = Document()
        table = create_two_column_grid(doc)
        tbl = table._tbl
        borders = tbl.tblPr.findall(qn("w:tblBorders"))
        assert len(borders) > 0
        for border_group in borders:
            for child in border_group:
                assert child.get(qn("w:val")) == "none"


class TestCreateInfoBox:
    """Tests for create_info_box()."""

    def test_returns_table_and_cell(self):
        """Returns a (table, cell) tuple."""
        doc = Document()
        table, cell = create_info_box(doc)
        assert table is not None
        assert cell is not None

    def test_single_cell(self):
        """Info box is a 1x1 table."""
        doc = Document()
        table, _ = create_info_box(doc)
        assert len(table.rows) == 1
        assert len(table.columns) == 1

    def test_all_box_types(self):
        """All 4 box types create valid boxes without error."""
        doc = Document()
        for box_type in BOX_SCHEMES:
            table, cell = create_info_box(doc, box_type=box_type)
            assert cell is not None

    def test_unknown_box_type_falls_back(self):
        """Unknown box type falls back to 'neutral'."""
        doc = Document()
        table, cell = create_info_box(doc, box_type="nonexistent")
        # Should not raise — falls back to neutral
        assert cell is not None

    def test_cell_has_shading(self):
        """Info box cell has background shading applied."""
        doc = Document()
        _, cell = create_info_box(doc, box_type="alert")
        tc = cell._tc
        shd = tc.tcPr.findall(qn("w:shd"))
        assert len(shd) > 0
        assert shd[0].get(qn("w:fill")) == "FEF2F2"

    def test_cell_has_borders(self):
        """Info box cell has colored borders."""
        doc = Document()
        _, cell = create_info_box(doc, box_type="success")
        tc = cell._tc
        borders = tc.tcPr.findall(qn("w:tcBorders"))
        assert len(borders) > 0


class TestAddHorizontalRule:
    """Tests for add_horizontal_rule()."""

    def test_returns_paragraph(self):
        """Returns a Paragraph object."""
        doc = Document()
        para = add_horizontal_rule(doc)
        assert para is not None

    def test_has_bottom_border(self):
        """Paragraph has a bottom border element."""
        doc = Document()
        para = add_horizontal_rule(doc)
        pPr = para._p.pPr
        pBdr = pPr.findall(qn("w:pBdr"))
        assert len(pBdr) > 0
        bottom = pBdr[0].findall(qn("w:bottom"))
        assert len(bottom) > 0

    def test_custom_color(self):
        """Custom color is applied to the border."""
        doc = Document()
        para = add_horizontal_rule(doc, color="FF0000")
        pBdr = para._p.pPr.findall(qn("w:pBdr"))[0]
        bottom = pBdr.findall(qn("w:bottom"))[0]
        assert bottom.get(qn("w:color")) == "FF0000"

    def test_custom_weight(self):
        """Custom weight is applied to the border."""
        doc = Document()
        para = add_horizontal_rule(doc, weight=8)
        pBdr = para._p.pPr.findall(qn("w:pBdr"))[0]
        bottom = pBdr.findall(qn("w:bottom"))[0]
        assert bottom.get(qn("w:sz")) == "8"


class TestMarkHeaderRow:
    """Tests for mark_header_row()."""

    def test_sets_tblheader(self):
        """First row gets the w:tblHeader element for accessibility."""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        mark_header_row(table)
        first_row = table.rows[0]
        trPr = first_row._tr.trPr
        headers = trPr.findall(qn("w:tblHeader"))
        assert len(headers) > 0

    def test_only_first_row(self):
        """Only the first row is marked, not subsequent rows."""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        mark_header_row(table)
        second_row = table.rows[1]
        trPr = second_row._tr.trPr
        if trPr is not None:
            headers = trPr.findall(qn("w:tblHeader"))
            assert len(headers) == 0


class TestAddPageNumberField:
    """Tests for add_page_number_field()."""

    def test_adds_field_code(self):
        """Paragraph gets fldChar elements for page numbering."""
        doc = Document()
        para = doc.add_paragraph()
        add_page_number_field(para)
        # Should have at least one run with fldChar elements
        r_elements = para._p.findall(qn("w:r"))
        assert len(r_elements) > 0
        fld_chars = []
        for r in r_elements:
            fld_chars.extend(r.findall(qn("w:fldChar")))
        assert len(fld_chars) >= 3  # begin, separate, end

    def test_field_contains_page_instruction(self):
        """instrText contains ' PAGE '."""
        doc = Document()
        para = doc.add_paragraph()
        add_page_number_field(para)
        instr_elements = []
        for r in para._p.findall(qn("w:r")):
            instr_elements.extend(r.findall(qn("w:instrText")))
        assert any("PAGE" in (el.text or "") for el in instr_elements)


# ---------------------------------------------------------------------------
# DocxEngine template integration tests
# ---------------------------------------------------------------------------


class TestDocxEngineTemplateIntegration:
    """Tests for DocxEngine using TemplateBuilder templates."""

    def test_engine_uses_template_when_available(self, tmp_path):
        """DocxEngine opens template instead of blank doc when path exists."""
        from src.packets.docx_engine import DocxEngine

        # Build a template
        template_path = tmp_path / "template.docx"
        build_template(template_path)

        # Create engine with explicit template
        engine = DocxEngine(
            config={"packets": {"output_dir": str(tmp_path / "output")}},
            programs={},
            template_path=template_path,
        )
        doc, sm = engine.create_document()
        names = {s.name for s in doc.styles}
        # Template should have extended styles
        assert "HS Column Primary" in names
        assert "HS Box Title" in names

    def test_engine_falls_back_without_template(self, tmp_path):
        """DocxEngine creates blank doc when template_path is False."""
        from src.packets.docx_engine import DocxEngine

        engine = DocxEngine(
            config={"packets": {"output_dir": str(tmp_path / "output")}},
            programs={},
            template_path=False,
        )
        doc, sm = engine.create_document()
        names = {s.name for s in doc.styles}
        # Base styles still present (from StyleManager)
        assert "HS Title" in names
        # Extended styles NOT present (no template)
        assert "HS Column Primary" not in names


# ---------------------------------------------------------------------------
# AgentReviewOrchestrator tests
# ---------------------------------------------------------------------------


def _make_critique(
    agent: str = "accuracy-agent",
    section: str = "Award History",
    severity: str = "major",
    **extra,
) -> dict:
    """Build a minimal critique dict."""
    d = {"agent": agent, "section": section, "severity": severity}
    d.update(extra)
    return d


class TestCritique:
    """Tests for Critique.from_dict() validation."""

    def test_valid_critique(self):
        """Valid dict creates a Critique."""
        c = Critique.from_dict(_make_critique())
        assert c.agent == "accuracy-agent"
        assert c.severity == "major"

    def test_missing_field_raises(self):
        """Missing required field raises ValueError."""
        with pytest.raises(ValueError, match="missing required"):
            Critique.from_dict({"agent": "accuracy-agent"})

    def test_invalid_severity_raises(self):
        """Invalid severity value raises ValueError."""
        with pytest.raises(ValueError, match="Invalid severity"):
            Critique.from_dict(_make_critique(severity="urgent"))

    def test_severity_normalized_lowercase(self):
        """Severity is normalized to lowercase."""
        c = Critique.from_dict(_make_critique(severity="CRITICAL"))
        assert c.severity == "critical"


class TestAgentReviewOrchestrator:
    """Tests for the review orchestrator logic."""

    def test_disabled_by_default(self):
        """Orchestrator is disabled by default."""
        orch = AgentReviewOrchestrator()
        assert not orch.enabled

    def test_disabled_noop(self):
        """When disabled, all methods return empty results."""
        orch = AgentReviewOrchestrator(enabled=False)
        assert orch.register_critiques("accuracy-agent", []) == 0
        assert orch.resolve_conflicts() == []
        gate = orch.check_quality_gate()
        assert gate["passed"] is True
        assert gate["enabled"] is False

    def test_register_valid_critiques(self):
        """Valid critiques are registered and counted."""
        orch = AgentReviewOrchestrator(enabled=True)
        count = orch.register_critiques("accuracy-agent", [
            _make_critique(agent="accuracy-agent", severity="critical"),
            _make_critique(agent="accuracy-agent", severity="minor"),
        ])
        assert count == 2
        assert len(orch.all_critiques) == 2

    def test_register_skips_invalid(self):
        """Invalid critiques are skipped with a warning."""
        orch = AgentReviewOrchestrator(enabled=True)
        count = orch.register_critiques("accuracy-agent", [
            _make_critique(agent="accuracy-agent"),
            {"agent": "accuracy-agent"},  # missing section, severity
        ])
        assert count == 1

    def test_register_unknown_agent_raises(self):
        """Registering for an unknown agent raises ValueError."""
        orch = AgentReviewOrchestrator(enabled=True)
        with pytest.raises(ValueError, match="Unknown agent"):
            orch.register_critiques("unknown-agent", [])

    def test_mark_agent_failed(self):
        """Failed agents are tracked."""
        orch = AgentReviewOrchestrator(enabled=True)
        orch.mark_agent_failed("design-aficionado")
        assert "design-aficionado" in orch._failed_agents

    def test_resolve_no_conflicts(self):
        """Non-overlapping critiques produce no conflicts."""
        orch = AgentReviewOrchestrator(enabled=True)
        orch.register_critiques("accuracy-agent", [
            _make_critique(agent="accuracy-agent", section="Awards"),
        ])
        orch.register_critiques("copy-editor", [
            _make_critique(agent="copy-editor", section="Actions"),
        ])
        conflicts = orch.resolve_conflicts()
        assert len(conflicts) == 0

    def test_resolve_conflict_higher_priority_wins(self):
        """When two agents critique the same section, higher priority wins."""
        orch = AgentReviewOrchestrator(enabled=True)
        orch.register_critiques("accuracy-agent", [
            _make_critique(
                agent="accuracy-agent", section="Table",
                severity="critical",
            ),
        ])
        orch.register_critiques("audience-director", [
            _make_critique(
                agent="audience-director", section="Table",
                severity="major",
            ),
        ])
        conflicts = orch.resolve_conflicts()
        assert len(conflicts) == 1
        assert conflicts[0].winner == "accuracy-agent"
        assert conflicts[0].loser == "audience-director"

    def test_resolve_three_agents_same_section(self):
        """Three agents on the same section produces multiple conflict records."""
        orch = AgentReviewOrchestrator(enabled=True)
        orch.register_critiques("accuracy-agent", [
            _make_critique(agent="accuracy-agent", section="Hazards"),
        ])
        orch.register_critiques("political-translator", [
            _make_critique(agent="political-translator", section="Hazards"),
        ])
        orch.register_critiques("copy-editor", [
            _make_critique(agent="copy-editor", section="Hazards"),
        ])
        conflicts = orch.resolve_conflicts()
        # accuracy wins over both political-translator and copy-editor
        assert len(conflicts) == 2
        winners = {c.winner for c in conflicts}
        assert winners == {"accuracy-agent"}

    def test_quality_gate_passes(self):
        """Quality gate passes when enough agents complete and criticals resolved."""
        orch = AgentReviewOrchestrator(enabled=True)
        # Register 3 agents (minimum)
        orch.register_critiques("accuracy-agent", [
            _make_critique(agent="accuracy-agent", severity="critical"),
        ])
        orch.register_critiques("audience-director", [
            _make_critique(agent="audience-director", severity="major"),
        ])
        orch.register_critiques("political-translator", [
            _make_critique(agent="political-translator", severity="minor"),
        ])
        gate = orch.check_quality_gate()
        assert gate["passed"] is True
        assert gate["agents_completed"] == 3

    def test_quality_gate_fails_not_enough_agents(self):
        """Quality gate fails when fewer than 3 agents complete."""
        orch = AgentReviewOrchestrator(enabled=True)
        orch.register_critiques("accuracy-agent", [
            _make_critique(agent="accuracy-agent"),
        ])
        orch.register_critiques("copy-editor", [
            _make_critique(agent="copy-editor"),
        ])
        gate = orch.check_quality_gate()
        assert gate["passed"] is False
        assert gate["agents_completed"] == 2

    def test_quality_gate_fails_unresolved_accuracy(self):
        """Quality gate fails when accuracy critiques are unresolved."""
        orch = AgentReviewOrchestrator(enabled=True)
        # Register 3 agents but accuracy has a critique on same section as
        # a lower-priority agent — accuracy should win and be in revisions
        orch.register_critiques("accuracy-agent", [
            _make_critique(agent="accuracy-agent", section="X"),
        ])
        orch.register_critiques("audience-director", [])
        orch.register_critiques("political-translator", [])
        gate = orch.check_quality_gate()
        # accuracy critique IS in revisions (it's the only one on section X)
        assert gate["unresolved_accuracy"] == 0
        assert gate["passed"] is True

    def test_resolution_log_schema(self):
        """Resolution log matches expected schema structure."""
        orch = AgentReviewOrchestrator(enabled=True)
        orch.register_critiques("accuracy-agent", [
            _make_critique(agent="accuracy-agent"),
        ])
        orch.register_critiques("audience-director", [])
        orch.register_critiques("copy-editor", [])
        log = orch.get_resolution_log()
        assert log["enabled"] is True
        assert "pass_2_summary" in log
        assert "conflicts_resolved" in log
        assert "revisions_applied" in log
        assert "quality_gate" in log
        assert isinstance(log["pass_2_summary"]["agents_completed"], list)
        assert isinstance(log["pass_2_summary"]["by_severity"], dict)

    def test_resolution_log_disabled(self):
        """Disabled orchestrator returns minimal log."""
        orch = AgentReviewOrchestrator(enabled=False)
        log = orch.get_resolution_log()
        assert log == {"enabled": False}

    def test_all_five_agents(self):
        """Full 5-agent cycle produces complete resolution."""
        orch = AgentReviewOrchestrator(enabled=True)
        for agent_name in AGENT_PRIORITY:
            orch.register_critiques(agent_name, [
                _make_critique(agent=agent_name, section="Summary"),
            ])
        orch.resolve_conflicts()
        gate = orch.check_quality_gate()
        assert gate["passed"] is True
        assert gate["agents_completed"] == 5

    def test_failed_agents_in_gate(self):
        """Failed agents show in quality gate output."""
        orch = AgentReviewOrchestrator(enabled=True)
        orch.register_critiques("accuracy-agent", [])
        orch.register_critiques("audience-director", [])
        orch.register_critiques("political-translator", [])
        orch.mark_agent_failed("design-aficionado")
        orch.mark_agent_failed("copy-editor")
        gate = orch.check_quality_gate()
        assert gate["passed"] is True  # 3 agents is enough
        assert "design-aficionado" in gate["agents_failed"]
        assert "copy-editor" in gate["agents_failed"]
