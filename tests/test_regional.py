"""Tests for regional aggregation system (Doc C/D generation).

Tests cover:
  - Tribe-to-region assignment by state overlap
  - Crosscutting region includes all Tribes
  - Award aggregation
  - Shared hazard detection and ranking
  - Congressional delegation overlap detection
  - Coverage gap identification
  - Multi-state Tribe assignment
  - RegionalContext dataclass population
"""

import json
from unittest.mock import MagicMock

import pytest

from src.packets.context import TribePacketContext
from src.packets.regional import RegionalAggregator, RegionalContext


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

SAMPLE_REGIONAL_CONFIG = {
    "version": "1.0",
    "description": "Test regional config",
    "regions": {
        "pnw": {
            "name": "Pacific Northwest / Columbia River Basin",
            "short_name": "Pacific Northwest",
            "states": ["WA", "OR", "ID", "MT"],
            "core_frame": "Drought, fisheries, snowpack",
            "key_programs": ["bia_tcr", "noaa_tribal"],
            "treaty_trust_angle": "Treaty fisheries, ceded lands",
        },
        "alaska": {
            "name": "Alaska",
            "short_name": "Alaska",
            "states": ["AK"],
            "core_frame": "Permafrost, coastal erosion",
            "key_programs": ["bia_tcr"],
            "treaty_trust_angle": "Existential displacement",
        },
        "crosscutting": {
            "name": "Cross-Cutting (All Regions)",
            "short_name": "Cross-Cutting",
            "states": [],
            "core_frame": "Science infrastructure",
            "key_programs": ["bia_tcr", "epa_gap"],
            "treaty_trust_angle": "All treaty/trust angles",
        },
    },
}


@pytest.fixture
def config_path(tmp_path):
    """Write sample regional config and return the path."""
    p = tmp_path / "regional_config.json"
    p.write_text(json.dumps(SAMPLE_REGIONAL_CONFIG), encoding="utf-8")
    return p


@pytest.fixture
def mock_registry():
    """Create a mock registry with Tribes in WA, AK, and multi-state."""
    registry = MagicMock()
    registry.get_all.return_value = [
        {"tribe_id": "tribe_wa_1", "name": "Tribe WA One", "states": ["WA"]},
        {"tribe_id": "tribe_wa_2", "name": "Tribe WA Two", "states": ["WA"]},
        {"tribe_id": "tribe_ak_1", "name": "Tribe AK One", "states": ["AK"]},
        {"tribe_id": "tribe_multi", "name": "Tribe Multi", "states": ["WA", "OR"]},
        {"tribe_id": "tribe_tx_1", "name": "Tribe TX One", "states": ["TX"]},
    ]
    return registry


@pytest.fixture
def aggregator(config_path, mock_registry):
    """Create a RegionalAggregator with test config and registry."""
    return RegionalAggregator(
        regional_config_path=config_path,
        registry=mock_registry,
    )


def _make_context(
    tribe_id: str,
    tribe_name: str,
    states: list[str] | None = None,
    awards: list[dict] | None = None,
    hazard_profile: dict | None = None,
    senators: list[dict] | None = None,
    representatives: list[dict] | None = None,
    economic_impact: dict | None = None,
) -> TribePacketContext:
    """Helper to create a TribePacketContext with specific fields."""
    return TribePacketContext(
        tribe_id=tribe_id,
        tribe_name=tribe_name,
        states=states or [],
        awards=awards or [],
        hazard_profile=hazard_profile or {},
        senators=senators or [],
        representatives=representatives or [],
        economic_impact=economic_impact or {},
        generated_at="2026-02-11T00:00:00Z",
    )


# ---------------------------------------------------------------------------
# Test: Tribe assignment by state
# ---------------------------------------------------------------------------


class TestTribeAssignment:
    """Tests for region-to-Tribe assignment logic."""

    def test_tribe_assignment_by_state(self, aggregator):
        """WA Tribes assigned to PNW, AK Tribes to Alaska."""
        pnw_ids = aggregator.get_tribe_ids_for_region("pnw")
        alaska_ids = aggregator.get_tribe_ids_for_region("alaska")

        assert "tribe_wa_1" in pnw_ids
        assert "tribe_wa_2" in pnw_ids
        assert "tribe_ak_1" in alaska_ids
        assert "tribe_ak_1" not in pnw_ids
        assert "tribe_wa_1" not in alaska_ids

    def test_crosscutting_includes_all_tribes(self, aggregator):
        """Crosscutting region includes all Tribes regardless of state."""
        cc_ids = aggregator.get_tribe_ids_for_region("crosscutting")
        assert len(cc_ids) == 5
        assert "tribe_wa_1" in cc_ids
        assert "tribe_ak_1" in cc_ids
        assert "tribe_tx_1" in cc_ids

    def test_multi_state_tribe_assignment(self, aggregator):
        """Tribe with states [WA, OR] is assigned to PNW."""
        pnw_ids = aggregator.get_tribe_ids_for_region("pnw")
        assert "tribe_multi" in pnw_ids

    def test_unmatched_state_not_assigned(self, aggregator):
        """TX Tribe is not assigned to PNW or Alaska."""
        pnw_ids = aggregator.get_tribe_ids_for_region("pnw")
        alaska_ids = aggregator.get_tribe_ids_for_region("alaska")
        assert "tribe_tx_1" not in pnw_ids
        assert "tribe_tx_1" not in alaska_ids

    def test_get_region_ids(self, aggregator):
        """get_region_ids returns all configured region IDs."""
        ids = aggregator.get_region_ids()
        assert set(ids) == {"pnw", "alaska", "crosscutting"}

    def test_assignment_cache(self, aggregator):
        """Second call returns cached result (same object)."""
        first = aggregator.get_tribe_ids_for_region("pnw")
        second = aggregator.get_tribe_ids_for_region("pnw")
        assert first is second


# ---------------------------------------------------------------------------
# Test: Award aggregation
# ---------------------------------------------------------------------------


class TestAwardAggregation:
    """Tests for regional award aggregation."""

    def test_award_aggregation(self, aggregator):
        """Three Tribes with known awards sum correctly."""
        contexts = [
            _make_context("t1", "Tribe 1", awards=[
                {"obligation": 100_000.0, "program_id": "bia_tcr"},
                {"obligation": 50_000.0, "program_id": "epa_gap"},
            ]),
            _make_context("t2", "Tribe 2", awards=[
                {"obligation": 200_000.0, "program_id": "bia_tcr"},
            ]),
            _make_context("t3", "Tribe 3", awards=[]),
        ]
        total, coverage = aggregator._aggregate_awards(contexts)
        assert total == 350_000.0
        assert coverage == 2  # t3 has no awards

    def test_award_aggregation_empty(self, aggregator):
        """No awards yields zero total and zero coverage."""
        contexts = [
            _make_context("t1", "Tribe 1"),
            _make_context("t2", "Tribe 2"),
        ]
        total, coverage = aggregator._aggregate_awards(contexts)
        assert total == 0.0
        assert coverage == 0

    def test_award_aggregation_handles_non_numeric(self, aggregator):
        """Non-numeric obligation values are safely ignored."""
        contexts = [
            _make_context("t1", "Tribe 1", awards=[
                {"obligation": "N/A", "program_id": "bia_tcr"},
                {"obligation": 100.0, "program_id": "epa_gap"},
            ]),
        ]
        total, coverage = aggregator._aggregate_awards(contexts)
        assert total == 100.0
        assert coverage == 1


# ---------------------------------------------------------------------------
# Test: Shared hazard detection
# ---------------------------------------------------------------------------


class TestSharedHazards:
    """Tests for shared hazard detection and ranking."""

    def test_shared_hazard_detection(self, aggregator):
        """Hazards shared by more Tribes rank higher."""
        contexts = [
            _make_context("t1", "Tribe 1", hazard_profile={
                "fema_nri": {
                    "top_hazards": [
                        {"type": "DRGT", "risk_score": 10.0},
                        {"type": "WFIR", "risk_score": 8.0},
                    ],
                    "composite": {"risk_score": 15.0},
                },
            }),
            _make_context("t2", "Tribe 2", hazard_profile={
                "fema_nri": {
                    "top_hazards": [
                        {"type": "DRGT", "risk_score": 12.0},
                        {"type": "RFLD", "risk_score": 5.0},
                    ],
                    "composite": {"risk_score": 14.0},
                },
            }),
            _make_context("t3", "Tribe 3", hazard_profile={
                "fema_nri": {
                    "top_hazards": [
                        {"type": "DRGT", "risk_score": 11.0},
                        {"type": "WFIR", "risk_score": 7.0},
                        {"type": "RFLD", "risk_score": 4.0},
                    ],
                    "composite": {"risk_score": 13.0},
                },
            }),
        ]
        shared = aggregator._find_shared_hazards(contexts)

        # DRGT appears in all 3 Tribes, should be first
        assert shared[0]["hazard_type"] == "DRGT"
        assert shared[0]["tribe_count"] == 3

        # WFIR and RFLD each appear in 2 Tribes
        second_types = {shared[1]["hazard_type"], shared[2]["hazard_type"]}
        assert second_types == {"WFIR", "RFLD"}

    def test_shared_hazard_max_five(self, aggregator):
        """Returns at most 5 hazard types."""
        # Create context with 7 different hazard types
        contexts = [
            _make_context("t1", "Tribe 1", hazard_profile={
                "fema_nri": {
                    "top_hazards": [
                        {"type": f"HAZ_{i}", "risk_score": float(i)}
                        for i in range(7)
                    ],
                    "composite": {"risk_score": 10.0},
                },
            }),
        ]
        shared = aggregator._find_shared_hazards(contexts)
        assert len(shared) <= 5

    def test_shared_hazards_empty(self, aggregator):
        """No hazard data yields empty shared list."""
        contexts = [_make_context("t1", "Tribe 1")]
        shared = aggregator._find_shared_hazards(contexts)
        assert shared == []

    def test_shared_hazards_sources_nesting(self, aggregator):
        """Handles sources.fema_nri nesting pattern from Phase 13 cache."""
        contexts = [
            _make_context("t1", "Tribe 1", hazard_profile={
                "sources": {
                    "fema_nri": {
                        "top_hazards": [
                            {"type": "DRGT", "risk_score": 10.0},
                        ],
                        "composite": {"risk_score": 12.0},
                    },
                },
            }),
        ]
        shared = aggregator._find_shared_hazards(contexts)
        assert len(shared) == 1
        assert shared[0]["hazard_type"] == "DRGT"


# ---------------------------------------------------------------------------
# Test: Delegation overlap
# ---------------------------------------------------------------------------


class TestDelegationOverlap:
    """Tests for congressional delegation overlap detection."""

    def test_delegation_overlap(self, aggregator):
        """Senator serving two Tribes detected as overlap."""
        shared_senator = {
            "bioguide_id": "S001",
            "formatted_name": "Sen. Shared",
            "state": "WA",
            "committees": [{"committee_name": "Appropriations"}],
        }
        contexts = [
            _make_context(
                "t1", "Tribe 1",
                senators=[shared_senator],
                representatives=[{
                    "bioguide_id": "R001",
                    "formatted_name": "Rep. Unique",
                    "committees": [],
                }],
            ),
            _make_context(
                "t2", "Tribe 2",
                senators=[shared_senator],
            ),
        ]
        overlap, total_sens, total_reps = aggregator._find_delegation_overlap(
            contexts
        )

        assert len(overlap) == 1
        assert overlap[0]["member_name"] == "Sen. Shared"
        assert overlap[0]["tribe_count"] == 2
        assert total_sens == 1
        assert total_reps == 1

    def test_no_delegation_overlap(self, aggregator):
        """No overlap when each member serves one Tribe."""
        contexts = [
            _make_context("t1", "Tribe 1", senators=[{
                "bioguide_id": "S001",
                "formatted_name": "Sen. Alpha",
                "committees": [],
            }]),
            _make_context("t2", "Tribe 2", senators=[{
                "bioguide_id": "S002",
                "formatted_name": "Sen. Beta",
                "committees": [],
            }]),
        ]
        overlap, total_sens, total_reps = aggregator._find_delegation_overlap(
            contexts
        )
        assert overlap == []
        assert total_sens == 2

    def test_delegation_overlap_empty(self, aggregator):
        """Empty delegation yields no overlap."""
        contexts = [_make_context("t1", "Tribe 1")]
        overlap, total_sens, total_reps = aggregator._find_delegation_overlap(
            contexts
        )
        assert overlap == []
        assert total_sens == 0
        assert total_reps == 0


# ---------------------------------------------------------------------------
# Test: Coverage gap identification
# ---------------------------------------------------------------------------


class TestCoverageGaps:
    """Tests for coverage gap identification."""

    def test_coverage_gap_identification(self, aggregator):
        """Mix of complete and partial Tribes identifies correct gaps."""
        contexts = [
            # Complete: has awards, hazards, delegation
            _make_context(
                "t1", "Tribe Complete",
                awards=[{"obligation": 100.0}],
                hazard_profile={
                    "fema_nri": {"composite": {"risk_score": 10.0}},
                },
                senators=[{"bioguide_id": "S001", "formatted_name": "Sen. X"}],
            ),
            # Missing awards
            _make_context(
                "t2", "Tribe No Awards",
                hazard_profile={
                    "fema_nri": {"composite": {"risk_score": 5.0}},
                },
                senators=[{"bioguide_id": "S002", "formatted_name": "Sen. Y"}],
            ),
            # Missing hazards
            _make_context(
                "t3", "Tribe No Hazards",
                awards=[{"obligation": 200.0}],
                senators=[{"bioguide_id": "S003", "formatted_name": "Sen. Z"}],
            ),
            # Missing delegation
            _make_context(
                "t4", "Tribe No Delegation",
                awards=[{"obligation": 50.0}],
                hazard_profile={
                    "fema_nri": {"composite": {"risk_score": 3.0}},
                },
            ),
        ]
        no_awards, no_hazards, no_delegation = aggregator._identify_gaps(contexts)

        assert no_awards == ["t2"]
        assert no_hazards == ["t3"]
        assert no_delegation == ["t4"]


# ---------------------------------------------------------------------------
# Test: Economic aggregation
# ---------------------------------------------------------------------------


class TestEconomicAggregation:
    """Tests for regional economic impact aggregation."""

    def test_economic_aggregation(self, aggregator):
        """Economic impacts sum across Tribes."""
        contexts = [
            _make_context("t1", "Tribe 1", economic_impact={
                "total_impact_low": 100_000.0,
                "total_impact_high": 200_000.0,
                "total_jobs_low": 5.0,
                "total_jobs_high": 10.0,
            }),
            _make_context("t2", "Tribe 2", economic_impact={
                "total_impact_low": 50_000.0,
                "total_impact_high": 80_000.0,
                "total_jobs_low": 3.0,
                "total_jobs_high": 6.0,
            }),
        ]
        econ = aggregator._aggregate_economics(contexts)
        assert econ["total_low"] == 150_000.0
        assert econ["total_high"] == 280_000.0
        assert econ["total_jobs_low"] == 8.0
        assert econ["total_jobs_high"] == 16.0

    def test_economic_aggregation_empty(self, aggregator):
        """No economic data yields zeros."""
        contexts = [_make_context("t1", "Tribe 1")]
        econ = aggregator._aggregate_economics(contexts)
        assert econ["total_low"] == 0.0
        assert econ["total_high"] == 0.0


# ---------------------------------------------------------------------------
# Test: Full aggregate() flow
# ---------------------------------------------------------------------------


class TestFullAggregation:
    """Tests for the full aggregate() method."""

    def test_aggregate_produces_regional_context(self, aggregator):
        """aggregate() returns a fully populated RegionalContext."""
        contexts = [
            _make_context(
                "t1", "Tribe 1",
                states=["WA"],
                awards=[{"obligation": 100_000.0, "program_id": "bia_tcr"}],
                hazard_profile={
                    "fema_nri": {
                        "top_hazards": [{"type": "DRGT", "risk_score": 10.0}],
                        "composite": {"risk_score": 15.0},
                    },
                },
                senators=[{
                    "bioguide_id": "S001",
                    "formatted_name": "Sen. Alpha",
                    "committees": [],
                }],
                economic_impact={
                    "total_impact_low": 180_000.0,
                    "total_impact_high": 240_000.0,
                    "total_jobs_low": 1.0,
                    "total_jobs_high": 2.0,
                },
            ),
        ]

        result = aggregator.aggregate("pnw", contexts)

        assert isinstance(result, RegionalContext)
        assert result.region_id == "pnw"
        assert result.region_name == "Pacific Northwest / Columbia River Basin"
        assert result.short_name == "Pacific Northwest"
        assert result.tribe_count == 1
        assert result.total_awards == 100_000.0
        assert result.award_coverage == 1
        assert result.hazard_coverage == 1
        assert result.congressional_coverage == 1
        assert result.composite_risk_score == 15.0
        assert len(result.top_shared_hazards) == 1
        assert result.top_shared_hazards[0]["hazard_type"] == "DRGT"
        assert result.total_senators == 1
        assert result.aggregate_economic_impact["total_low"] == 180_000.0
        assert result.generated_at  # non-empty

    def test_aggregate_empty_contexts(self, aggregator):
        """aggregate() handles empty context list gracefully."""
        result = aggregator.aggregate("pnw", [])
        assert result.tribe_count == 0
        assert result.total_awards == 0.0
        assert result.top_shared_hazards == []
        assert result.delegation_overlap == []


# ---------------------------------------------------------------------------
# Test: Composite risk score
# ---------------------------------------------------------------------------


class TestCompositeRisk:
    """Tests for composite risk score calculation."""

    def test_composite_risk_average(self, aggregator):
        """Composite risk is averaged across Tribes with data."""
        contexts = [
            _make_context("t1", "Tribe 1", hazard_profile={
                "fema_nri": {"composite": {"risk_score": 10.0}},
            }),
            _make_context("t2", "Tribe 2", hazard_profile={
                "fema_nri": {"composite": {"risk_score": 20.0}},
            }),
            _make_context("t3", "Tribe 3"),  # No hazard data
        ]
        score = aggregator._compute_composite_risk(contexts)
        assert score == 15.0  # average of 10 and 20

    def test_composite_risk_no_data(self, aggregator):
        """Returns 0.0 when no Tribes have risk scores."""
        contexts = [_make_context("t1", "Tribe 1")]
        score = aggregator._compute_composite_risk(contexts)
        assert score == 0.0


# ---------------------------------------------------------------------------
# Test: Regional DOCX section renderers
# ---------------------------------------------------------------------------


class TestRegionalDocxSections:
    """Tests for regional document section renderers."""

    def _make_regional_ctx(self):
        """Create a sample RegionalContext for rendering tests."""
        return RegionalContext(
            region_id="pnw",
            region_name="Pacific Northwest / Columbia River Basin",
            short_name="Pacific Northwest",
            core_frame="Drought, fisheries, snowpack",
            treaty_trust_angle="Treaty fisheries, ceded lands",
            key_programs=["bia_tcr", "noaa_tribal"],
            states=["WA", "OR", "ID", "MT"],
            tribe_count=3,
            tribes=[
                {"tribe_id": "t1", "tribe_name": "Tribe One", "states": ["WA"]},
                {"tribe_id": "t2", "tribe_name": "Tribe Two", "states": ["OR"]},
                {"tribe_id": "t3", "tribe_name": "Tribe Three", "states": ["WA", "ID"]},
            ],
            total_awards=500_000.0,
            award_coverage=2,
            hazard_coverage=3,
            congressional_coverage=3,
            top_shared_hazards=[
                {"hazard_type": "DRGT", "tribe_count": 3, "avg_score": 12.5},
                {"hazard_type": "WFIR", "tribe_count": 2, "avg_score": 8.0},
            ],
            composite_risk_score=14.5,
            aggregate_economic_impact={
                "total_low": 900_000.0,
                "total_high": 1_200_000.0,
                "total_jobs_low": 7.0,
                "total_jobs_high": 12.0,
            },
            delegation_overlap=[{
                "member_name": "Sen. Shared",
                "role": "Senator",
                "tribe_count": 2,
                "tribe_ids": ["t1", "t2"],
                "committees": ["Appropriations"],
            }],
            total_senators=4,
            total_representatives=6,
            tribes_without_awards=["t3"],
            tribes_without_hazards=[],
            tribes_without_delegation=[],
            generated_at="2026-02-11T00:00:00Z",
        )

    def test_regional_cover_page_renders(self):
        """Cover page renders without error for Doc C."""
        from docx import Document as DocxDocument
        from src.packets.doc_types import DOC_C
        from src.packets.docx_styles import StyleManager
        from src.packets.docx_regional_sections import (
            render_regional_cover_page,
        )

        doc = DocxDocument()
        sm = StyleManager(doc)
        ctx = self._make_regional_ctx()
        render_regional_cover_page(doc, ctx, sm, DOC_C)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "CONFIDENTIAL" in text
        assert "Pacific Northwest" in text

    def test_regional_cover_page_congressional(self):
        """Cover page for Doc D has no CONFIDENTIAL banner."""
        from docx import Document as DocxDocument
        from src.packets.doc_types import DOC_D
        from src.packets.docx_styles import StyleManager
        from src.packets.docx_regional_sections import (
            render_regional_cover_page,
        )

        doc = DocxDocument()
        sm = StyleManager(doc)
        ctx = self._make_regional_ctx()
        render_regional_cover_page(doc, ctx, sm, DOC_D)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "CONFIDENTIAL" not in text
        assert "Pacific Northwest" in text

    def test_regional_exec_summary_internal(self):
        """Internal exec summary includes coverage gap analysis."""
        from docx import Document as DocxDocument
        from src.packets.doc_types import DOC_C
        from src.packets.docx_styles import StyleManager
        from src.packets.docx_regional_sections import (
            render_regional_executive_summary,
        )

        doc = DocxDocument()
        sm = StyleManager(doc)
        ctx = self._make_regional_ctx()
        render_regional_executive_summary(doc, ctx, sm, DOC_C)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Coverage Gap" in text
        assert "Shared Delegation" in text

    def test_regional_exec_summary_congressional(self):
        """Congressional exec summary is evidence-only."""
        from docx import Document as DocxDocument
        from src.packets.doc_types import DOC_D
        from src.packets.docx_styles import StyleManager
        from src.packets.docx_regional_sections import (
            render_regional_executive_summary,
        )

        doc = DocxDocument()
        sm = StyleManager(doc)
        ctx = self._make_regional_ctx()
        render_regional_executive_summary(doc, ctx, sm, DOC_D)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Federal Investment Overview" in text
        assert "Coverage Gap" not in text
        assert "Shared Delegation" not in text

    def test_regional_hazard_synthesis(self):
        """Hazard synthesis renders shared hazards table."""
        from docx import Document as DocxDocument
        from src.packets.docx_styles import StyleManager
        from src.packets.docx_regional_sections import (
            render_regional_hazard_synthesis,
        )

        doc = DocxDocument()
        sm = StyleManager(doc)
        ctx = self._make_regional_ctx()
        render_regional_hazard_synthesis(doc, ctx, sm)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Regional Hazard Synthesis" in text
        assert "DRGT" in doc.tables[0].rows[1].cells[0].text

    def test_regional_delegation_internal_overlap(self):
        """Internal delegation section shows overlap table."""
        from docx import Document as DocxDocument
        from src.packets.doc_types import DOC_C
        from src.packets.docx_styles import StyleManager
        from src.packets.docx_regional_sections import (
            render_regional_delegation,
        )

        doc = DocxDocument()
        sm = StyleManager(doc)
        ctx = self._make_regional_ctx()
        render_regional_delegation(doc, ctx, sm, DOC_C)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Delegation Overlap Analysis" in text
        # Check overlap table
        assert len(doc.tables) >= 1
        assert "Sen. Shared" in doc.tables[0].rows[1].cells[0].text

    def test_regional_delegation_congressional_no_overlap(self):
        """Congressional delegation has no overlap analysis."""
        from docx import Document as DocxDocument
        from src.packets.doc_types import DOC_D
        from src.packets.docx_styles import StyleManager
        from src.packets.docx_regional_sections import (
            render_regional_delegation,
        )

        doc = DocxDocument()
        sm = StyleManager(doc)
        ctx = self._make_regional_ctx()
        render_regional_delegation(doc, ctx, sm, DOC_D)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Delegation Overlap Analysis" not in text

    def test_regional_appendix_lists_tribes(self):
        """Appendix lists all Tribes in the region."""
        from docx import Document as DocxDocument
        from src.packets.docx_styles import StyleManager
        from src.packets.docx_regional_sections import (
            render_regional_appendix,
        )

        doc = DocxDocument()
        sm = StyleManager(doc)
        ctx = self._make_regional_ctx()
        render_regional_appendix(doc, ctx, sm)

        # Check table has 3 tribe rows + 1 header
        assert len(doc.tables) >= 1
        assert len(doc.tables[0].rows) == 4  # 1 header + 3 tribes

    def test_regional_award_landscape(self):
        """Award landscape shows aggregate and gap data."""
        from docx import Document as DocxDocument
        from src.packets.docx_styles import StyleManager
        from src.packets.docx_regional_sections import (
            render_regional_award_landscape,
        )

        doc = DocxDocument()
        sm = StyleManager(doc)
        ctx = self._make_regional_ctx()
        render_regional_award_landscape(doc, ctx, sm)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "$500,000" in text
        assert "2 of 3" in text
        assert "Investment Gap" in text
