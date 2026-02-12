"""Tests for congressional bill intelligence DOCX rendering (Plan 15-05).

Validates render_bill_intelligence_section for audience-differentiated
output (Doc A internal briefing vs. Doc B facts-only), confidence badges,
bill sorting, edge cases, and DocxEngine integration.

Also tests render_regional_congressional_section for shared-bill
aggregation across multiple Tribes.

30+ tests covering all rendering paths.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from src.packets.context import TribePacketContext
from src.packets.doc_types import DOC_A, DOC_B, DOC_C, DOC_D
from src.packets.docx_engine import DocxEngine
from src.packets.docx_sections import render_bill_intelligence_section
from src.packets.docx_styles import StyleManager
from src.packets.economic import (
    EconomicImpactCalculator,
    TribeEconomicSummary,
)


# ---------------------------------------------------------------------------
# Test helpers
# ---------------------------------------------------------------------------


def _mock_programs() -> list[dict]:
    """Return a minimal list of mock programs."""
    return [
        {
            "id": "bia_tcr",
            "name": "BIA Tribal Climate Resilience",
            "agency": "Bureau of Indian Affairs",
            "federal_home": "DOI / BIA",
            "priority": "critical",
            "access_type": "direct_service",
            "funding_type": "Discretionary",
            "cfda": "15.156",
            "ci_status": "AT_RISK",
            "ci_determination": "FY26 budget uncertainty.",
            "description": "Annual awards for climate adaptation planning.",
            "advocacy_lever": "Increase TCR funding to $50M annually",
            "tightened_language": "TCR faces flat funding.",
            "keywords": ["climate", "resilience"],
            "search_queries": ["tribal climate resilience"],
            "confidence_index": 0.85,
        },
    ]


def _mock_context(
    congressional_intel: dict | None = None,
    **overrides,
) -> TribePacketContext:
    """Return a TribePacketContext with optional congressional intel."""
    defaults = dict(
        tribe_id="epa_001",
        tribe_name="Alpha Tribe",
        states=["WA"],
        ecoregions=["pacific_northwest"],
        districts=[{"district": "WA-06", "state": "WA", "overlap_pct": 80.0}],
        senators=[
            {
                "bioguide_id": "S001",
                "formatted_name": "Sen. Jane Smith (D-WA)",
                "state": "WA",
                "committees": [
                    {
                        "committee_id": "SSAP",
                        "committee_name": "Senate Appropriations Committee",
                        "role": "member",
                    },
                ],
            },
        ],
        representatives=[
            {
                "bioguide_id": "R001",
                "formatted_name": "Rep. Carol Lee (D-WA-06)",
                "state": "WA",
                "district": "WA-06",
                "committees": [],
            },
        ],
        awards=[
            {
                "award_id": "A001",
                "cfda": "15.156",
                "program_id": "bia_tcr",
                "obligation": 100000.0,
            },
        ],
        hazard_profile={
            "sources": {
                "fema_nri": {
                    "composite": {"risk_rating": "Relatively High", "risk_score": 28.5},
                    "top_hazards": [
                        {"type": "Wildfire", "code": "WFIR", "risk_score": 45.0},
                    ],
                },
                "usfs_wildfire": {},
            },
        },
        congressional_intel=congressional_intel or {},
        generated_at="2026-02-10T12:00:00+00:00",
    )
    defaults.update(overrides)
    return TribePacketContext(**defaults)


def _mock_bill(
    bill_id: str = "119-HR-1234",
    title: str = "Tribal Climate Resilience Act",
    relevance_score: float = 0.85,
    bill_type: str = "HR",
    bill_number: int = 1234,
    **overrides,
) -> dict:
    """Return a mock bill dict for congressional intel."""
    data = {
        "bill_id": bill_id,
        "bill_type": bill_type,
        "bill_number": bill_number,
        "title": title,
        "relevance_score": relevance_score,
        "sponsor": {"name": "Jane Doe", "party": "D", "state": "WA"},
        "cosponsor_count": 5,
        "matched_programs": ["bia_tcr"],
        "subjects": ["Native Americans", "Climate Change"],
        "latest_action": {
            "text": "Referred to Committee on Natural Resources.",
            "date": "2026-01-15",
        },
        "committees": [{"name": "House Committee on Natural Resources"}],
    }
    data.update(overrides)
    return data


def _mock_economic_summary() -> TribeEconomicSummary:
    """Return a TribeEconomicSummary with minimal data."""
    calc = EconomicImpactCalculator()
    return calc.compute(
        tribe_id="epa_001",
        tribe_name="Alpha Tribe",
        awards=[{"obligation": 100000.0, "program_id": "bia_tcr", "cfda": "15.156"}],
        districts=[{"district": "WA-06", "overlap_pct": 80.0}],
        programs={p["id"]: p for p in _mock_programs()},
    )


def _make_engine(tmp_path: Path, doc_type_config=None) -> DocxEngine:
    """Create a DocxEngine with tmp_path output directory."""
    output_dir = tmp_path / "outputs" / "packets"
    config = {"packets": {"output_dir": str(output_dir)}}
    programs = {p["id"]: p for p in _mock_programs()}
    return DocxEngine(config, programs, doc_type_config=doc_type_config)


def _render_bill_section_to_text(
    congressional_intel: dict,
    doc_type_config=None,
) -> str:
    """Render bill intelligence section and return all paragraph text."""
    doc = DocxDocument()
    sm = StyleManager(doc)
    context = _mock_context(congressional_intel=congressional_intel)
    render_bill_intelligence_section(
        doc, context, sm, doc_type_config=doc_type_config
    )
    return "\n".join(p.text for p in doc.paragraphs)


def _get_all_text(path: Path) -> str:
    """Read a DOCX and return all paragraph text joined."""
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _get_all_doc_text(doc) -> str:
    """Extract all text from a Document including paragraphs and table cells."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ===========================================================================
# render_bill_intelligence_section Tests (~20 tests)
# ===========================================================================


class TestBillIntelligenceRendering:
    """Validate render_bill_intelligence_section output."""

    def test_empty_bills_shows_no_legislation(self):
        """Empty bills list renders 'no tracked legislation' message."""
        text = _render_bill_section_to_text({"bills": []})
        assert "No tracked legislation" in text

    def test_empty_intel_dict_shows_no_legislation(self):
        """Missing congressional_intel shows 'no tracked legislation'."""
        text = _render_bill_section_to_text({})
        assert "No tracked legislation" in text

    def test_doc_a_heading_is_internal(self):
        """Doc A heading says 'Legislation Affecting Your Programs'."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_A)
        assert "Legislation Affecting Your Programs" in text

    def test_doc_b_heading_is_congressional(self):
        """Doc B heading says 'Relevant Legislation'."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_B)
        assert "Relevant Legislation" in text

    def test_doc_a_has_talking_points(self):
        """Doc A renders talking points for top bills."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_A)
        assert "Talking Points" in text

    def test_doc_b_no_talking_points(self):
        """Doc B contains ZERO talking points."""
        intel = {
            "bills": [_mock_bill(relevance_score=0.8)],
            "scan_date": "2026-02-10",
        }
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_B)
        assert "Talking Points" not in text

    def test_doc_b_no_strategy_words(self):
        """Doc B contains ZERO strategy/timing/urgency words."""
        intel = {
            "bills": [_mock_bill(relevance_score=0.8)],
            "scan_date": "2026-02-10",
        }
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_B)
        text_lower = text.lower()

        # Strategy words that should NEVER appear in Doc B
        forbidden = [
            "talking points",
            "timing:",
            "active movement",
            "moderate activity",
            "monitoring --",
        ]
        for word in forbidden:
            assert word not in text_lower, f"Doc B contains forbidden term: {word}"

    def test_confidence_badge_appears(self):
        """Confidence badge [HIGH/MEDIUM/LOW] appears in heading."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        doc = DocxDocument()
        sm = StyleManager(doc)
        context = _mock_context(congressional_intel=intel)
        render_bill_intelligence_section(doc, context, sm)

        # Check for badge in heading paragraph runs
        heading_found = False
        for para in doc.paragraphs:
            full_text = "".join(r.text for r in para.runs)
            if "Relevant Legislation" in full_text or "Legislation Affecting" in full_text:
                if "[HIGH]" in full_text or "[MEDIUM]" in full_text or "[LOW]" in full_text:
                    heading_found = True
                    break
        assert heading_found, "Confidence badge not found in heading"

    def test_bills_sorted_by_relevance_descending(self):
        """Bills are sorted by relevance_score descending."""
        intel = {
            "bills": [
                _mock_bill(bill_id="119-HR-LOW", relevance_score=0.3, bill_number=1),
                _mock_bill(bill_id="119-HR-HIGH", relevance_score=0.95, bill_number=2),
                _mock_bill(bill_id="119-HR-MID", relevance_score=0.6, bill_number=3),
            ],
            "scan_date": "2026-02-10",
        }
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_A)

        # HIGH should appear before MID, MID before LOW
        pos_high = text.find("HR 2")
        pos_mid = text.find("HR 3")
        pos_low = text.find("HR 1")
        assert pos_high < pos_mid < pos_low

    def test_single_bill_renders(self):
        """Single bill renders without error."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel)
        assert "HR 1234" in text
        assert "Tribal Climate Resilience Act" in text

    def test_ten_plus_bills_renders(self):
        """10+ bills all render without error."""
        bills = [
            _mock_bill(
                bill_id=f"119-HR-{i}",
                title=f"Bill {i}",
                bill_number=i,
                relevance_score=round(0.5 + i * 0.04, 2),
            )
            for i in range(1, 13)
        ]
        intel = {"bills": bills, "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_A)

        # Top 3 get full briefing, rest get compact cards
        assert "Talking Points" in text  # Top bills have talking points
        assert "Additional Tracked Legislation" in text  # Remaining bills section

    def test_doc_a_full_briefing_has_sponsor(self):
        """Doc A full briefing shows sponsor info."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_A)
        assert "Jane Doe" in text
        assert "cosponsor" in text.lower()

    def test_doc_a_full_briefing_has_committees(self):
        """Doc A full briefing shows committee referrals."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_A)
        assert "Natural Resources" in text

    def test_doc_a_full_briefing_has_affected_programs(self):
        """Doc A full briefing shows affected programs."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_A)
        assert "Affected Programs" in text
        assert "bia_tcr" in text

    def test_doc_b_filters_low_relevance(self):
        """Doc B filters out bills with relevance < 0.5."""
        intel = {
            "bills": [
                _mock_bill(bill_id="119-HR-LOW", title="Low Bill", relevance_score=0.3, bill_number=1),
                _mock_bill(bill_id="119-HR-HIGH", title="High Bill", relevance_score=0.8, bill_number=2),
            ],
            "scan_date": "2026-02-10",
        }
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_B)
        assert "High Bill" in text
        assert "Low Bill" not in text

    def test_doc_b_all_below_threshold_shows_message(self):
        """Doc B with all bills below 0.5 shows 'no direct relevance' message."""
        intel = {
            "bills": [
                _mock_bill(relevance_score=0.2, bill_number=1),
                _mock_bill(relevance_score=0.1, bill_number=2),
            ],
            "scan_date": "2026-02-10",
        }
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_B)
        assert "No legislation with direct program relevance" in text

    def test_doc_b_facts_only_has_sponsor(self):
        """Doc B facts-only entry shows sponsor name."""
        intel = {
            "bills": [_mock_bill(relevance_score=0.8)],
            "scan_date": "2026-02-10",
        }
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_B)
        assert "Sponsor" in text
        assert "Jane Doe" in text

    def test_default_no_dtc_compact_cards(self):
        """Default (no doc_type_config) renders compact cards."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel, doc_type_config=None)
        # Compact cards show bill number + title but not talking points
        assert "HR 1234" in text
        assert "Talking Points" not in text

    def test_doc_a_timing_note_rendered(self):
        """Doc A full briefing includes timing note."""
        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_A)
        # Timing note depends on action date recency
        # With action date "2026-01-15", it should show some timing note
        assert "Timing:" in text or "movement" in text.lower() or "Monitoring" in text


# ===========================================================================
# DocxEngine generate() ordering tests (~5 tests)
# ===========================================================================


class TestDocxEngineBillIntelligenceOrdering:
    """Verify bill intelligence appears in correct document position."""

    def test_bill_intelligence_appears_in_document(self, tmp_path):
        """Bill intelligence section appears in generated document."""
        intel = {
            "bills": [_mock_bill()],
            "scan_date": "2026-02-10",
        }
        engine = _make_engine(tmp_path, doc_type_config=DOC_A)
        context = _mock_context(congressional_intel=intel)
        programs = _mock_programs()
        economic = _mock_economic_summary()
        asks = [
            {
                "id": "ask_1",
                "name": "Multi-Year Funding",
                "description": "Multi-year authorization.",
                "target": "Congress",
                "urgency": "FY26",
                "programs": ["bia_tcr"],
            },
        ]
        path = engine.generate(
            context, programs, economic, asks, doc_type_config=DOC_A
        )
        text = _get_all_text(path)

        assert "Legislation Affecting Your Programs" in text

    def test_bill_intelligence_before_delegation(self, tmp_path):
        """Bill intelligence appears before delegation section."""
        intel = {
            "bills": [_mock_bill()],
            "scan_date": "2026-02-10",
        }
        engine = _make_engine(tmp_path, doc_type_config=DOC_A)
        context = _mock_context(congressional_intel=intel)
        programs = _mock_programs()
        economic = _mock_economic_summary()
        asks = []
        path = engine.generate(
            context, programs, economic, asks, doc_type_config=DOC_A
        )
        text = _get_all_text(path)

        pos_bills = text.find("Legislation Affecting Your Programs")
        pos_delegation = text.find("Congressional Delegation")
        assert pos_bills < pos_delegation, "Bills section should come before delegation"

    def test_missing_congressional_intel_no_crash(self, tmp_path):
        """Missing congressional_intel doesn't crash document generation."""
        engine = _make_engine(tmp_path, doc_type_config=DOC_B)
        context = _mock_context(congressional_intel={})
        programs = _mock_programs()
        economic = _mock_economic_summary()
        asks = []
        path = engine.generate(
            context, programs, economic, asks, doc_type_config=DOC_B
        )
        text = _get_all_text(path)

        assert "No tracked legislation" in text

    def test_none_congressional_intel_no_crash(self, tmp_path):
        """None congressional_intel dict doesn't crash generation."""
        engine = _make_engine(tmp_path)
        # Context with no congressional_intel at all
        context = TribePacketContext(
            tribe_id="epa_safe",
            tribe_name="Safe Tribe",
            states=["OR"],
            generated_at="2026-02-12T00:00:00+00:00",
        )
        programs = _mock_programs()
        economic = TribeEconomicSummary(tribe_id="epa_safe", tribe_name="Safe Tribe")
        asks = []
        path = engine.generate(context, programs, economic, asks)
        text = _get_all_text(path)

        assert "No tracked legislation" in text

    def test_doc_b_bill_section_in_output(self, tmp_path):
        """Doc B includes bill intelligence section with facts."""
        intel = {
            "bills": [_mock_bill(relevance_score=0.9)],
            "scan_date": "2026-02-10",
        }
        engine = _make_engine(tmp_path, doc_type_config=DOC_B)
        context = _mock_context(congressional_intel=intel)
        programs = _mock_programs()
        economic = _mock_economic_summary()
        asks = []
        path = engine.generate(
            context, programs, economic, asks, doc_type_config=DOC_B
        )
        text = _get_all_text(path)

        assert "Relevant Legislation" in text
        assert "HR 1234" in text


# ===========================================================================
# Regional congressional section tests (~5 tests)
# ===========================================================================


class TestRegionalCongressionalSection:
    """Validate render_regional_congressional_section output."""

    def _make_regional_context(self, tribes_intel: list[dict]):
        """Create a mock RegionalContext with per-Tribe intel."""
        from src.packets.regional import RegionalContext

        tribes = []
        for i, intel in enumerate(tribes_intel):
            tribes.append({
                "tribe_id": f"epa_{i:09d}",
                "tribe_name": f"Tribe {i}",
                "states": ["WA"],
                "congressional_intel": intel,
            })

        return RegionalContext(
            region_id="pnw",
            region_name="Pacific Northwest",
            short_name="PNW",
            core_frame="Wildfire and flooding dominate.",
            treaty_trust_angle="Treaty obligations require action.",
            key_programs=["bia_tcr"],
            states=["WA", "OR"],
            tribe_count=len(tribes),
            tribes=tribes,
            total_awards=200000.0,
            award_coverage=len(tribes),
            hazard_coverage=len(tribes),
            congressional_coverage=len(tribes),
            top_shared_hazards=[],
            composite_risk_score=0.0,
            aggregate_economic_impact={},
            delegation_overlap=[],
            total_senators=2,
            total_representatives=3,
            tribes_without_awards=[],
            tribes_without_hazards=[],
            tribes_without_delegation=[],
            generated_at="2026-02-12T00:00:00+00:00",
        )

    def test_multiple_tribes_shared_bill(self):
        """Shared bill across Tribes aggregates correctly."""
        from src.packets.docx_regional_sections import render_regional_congressional_section

        shared_bill = _mock_bill(bill_id="119-HR-999", bill_number=999, relevance_score=0.8)
        tribe1_intel = {"bills": [shared_bill], "scan_date": "2026-02-10"}
        tribe2_intel = {"bills": [shared_bill], "scan_date": "2026-02-10"}

        ctx = self._make_regional_context([tribe1_intel, tribe2_intel])

        doc = DocxDocument()
        sm = StyleManager(doc)
        render_regional_congressional_section(doc, ctx, sm)

        # Bill data is in table cells, not paragraphs
        all_text = _get_all_doc_text(doc)
        assert "HR 999" in all_text or "119-HR-999" in all_text
        # Tribes Affected should be "2"
        assert "2" in all_text

    def test_empty_bills_shows_message(self):
        """No bills across all Tribes shows 'no legislation' message."""
        from src.packets.docx_regional_sections import render_regional_congressional_section

        ctx = self._make_regional_context([{"bills": []}, {"bills": []}])

        doc = DocxDocument()
        sm = StyleManager(doc)
        render_regional_congressional_section(doc, ctx, sm)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "No tracked legislation" in text

    def test_confidence_badge_in_heading(self):
        """Regional section heading includes confidence badge."""
        from src.packets.docx_regional_sections import render_regional_congressional_section

        intel = {"bills": [_mock_bill()], "scan_date": "2026-02-10"}
        ctx = self._make_regional_context([intel])

        doc = DocxDocument()
        sm = StyleManager(doc)
        render_regional_congressional_section(doc, ctx, sm)

        heading_text = ""
        for para in doc.paragraphs:
            full = "".join(r.text for r in para.runs)
            if "Regional Legislation Overview" in full:
                heading_text = full
                break
        assert "[HIGH]" in heading_text or "[MEDIUM]" in heading_text or "[LOW]" in heading_text

    def test_doc_c_coordination_note(self):
        """Doc C (internal) includes coordination note for shared bills."""
        from src.packets.docx_regional_sections import render_regional_congressional_section

        shared_bill = _mock_bill(bill_id="119-HR-SHARED", relevance_score=0.9)
        intel1 = {"bills": [shared_bill], "scan_date": "2026-02-10"}
        intel2 = {"bills": [shared_bill], "scan_date": "2026-02-10"}
        ctx = self._make_regional_context([intel1, intel2])

        doc = DocxDocument()
        sm = StyleManager(doc)
        render_regional_congressional_section(doc, ctx, sm, doc_type_config=DOC_C)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "coordinated" in text.lower() or "Coordinated" in text

    def test_unique_bills_from_different_tribes(self):
        """Unique bills from different Tribes appear separately."""
        from src.packets.docx_regional_sections import render_regional_congressional_section

        bill1 = _mock_bill(bill_id="119-HR-100", title="Bill A", bill_number=100, relevance_score=0.7)
        bill2 = _mock_bill(bill_id="119-S-200", title="Bill B", bill_type="S", bill_number=200, relevance_score=0.6)
        intel1 = {"bills": [bill1], "scan_date": "2026-02-10"}
        intel2 = {"bills": [bill2], "scan_date": "2026-02-10"}

        ctx = self._make_regional_context([intel1, intel2])

        doc = DocxDocument()
        sm = StyleManager(doc)
        render_regional_congressional_section(doc, ctx, sm)

        # Bill data is in table cells, not paragraphs
        all_text = _get_all_doc_text(doc)
        assert "HR 100" in all_text or "119-HR-100" in all_text
        assert "S 200" in all_text or "119-S-200" in all_text


# ===========================================================================
# Air gap verification for bill rendering
# ===========================================================================


class TestBillRenderingAirGap:
    """Ensure Doc B bill rendering has zero strategy content."""

    def test_doc_b_zero_strategy_words_comprehensive(self):
        """Comprehensive check: Doc B has absolutely zero strategy language."""
        intel = {
            "bills": [
                _mock_bill(relevance_score=0.9),
                _mock_bill(
                    bill_id="119-S-567",
                    title="Senate Climate Bill",
                    bill_type="S",
                    bill_number=567,
                    relevance_score=0.7,
                ),
            ],
            "scan_date": "2026-02-10",
        }
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_B)
        text_lower = text.lower()

        # Comprehensive list of strategy words that must be absent
        strategy_terms = [
            "talking points",
            "timing:",
            "active movement",
            "moderate activity",
            "monitoring --",
            "strategic",
            "leverage",
            "approach",
            "key subject areas",
            "underscores the importance",
            "relies on for climate",
        ]

        violations = [t for t in strategy_terms if t in text_lower]
        assert not violations, (
            f"Doc B bill section contains strategy terms: {violations}"
        )

    def test_doc_b_has_facts_only_content(self):
        """Doc B bill entries have exactly: title, status, sponsor, programs."""
        intel = {
            "bills": [_mock_bill(relevance_score=0.8)],
            "scan_date": "2026-02-10",
        }
        text = _render_bill_section_to_text(intel, doc_type_config=DOC_B)

        # Facts that SHOULD be present
        assert "HR 1234" in text
        assert "Tribal Climate Resilience Act" in text
        assert "Sponsor" in text
        assert "Affected Programs" in text
