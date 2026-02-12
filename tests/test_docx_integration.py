"""Integration tests for full DOCX generation pipeline (Plan 07-04).

Tests end-to-end flow: PacketOrchestrator.generate_packet() produces a
valid .docx file with cover page, Hot Sheets, and appendix. Exercises
both the standard path (Tribe with awards and hazards) and edge cases
(zero-award zero-hazard Tribes, invalid Tribe names).

All tests use tmp_path fixtures with mock data -- no network or real
data files required.
"""

import json
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _write_json(path: Path, data) -> None:
    """Write data as JSON to the given path, creating parent dirs."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f)


# ---------------------------------------------------------------------------
# Mock data builders
# ---------------------------------------------------------------------------

def _mock_programs() -> list[dict]:
    """Return a list of 4 mock programs (2 critical, 2 high)."""
    return [
        {
            "id": "bia_tcr",
            "name": "BIA Tribal Climate Resilience",
            "agency": "Bureau of Indian Affairs",
            "federal_home": "DOI / BIA",
            "priority": "critical",
            "access_type": "direct_service",
            "funding_type": "Discretionary",
            "cfda": "15.156",
            "ci_status": "AT_RISK",
            "ci_determination": "FY26 budget uncertainty.",
            "description": "Annual awards to Tribes for climate adaptation planning.",
            "advocacy_lever": "Increase TCR funding to $50M annually",
            "tightened_language": "TCR faces flat funding despite rising demand.",
            "keywords": ["climate", "resilience"],
            "search_queries": ["tribal climate resilience"],
            "confidence_index": 0.85,
        },
        {
            "id": "fema_bric",
            "name": "FEMA BRIC",
            "agency": "FEMA",
            "federal_home": "DHS / FEMA",
            "priority": "critical",
            "access_type": "competitive",
            "funding_type": "Discretionary",
            "cfda": "97.047",
            "ci_status": "FLAGGED",
            "ci_determination": "BRIC terminated April 2025.",
            "description": "Building Resilient Infrastructure and Communities.",
            "advocacy_lever": "Restore BRIC funding and Tribal set-aside",
            "tightened_language": "BRIC program terminated; Tribal projects stranded.",
            "keywords": ["mitigation", "resilience"],
            "search_queries": ["BRIC tribal"],
            "confidence_index": 0.30,
        },
        {
            "id": "epa_gap",
            "name": "EPA Indian Environmental GAP",
            "agency": "EPA",
            "federal_home": "EPA / OEJ",
            "priority": "high",
            "access_type": "direct_service",
            "funding_type": "Discretionary",
            "cfda": "66.926",
            "ci_status": "STABLE",
            "ci_determination": "Permanent authorization.",
            "description": "General Assistance Program for Tribal environmental capacity.",
            "advocacy_lever": "Expand GAP funding for climate monitoring",
            "keywords": ["environmental", "capacity"],
            "search_queries": ["EPA GAP tribal"],
            "confidence_index": 0.70,
        },
        {
            "id": "doe_indian_energy",
            "name": "DOE Indian Energy",
            "agency": "Department of Energy",
            "federal_home": "DOE / OIE",
            "priority": "high",
            "access_type": "competitive",
            "funding_type": "Discretionary",
            "cfda": "81.104",
            "ci_status": "UNCERTAIN",
            "ci_determination": "IIJA supplement timing unclear.",
            "description": "Energy development and efficiency for Tribal communities.",
            "advocacy_lever": "Streamline Indian energy project approvals",
            "keywords": ["energy", "tribal"],
            "search_queries": ["DOE indian energy"],
            "confidence_index": 0.55,
        },
    ]


def _mock_registry() -> dict:
    """Return mock tribal registry data."""
    return {
        "metadata": {"total_tribes": 2, "placeholder": False},
        "tribes": [
            {
                "tribe_id": "epa_001",
                "bia_code": "100",
                "name": "Alpha Tribe of Arizona",
                "states": ["AZ"],
                "alternate_names": ["Alpha AZ Tribe"],
                "epa_region": "9",
                "bia_recognized": True,
            },
            {
                "tribe_id": "epa_002",
                "bia_code": "200",
                "name": "Beta Band of Montana",
                "states": ["MT"],
                "alternate_names": [],
                "epa_region": "8",
                "bia_recognized": True,
            },
        ],
    }


def _mock_ecoregion() -> dict:
    """Return mock ecoregion config."""
    return {
        "metadata": {"version": "test"},
        "ecoregions": {
            "southwest": {
                "name": "Southwest",
                "states": ["AZ", "NM", "NV"],
                "priority_programs": ["bia_tcr", "fema_bric"],
            },
            "northern_rockies": {
                "name": "Northern Rockies",
                "states": ["MT", "ID", "WY"],
                "priority_programs": ["usda_wildfire", "bia_tcr"],
            },
        },
    }


def _mock_congress() -> dict:
    """Return mock congressional cache."""
    return {
        "metadata": {
            "congress_session": "119",
            "total_members": 3,
            "total_committees": 1,
            "total_tribes_mapped": 2,
        },
        "members": {
            "S001": {
                "bioguide_id": "S001",
                "name": "Smith, Jane",
                "state": "AZ",
                "district": None,
                "party": "Democratic",
                "party_abbr": "D",
                "chamber": "Senate",
                "formatted_name": "Sen. Jane Smith (D-AZ)",
            },
            "R001": {
                "bioguide_id": "R001",
                "name": "Lee, Carol",
                "state": "AZ",
                "district": 2,
                "party": "Democratic",
                "party_abbr": "D",
                "chamber": "House",
                "formatted_name": "Rep. Carol Lee (D-AZ-02)",
            },
        },
        "committees": {
            "SLIA": {
                "name": "Senate Committee on Indian Affairs",
                "chamber": "Senate",
                "members": [
                    {"bioguide_id": "S001", "name": "Smith, Jane", "role": "member"},
                ],
            },
        },
        "delegations": {
            "epa_001": {
                "tribe_id": "epa_001",
                "districts": [
                    {
                        "district": "AZ-02",
                        "state": "AZ",
                        "aiannh_name": "Alpha Reservation",
                        "overlap_pct": 95.0,
                    },
                ],
                "states": ["AZ"],
                "senators": [
                    {
                        "bioguide_id": "S001",
                        "formatted_name": "Sen. Jane Smith (D-AZ)",
                        "committees": [
                            {
                                "committee_id": "SLIA",
                                "committee_name": "Senate Committee on Indian Affairs",
                                "role": "member",
                            }
                        ],
                    },
                ],
                "representatives": [
                    {
                        "bioguide_id": "R001",
                        "formatted_name": "Rep. Carol Lee (D-AZ-02)",
                        "committees": [],
                    },
                ],
            },
            "epa_002": {
                "tribe_id": "epa_002",
                "districts": [],
                "states": ["MT"],
                "senators": [],
                "representatives": [],
            },
        },
    }


def _mock_awards(tribe_id: str) -> dict:
    """Return mock award cache for a Tribe."""
    if tribe_id == "epa_001":
        return {
            "tribe_id": "epa_001",
            "awards": [
                {
                    "award_id": "A001",
                    "cfda": "15.156",
                    "program_id": "bia_tcr",
                    "obligation": 500000.0,
                    "start_date": "2024-01-01",
                    "end_date": "2025-01-01",
                },
                {
                    "award_id": "A002",
                    "cfda": "97.047",
                    "program_id": "fema_bric",
                    "obligation": 250000.0,
                    "start_date": "2023-06-01",
                    "end_date": "2024-06-01",
                },
            ],
        }
    # epa_002 has no awards
    return {"tribe_id": tribe_id, "awards": []}


def _mock_hazards(tribe_id: str) -> dict:
    """Return mock hazard profile for a Tribe."""
    if tribe_id == "epa_001":
        return {
            "tribe_id": "epa_001",
            "sources": {
                "fema_nri": {
                    "composite": {
                        "risk_rating": "Relatively High",
                        "risk_score": 28.5,
                        "sovi_rating": "Relatively Moderate",
                        "sovi_score": 15.2,
                    },
                    "top_hazards": [
                        {
                            "type": "Wildfire",
                            "code": "WFIR",
                            "risk_score": 45.0,
                            "risk_rating": "Very High",
                        },
                        {
                            "type": "Drought",
                            "code": "DRGT",
                            "risk_score": 30.0,
                            "risk_rating": "Relatively High",
                        },
                    ],
                },
                "usfs_wildfire": {
                    "risk_to_homes": 0.85,
                    "likelihood": 0.45,
                },
            },
        }
    # epa_002 has no hazard data
    return {"tribe_id": tribe_id, "sources": {}}


def _mock_graph_schema() -> dict:
    """Return mock graph_schema.json with structural asks."""
    return {
        "authorities": [],
        "funding_vehicles": [],
        "barriers": [],
        "structural_asks": [
            {
                "id": "ask_multi_year",
                "name": "Multi-Year Funding Stability",
                "description": "Shift from annual to multi-year authorization.",
                "target": "Congress",
                "urgency": "FY26",
                "mitigates": ["bar_appropriations_volatility"],
                "programs": ["bia_tcr", "epa_gap", "doe_indian_energy"],
            },
            {
                "id": "ask_match_waivers",
                "name": "Match/Cost-Share Waivers",
                "description": "Categorical Tribal exemptions from cost-share.",
                "target": "Congress",
                "urgency": "Immediate",
                "mitigates": ["bar_cost_share"],
                "programs": ["fema_bric"],
            },
        ],
    }


# ---------------------------------------------------------------------------
# Main fixture: full Phase 7 config
# ---------------------------------------------------------------------------

@pytest.fixture
def full_phase7_config(tmp_path):
    """Build a complete mock environment for DOCX generation.

    Creates all required data files in tmp_path, returns config dict
    and programs list ready for PacketOrchestrator.
    """
    # Write all mock data files
    _write_json(tmp_path / "tribal_registry.json", _mock_registry())
    _write_json(tmp_path / "ecoregion_config.json", _mock_ecoregion())
    _write_json(tmp_path / "congressional_cache.json", _mock_congress())

    # Per-Tribe caches
    award_dir = tmp_path / "award_cache"
    hazard_dir = tmp_path / "hazard_profiles"
    for tid in ("epa_001", "epa_002"):
        _write_json(award_dir / f"{tid}.json", _mock_awards(tid))
        _write_json(hazard_dir / f"{tid}.json", _mock_hazards(tid))

    # Graph schema
    _write_json(tmp_path / "graph_schema.json", _mock_graph_schema())

    output_dir = tmp_path / "outputs" / "packets"

    config = {
        "packets": {
            "tribal_registry": {"data_path": str(tmp_path / "tribal_registry.json")},
            "ecoregion": {"data_path": str(tmp_path / "ecoregion_config.json")},
            "congressional_cache": {"data_path": str(tmp_path / "congressional_cache.json")},
            "output_dir": str(output_dir),
            "awards": {
                "alias_path": str(tmp_path / "tribal_aliases.json"),
                "cache_dir": str(award_dir),
                "fuzzy_threshold": 85,
            },
            "hazards": {
                "nri_dir": str(tmp_path / "nri"),
                "usfs_dir": str(tmp_path / "usfs"),
                "cache_dir": str(hazard_dir),
                "crosswalk_path": str(tmp_path / "crosswalk.json"),
            },
            "docx": {
                "enabled": True,
                "output_dir": str(output_dir),
            },
        },
    }

    programs = _mock_programs()

    # Return config, programs, tmp_path, and output_dir for tests
    return {
        "config": config,
        "programs": programs,
        "tmp_path": tmp_path,
        "output_dir": output_dir,
    }


def _make_orchestrator(phase7):
    """Create a PacketOrchestrator from full_phase7_config fixture data.

    Patches _load_structural_asks to use the mock graph_schema.json in tmp_path.
    """
    from src.packets.orchestrator import PacketOrchestrator

    orch = PacketOrchestrator(phase7["config"], phase7["programs"])

    # Patch _load_structural_asks to read from our tmp_path
    graph_path = phase7["tmp_path"] / "graph_schema.json"
    def patched_load():
        with open(graph_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get("structural_asks", [])

    orch._load_structural_asks = patched_load
    return orch


# ===========================================================================
# Integration tests
# ===========================================================================


class TestGeneratePacket:
    """End-to-end tests for PacketOrchestrator.generate_packet()."""

    def test_generate_packet_single_tribe(self, full_phase7_config):
        """Generate a full .docx for a Tribe with awards and hazards."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")

        assert path is not None
        assert path.exists()
        assert path.suffix == ".docx"
        assert path.stat().st_size > 0

        # Verify python-docx can re-read the file
        doc = Document(str(path))
        assert len(doc.paragraphs) > 0

    def test_generate_packet_zero_awards_zero_hazards(self, full_phase7_config):
        """Generate for a Tribe with no awards and no hazard data."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Beta Band of Montana")

        assert path is not None
        assert path.exists()
        assert path.suffix == ".docx"

        # Verify valid DOCX
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Should have First-Time Applicant content
        assert "First-Time Applicant" in all_text

    def test_generate_packet_returns_path(self, full_phase7_config):
        """generate_packet() returns a Path object to the .docx file."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")

        assert isinstance(path, Path)
        assert path.name == "epa_001.docx"

    def test_generate_packet_invalid_tribe(self, full_phase7_config):
        """generate_packet() returns None for nonexistent Tribe."""
        orch = _make_orchestrator(full_phase7_config)
        result = orch.generate_packet("Nonexistent Tribe XYZ123")

        assert result is None


class TestDocxContent:
    """Verify generated .docx has expected structural content."""

    def test_docx_has_cover_page(self, full_phase7_config):
        """Cover page contains Tribe name and metadata."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")
        doc = Document(str(path))

        # Collect first 15 paragraphs text
        first_paras = [p.text for p in doc.paragraphs[:15]]
        combined = "\n".join(first_paras)

        assert "Alpha Tribe of Arizona" in combined
        assert "FY26 Climate Resilience Program Priorities" in combined
        assert "State(s): AZ" in combined
        assert "Ecoregion(s): southwest" in combined
        assert "Congressional Session: 119" in combined

    def test_docx_has_hotsheets(self, full_phase7_config):
        """Hot Sheets contain program names from inventory."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        # All 4 programs should appear as Hot Sheets (4 programs, all relevant)
        assert "BIA Tribal Climate Resilience" in all_text
        assert "FEMA BRIC" in all_text
        assert "EPA Indian Environmental GAP" in all_text

    def test_docx_has_appendix_when_omitted(self, full_phase7_config):
        """Appendix content present when programs are omitted.

        With only 4 programs and MIN_PROGRAMS=8, all are included --
        so appendix says "All tracked programs are included."
        """
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        # With 4 programs and min=8, all are included (no omitted)
        assert "Appendix" in all_text or "All tracked programs are included" in all_text

    def test_docx_has_delegation_section(self, full_phase7_config):
        """Hot Sheets include delegation info."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Sen. Jane Smith (D-AZ)" in all_text
        assert "Your Delegation" in all_text

    def test_docx_has_economic_impact(self, full_phase7_config):
        """Economic impact sections appear for programs with awards."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Should have economic content (BEA multiplier language)
        assert "economic" in all_text.lower() or "BEA" in all_text

    def test_docx_has_advocacy_language(self, full_phase7_config):
        """Hot Sheets include advocacy position sections."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Should have advocacy content
        assert "Advocacy Position" in all_text


class TestEconomicSummaryInContext:
    """Verify economic impact is stored in context after generation."""

    def test_economic_summary_populated(self, full_phase7_config):
        """After generation, context.economic_impact is populated."""
        orch = _make_orchestrator(full_phase7_config)

        # Build context directly to inspect it
        tribe = orch.registry.resolve("Alpha Tribe of Arizona")
        context = orch._build_context(tribe)

        # Before generation, economic_impact is empty
        assert context.economic_impact == {}

        # Generate (which populates economic_impact)
        orch.generate_packet_from_context(context, tribe)

        # After generation, economic_impact should be populated
        assert context.economic_impact != {}
        assert "total_obligation" in context.economic_impact
        assert context.economic_impact["total_obligation"] > 0

    def test_economic_summary_zero_award_tribe(self, full_phase7_config):
        """Zero-award Tribe gets benchmark-based economic summary."""
        orch = _make_orchestrator(full_phase7_config)

        tribe = orch.registry.resolve("Beta Band of Montana")
        context = orch._build_context(tribe)

        orch.generate_packet_from_context(context, tribe)

        # Should still have economic data (benchmark path)
        assert context.economic_impact != {}
        assert "by_program" in context.economic_impact


class TestConfigDocx:
    """Verify config loading with docx section."""

    def test_config_has_docx_section(self, full_phase7_config):
        """Config contains packets.docx section."""
        config = full_phase7_config["config"]
        docx_cfg = config.get("packets", {}).get("docx", {})

        assert docx_cfg.get("enabled") is True
        assert "output_dir" in docx_cfg

    def test_config_json_has_docx(self):
        """scanner_config.json has packets.docx section."""
        from src.paths import SCANNER_CONFIG_PATH
        config_path = SCANNER_CONFIG_PATH
        if not config_path.exists():
            pytest.skip("scanner_config.json not found")

        with open(config_path, "r", encoding="utf-8") as f:
            config = json.load(f)

        docx_cfg = config.get("packets", {}).get("docx", {})
        assert docx_cfg.get("enabled") is True
        assert docx_cfg.get("output_dir") == "outputs/packets"


class TestDocxReReadability:
    """Verify generated .docx files are valid and re-readable."""

    def test_reopen_with_python_docx(self, full_phase7_config):
        """Generated .docx opens correctly with python-docx Document()."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")

        # Must not raise
        doc = Document(str(path))

        # Must have meaningful content
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) >= 10, f"Expected >= 10 non-empty paragraphs, got {len(texts)}"

    def test_reopen_zero_award_tribe(self, full_phase7_config):
        """Zero-award Tribe .docx opens correctly."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Beta Band of Montana")

        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) >= 5


class TestStructuralAsks:
    """Verify structural asks are loaded and rendered."""

    def test_structural_asks_loaded(self, full_phase7_config):
        """Orchestrator loads structural asks from graph_schema.json."""
        orch = _make_orchestrator(full_phase7_config)
        asks = orch._load_structural_asks()

        assert len(asks) == 2
        assert asks[0]["id"] == "ask_multi_year"

    def test_structural_asks_in_docx(self, full_phase7_config):
        """Structural asks appear in generated .docx for matching programs."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Multi-year ask should appear (matches bia_tcr, epa_gap, doe_indian_energy)
        assert "Multi-Year Funding Stability" in all_text


class TestTableOfContents:
    """Verify table of contents rendering."""

    def test_toc_lists_programs(self, full_phase7_config):
        """TOC lists program names with numbering."""
        orch = _make_orchestrator(full_phase7_config)
        path = orch.generate_packet("Alpha Tribe of Arizona")
        doc = Document(str(path))

        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Program Hot Sheets" in all_text
        # Should reference the program count
        assert "program analyses" in all_text


class TestFullSuiteNoRegression:
    """Meta-test: verify existing test modules import and have expected tests."""

    def test_existing_test_modules_importable(self):
        """All existing test modules import without error."""
        import importlib

        modules = [
            "tests.test_packets",
            "tests.test_docx_styles",
            "tests.test_economic",
            "tests.test_docx_hotsheet",
            "tests.test_decision_engine",
        ]
        for mod_name in modules:
            mod = importlib.import_module(mod_name)
            assert mod is not None, f"Failed to import {mod_name}"

    def test_minimum_test_count(self):
        """Verify we have >= 193 tests across all test files via collection."""
        from src.paths import PROJECT_ROOT
        result = subprocess.run(
            [sys.executable, "-m", "pytest", "tests/", "--collect-only", "-q"],
            capture_output=True,
            text=True,
            cwd=str(PROJECT_ROOT),
            timeout=60,
        )

        output = result.stdout.strip()
        # Last line format: "N tests collected" or "N tests / N errors"
        last_line = output.split("\n")[-1] if output else ""

        # Extract test count from collection output
        parts = last_line.split()
        if parts and parts[0].isdigit():
            count = int(parts[0])
            assert count >= 214, (
                f"Expected >= 214 tests collected, got {count}. "
                f"Output: {last_line}"
            )
        else:
            # Fallback: count lines that look like test items
            test_lines = [
                line for line in output.split("\n")
                if "::" in line and "test_" in line
            ]
            assert len(test_lines) >= 214, (
                f"Expected >= 214 test items, got {len(test_lines)}"
            )


# ===========================================================================
# Phase 16 critical coverage gap tests
# ===========================================================================


class TestRenderChangeTrackingFont:
    """Verify render_change_tracking heading uses Arial font (GAP-002).

    ACCURACY-006 found that document.add_heading() bypasses StyleManager
    Arial override. Fix changed to document.add_paragraph(..., style='Heading 2').
    This test verifies the heading font is correct.
    """

    def test_change_tracking_heading_uses_heading2_style(self, full_phase7_config):
        """render_change_tracking heading uses Heading 2 style with Arial font."""
        from src.packets.docx_sections import render_change_tracking
        from src.packets.docx_styles import StyleManager

        doc = Document()
        sm = StyleManager(doc)

        changes = [
            {"type": "ci_status_change", "description": "BIA TCR moved from STABLE to AT_RISK"},
            {"type": "new_award", "description": "New FEMA BRIC award: $250,000"},
        ]

        render_change_tracking(doc, changes, "2026-01-01T00:00:00Z", sm)

        # Find the heading paragraph
        heading_para = None
        for para in doc.paragraphs:
            if "Since Last Packet" in para.text:
                heading_para = para
                break

        assert heading_para is not None, "Change tracking heading not found"

        # Verify it uses the Heading 2 style (which StyleManager overrides to Arial)
        assert heading_para.style.name == "Heading 2", (
            f"Expected 'Heading 2' style, got '{heading_para.style.name}'"
        )
        assert heading_para.style.font.name == "Arial", (
            f"Expected Arial font on Heading 2, got '{heading_para.style.font.name}'"
        )

    def test_change_tracking_empty_changes_not_rendered(self, full_phase7_config):
        """render_change_tracking with empty changes list renders nothing."""
        from src.packets.docx_sections import render_change_tracking
        from src.packets.docx_styles import StyleManager

        doc = Document()
        sm = StyleManager(doc)

        render_change_tracking(doc, [], "2026-01-01T00:00:00Z", sm)

        # No paragraphs added (empty changes = no section)
        assert all("Since Last Packet" not in p.text for p in doc.paragraphs)
