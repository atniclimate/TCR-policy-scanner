"""Tests for PacketChangeTracker and render_change_tracking.

Covers first-generation behavior, change detection for all 5 change types,
state persistence, corruption handling, security (path traversal, size limits),
atomic write cleanup, and DOCX rendering of the change tracking section.
"""

import json
import os
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from src.packets.change_tracker import MAX_STATE_FILE_SIZE, PacketChangeTracker
from src.packets.context import TribePacketContext
from src.packets.docx_sections import render_change_tracking


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_context(
    tribe_id: str = "epa_123",
    tribe_name: str = "Test Tribe",
    awards: list | None = None,
    hazard_profile: dict | None = None,
) -> TribePacketContext:
    """Create a minimal TribePacketContext for testing."""
    return TribePacketContext(
        tribe_id=tribe_id,
        tribe_name=tribe_name,
        awards=awards or [],
        hazard_profile=hazard_profile or {},
    )


def _make_programs() -> dict:
    """Create a minimal programs dict for testing."""
    return {
        "prog_a": {"id": "prog_a", "name": "Program A", "ci_status": "active"},
        "prog_b": {"id": "prog_b", "name": "Program B", "ci_status": "at_risk"},
    }


# ---------------------------------------------------------------------------
# Test 1: First generation -- no previous state
# ---------------------------------------------------------------------------


class TestFirstGeneration:
    def test_first_generation_no_previous(self, tmp_path: Path) -> None:
        """First generation returns None for load_previous (no state file)."""
        tracker = PacketChangeTracker(state_dir=tmp_path)
        result = tracker.load_previous("epa_123")
        assert result is None


# ---------------------------------------------------------------------------
# Test 2-5: Change detection
# ---------------------------------------------------------------------------


class TestChangeDetection:
    def test_ci_status_change_detected(self) -> None:
        """Detects CI status change for a program."""
        tracker = PacketChangeTracker()
        previous = {
            "program_states": {"prog_a": "active", "prog_b": "at_risk"},
            "total_awards": 0,
            "total_obligation": 0.0,
            "top_hazards": [],
            "advocacy_goal": "new_applicant",
        }
        current = {
            "program_states": {"prog_a": "active", "prog_b": "funded"},
            "total_awards": 0,
            "total_obligation": 0.0,
            "top_hazards": [],
            "advocacy_goal": "new_applicant",
        }
        changes = tracker.diff(previous, current)
        ci_changes = [c for c in changes if c["type"] == "ci_status_change"]
        assert len(ci_changes) == 1
        assert "prog_b" in ci_changes[0]["description"]
        assert "at_risk" in ci_changes[0]["description"]
        assert "funded" in ci_changes[0]["description"]

    def test_award_total_change_detected(self) -> None:
        """Detects change in total obligation amount."""
        tracker = PacketChangeTracker()
        previous = {
            "program_states": {},
            "total_awards": 2,
            "total_obligation": 100000.0,
            "top_hazards": [],
            "advocacy_goal": "renewal",
        }
        current = {
            "program_states": {},
            "total_awards": 2,
            "total_obligation": 250000.0,
            "top_hazards": [],
            "advocacy_goal": "renewal",
        }
        changes = tracker.diff(previous, current)
        award_changes = [c for c in changes if c["type"] == "award_total_change"]
        assert len(award_changes) == 1
        assert "$100,000" in award_changes[0]["description"]
        assert "$250,000" in award_changes[0]["description"]

    def test_advocacy_goal_shift_detected(self) -> None:
        """Detects advocacy goal shift (new_applicant -> renewal)."""
        tracker = PacketChangeTracker()
        previous = {
            "program_states": {},
            "total_awards": 0,
            "total_obligation": 0.0,
            "top_hazards": [],
            "advocacy_goal": "new_applicant",
        }
        current = {
            "program_states": {},
            "total_awards": 1,
            "total_obligation": 50000.0,
            "top_hazards": [],
            "advocacy_goal": "renewal",
        }
        changes = tracker.diff(previous, current)
        goal_changes = [c for c in changes if c["type"] == "advocacy_goal_shift"]
        assert len(goal_changes) == 1
        assert "new_applicant" in goal_changes[0]["description"]
        assert "renewal" in goal_changes[0]["description"]

    def test_new_threat_detected(self) -> None:
        """Detects new hazard threat appearing in top hazards."""
        tracker = PacketChangeTracker()
        previous = {
            "program_states": {},
            "total_awards": 0,
            "total_obligation": 0.0,
            "top_hazards": ["Flooding", "Wildfire"],
            "advocacy_goal": "new_applicant",
        }
        current = {
            "program_states": {},
            "total_awards": 0,
            "total_obligation": 0.0,
            "top_hazards": ["Flooding", "Wildfire", "Drought"],
            "advocacy_goal": "new_applicant",
        }
        changes = tracker.diff(previous, current)
        threat_changes = [c for c in changes if c["type"] == "new_threat"]
        assert len(threat_changes) == 1
        assert "Drought" in threat_changes[0]["description"]

    def test_new_awards_detected(self) -> None:
        """Detects new awards added since last generation."""
        tracker = PacketChangeTracker()
        previous = {
            "program_states": {},
            "total_awards": 1,
            "total_obligation": 50000.0,
            "top_hazards": [],
            "advocacy_goal": "renewal",
        }
        current = {
            "program_states": {},
            "total_awards": 3,
            "total_obligation": 150000.0,
            "top_hazards": [],
            "advocacy_goal": "renewal",
        }
        changes = tracker.diff(previous, current)
        new_awards = [c for c in changes if c["type"] == "new_award"]
        assert len(new_awards) == 1
        assert "2 new award(s)" in new_awards[0]["description"]


# ---------------------------------------------------------------------------
# Test 6: No changes
# ---------------------------------------------------------------------------


class TestNoChanges:
    def test_no_changes_empty_list(self) -> None:
        """Identical states produce empty changes list."""
        tracker = PacketChangeTracker()
        state = {
            "program_states": {"prog_a": "active"},
            "total_awards": 1,
            "total_obligation": 50000.0,
            "top_hazards": ["Flooding"],
            "advocacy_goal": "renewal",
        }
        changes = tracker.diff(state, state)
        assert changes == []


# ---------------------------------------------------------------------------
# Tests 7-8: State persistence
# ---------------------------------------------------------------------------


class TestStatePersistence:
    def test_state_file_written_correctly(self, tmp_path: Path) -> None:
        """save_current writes valid JSON to the correct path."""
        tracker = PacketChangeTracker(state_dir=tmp_path)
        state = {
            "tribe_id": "epa_123",
            "generated_at": "2026-02-10T00:00:00+00:00",
            "program_states": {"prog_a": "active"},
            "total_awards": 1,
            "total_obligation": 50000.0,
            "top_hazards": ["Flooding"],
            "advocacy_goal": "renewal",
        }
        tracker.save_current("epa_123", state)

        saved_path = tmp_path / "epa_123.json"
        assert saved_path.exists()

        with open(saved_path, "r", encoding="utf-8") as f:
            loaded = json.load(f)
        assert loaded["tribe_id"] == "epa_123"
        assert loaded["total_awards"] == 1
        assert loaded["program_states"]["prog_a"] == "active"

    def test_state_roundtrip(self, tmp_path: Path) -> None:
        """State can be saved and loaded back identically."""
        tracker = PacketChangeTracker(state_dir=tmp_path)
        context = _make_context(
            awards=[{"obligation": 50000.0, "program_id": "prog_a"}],
            hazard_profile={
                "fema_nri": {
                    "top_hazards": [
                        {"type": "Flooding", "risk_score": 50.0},
                    ]
                }
            },
        )
        programs = _make_programs()
        current = tracker.compute_current(context, programs)
        tracker.save_current("epa_123", current)

        loaded = tracker.load_previous("epa_123")
        assert loaded is not None
        assert loaded["tribe_id"] == "epa_123"
        assert loaded["total_awards"] == 1
        assert loaded["total_obligation"] == 50000.0
        assert "Flooding" in loaded["top_hazards"]
        assert loaded["program_states"]["prog_a"] == "active"


# ---------------------------------------------------------------------------
# Tests 9-10: Corruption and size limit handling
# ---------------------------------------------------------------------------


class TestCorruptionHandling:
    def test_corrupt_state_file_handled(self, tmp_path: Path) -> None:
        """Corrupt JSON file returns None (treated as first generation)."""
        tracker = PacketChangeTracker(state_dir=tmp_path)
        corrupt_path = tmp_path / "epa_corrupt.json"
        corrupt_path.write_text("NOT VALID JSON {{{{", encoding="utf-8")

        result = tracker.load_previous("epa_corrupt")
        assert result is None

    def test_oversized_state_file_handled(self, tmp_path: Path) -> None:
        """Oversized state file returns None (treated as first generation)."""
        tracker = PacketChangeTracker(state_dir=tmp_path)
        oversized_path = tmp_path / "epa_big.json"
        # Write slightly over the limit
        oversized_path.write_text(
            "x" * (MAX_STATE_FILE_SIZE + 1), encoding="utf-8"
        )

        result = tracker.load_previous("epa_big")
        assert result is None


# ---------------------------------------------------------------------------
# Test 11: Path traversal protection
# ---------------------------------------------------------------------------


class TestSecurity:
    def test_path_traversal_rejected(self, tmp_path: Path) -> None:
        """Path traversal attempts are rejected with ValueError."""
        tracker = PacketChangeTracker(state_dir=tmp_path)

        with pytest.raises(ValueError, match="Invalid tribe_id"):
            tracker._safe_path("..")

        with pytest.raises(ValueError, match="Invalid tribe_id"):
            tracker._safe_path(".")


# ---------------------------------------------------------------------------
# Test 12: Atomic write cleanup
# ---------------------------------------------------------------------------


class TestAtomicWrite:
    def test_atomic_write_cleanup(self, tmp_path: Path) -> None:
        """Verify no temp files are left after successful save."""
        tracker = PacketChangeTracker(state_dir=tmp_path)
        state = {"tribe_id": "epa_123", "generated_at": "2026-01-01"}
        tracker.save_current("epa_123", state)

        # Only the target file should exist (no .json temp files)
        json_files = list(tmp_path.glob("*.json"))
        assert len(json_files) == 1
        assert json_files[0].name == "epa_123.json"

        # No leftover temp files
        tmp_files = list(tmp_path.glob("*.tmp*"))
        assert len(tmp_files) == 0


# ---------------------------------------------------------------------------
# Tests 13-14: DOCX rendering
# ---------------------------------------------------------------------------


class TestRenderChangeTracking:
    def test_render_change_tracking_with_changes(self) -> None:
        """render_change_tracking adds content when changes exist."""
        doc = Document()
        style_manager = MagicMock()
        changes = [
            {"type": "ci_status_change", "description": "Program X: active -> at_risk"},
            {"type": "new_award", "description": "2 new awards recorded"},
            {"type": "new_threat", "description": "New hazard: Drought"},
        ]
        previous_date = "2026-01-15T00:00:00+00:00"

        render_change_tracking(doc, changes, previous_date, style_manager)

        # Extract all paragraph text
        all_text = [p.text for p in doc.paragraphs]
        full_text = " ".join(all_text)

        # Verify heading with date
        assert "Since Last Packet" in full_text
        assert "2026-01-15" in full_text

        # Verify type labels rendered
        assert "Confidence Index Changes" in full_text
        assert "New Awards" in full_text
        assert "New Hazard Threats" in full_text

        # Verify descriptions rendered
        assert "Program X: active -> at_risk" in full_text
        assert "2 new awards recorded" in full_text
        assert "New hazard: Drought" in full_text

        # Verify summary
        assert "3 change(s) detected" in full_text

    def test_render_change_tracking_empty_changes(self) -> None:
        """render_change_tracking adds NO content for empty changes list."""
        doc = Document()
        style_manager = MagicMock()

        render_change_tracking(doc, [], "2026-01-15", style_manager)

        # Document should have no user-added paragraphs
        # (python-docx Document() starts with one default empty paragraph)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) == 0


# ---------------------------------------------------------------------------
# Test 15: New awards detected (via compute_current + diff)
# ---------------------------------------------------------------------------


class TestComputeCurrentIntegration:
    def test_compute_current_captures_state(self, tmp_path: Path) -> None:
        """compute_current captures all expected state fields."""
        tracker = PacketChangeTracker(state_dir=tmp_path)
        context = _make_context(
            awards=[
                {"obligation": 30000.0, "program_id": "prog_a"},
                {"obligation": 20000.0, "program_id": "prog_b"},
            ],
            hazard_profile={
                "fema_nri": {
                    "top_hazards": [
                        {"type": "Wildfire"},
                        {"type": "Flooding"},
                    ]
                }
            },
        )
        programs = _make_programs()
        state = tracker.compute_current(context, programs)

        assert state["tribe_id"] == "epa_123"
        assert "generated_at" in state
        assert state["total_awards"] == 2
        assert state["total_obligation"] == 50000.0
        assert "Wildfire" in state["top_hazards"]
        assert "Flooding" in state["top_hazards"]
        assert state["program_states"]["prog_a"] == "active"
        assert state["program_states"]["prog_b"] == "at_risk"
        assert state["advocacy_goal"] == "renewal"
