"""Tests for Hot Sheet renderer (docx_hotsheet.py).

Covers: HotSheetRenderer section rendering for all CI statuses,
award history (with awards + zero-award first-time applicant),
hazard relevance, economic impact, advocacy language (tightened +
fallback), Key Ask callout, structural asks linkage, delegation
committee-match flags, and CFDA normalization edge cases.
"""

import pytest
from docx import Document

from src.packets.context import TribePacketContext
from src.packets.docx_hotsheet import (
    FRAMING_BY_STATUS,
    RELEVANT_COMMITTEE_KEYWORDS,
    HotSheetRenderer,
)
from src.packets.docx_styles import CI_STATUS_COLORS, StyleManager
from src.packets.economic import (
    EconomicImpactCalculator,
    ProgramEconomicImpact,
    TribeEconomicSummary,
)


# ---------------------------------------------------------------------------
# Fixtures (shared mock data)
# ---------------------------------------------------------------------------


@pytest.fixture
def mock_context():
    """TribePacketContext with awards and hazard data."""
    return TribePacketContext(
        tribe_id="epa_001",
        tribe_name="Alpha Tribe of Arizona",
        states=["AZ"],
        ecoregions=["southwest"],
        districts=[
            {"district": "AZ-02", "state": "AZ", "overlap_pct": 95.0},
        ],
        senators=[
            {
                "formatted_name": "Sen. Jane Smith (D-AZ)",
                "committees": [
                    {
                        "committee_name": "Senate Committee on Indian Affairs",
                        "role": "member",
                    }
                ],
            }
        ],
        representatives=[
            {
                "formatted_name": "Rep. Carol Lee (D-AZ-02)",
                "committees": [
                    {
                        "committee_name": "House Committee on Appropriations",
                        "role": "member",
                    }
                ],
            }
        ],
        awards=[
            {
                "award_id": "A001",
                "cfda": "15.156",
                "program_id": "bia_tcr",
                "obligation": 500000.0,
                "start_date": "2024-01-01",
                "end_date": "2025-01-01",
            }
        ],
        hazard_profile={
            "fema_nri": {
                "top_hazards": [
                    {
                        "type": "Wildfire",
                        "code": "WFIR",
                        "risk_score": 92.0,
                        "risk_rating": "Very High",
                        "eal_total": 500000.0,
                    },
                ],
                "composite": {"risk_score": 85.5},
            },
            "usfs_wildfire": {"risk_to_homes": 0.85},
        },
    )


@pytest.fixture
def mock_context_zero_awards():
    """TribePacketContext with no awards and no hazard data."""
    return TribePacketContext(
        tribe_id="epa_002",
        tribe_name="Beta Tribe of Montana",
        states=["MT"],
        ecoregions=["northern_rockies"],
        districts=[
            {"district": "MT-02", "state": "MT", "overlap_pct": 80.0},
        ],
        senators=[
            {
                "formatted_name": "Sen. Bob Johnson (R-MT)",
                "committees": [
                    {
                        "committee_name": "Senate Committee on Commerce",
                        "role": "member",
                    }
                ],
            }
        ],
        representatives=[],
        awards=[],
        hazard_profile={},
    )


@pytest.fixture
def mock_program_stable():
    """Program with STABLE CI status and tightened_language."""
    return {
        "id": "bia_tcr",
        "name": "BIA Tribal Climate Resilience",
        "agency": "BIA",
        "ci_status": "STABLE",
        "tightened_language": "Protect the $34.291M base allocation.",
        "advocacy_lever": "Protect/expand the line",
        "description": "Core BIA program for Tribal climate adaptation.",
        "priority": "critical",
        "cfda": "15.156",
        "access_type": "direct",
        "funding_type": "Discretionary",
        "confidence_index": 0.88,
        "federal_home": "Bureau of Indian Affairs",
        "ci_determination": "FY26 funded at $34.291M.",
    }


@pytest.fixture
def mock_program_at_risk():
    """Program with AT_RISK CI status and tightened_language."""
    return {
        "id": "irs_elective_pay",
        "name": "IRS Elective Pay (Direct Pay)",
        "agency": "IRS",
        "ci_status": "AT_RISK",
        "tightened_language": "Defend IRA clean energy direct-pay for Tribes.",
        "advocacy_lever": "Defend IRA elective pay provisions",
        "description": "IRA clean energy tax credit direct pay for Tribes.",
        "priority": "critical",
        "cfda": "21.IRA",
        "access_type": "direct",
        "funding_type": "Mandatory",
        "confidence_index": 0.45,
        "federal_home": "Internal Revenue Service",
        "ci_determination": "At risk of repeal in budget reconciliation.",
    }


@pytest.fixture
def mock_program_no_tightened():
    """Program WITHOUT tightened_language field."""
    return {
        "id": "epa_stag",
        "name": "EPA STAG Tribal Grants",
        "agency": "EPA",
        "ci_status": "UNCERTAIN",
        "advocacy_lever": "Push for dedicated Tribal set-aside",
        "description": "State and Tribal Assistance Grants program.",
        "priority": "high",
        "cfda": "66.123",
        "access_type": "competitive",
        "funding_type": "Discretionary",
    }


@pytest.fixture
def mock_structural_asks():
    """Structural asks from graph_schema.json."""
    return [
        {
            "id": "ask_multi_year",
            "name": "Multi-Year Funding Stability",
            "description": "Shift from annual discretionary to multi-year authorization",
            "target": "Congress",
            "urgency": "FY26",
            "programs": [
                "bia_tcr",
                "epa_gap",
                "doe_indian_energy",
                "noaa_tribal",
                "epa_stag",
            ],
        },
        {
            "id": "ask_match_waivers",
            "name": "Match/Cost-Share Waivers",
            "description": "Categorical Tribal exemptions from cost-share requirements",
            "target": "Congress",
            "urgency": "Immediate",
            "programs": [
                "usda_wildfire",
                "usbr_watersmart",
                "fema_tribal_mitigation",
                "fema_bric",
            ],
        },
    ]


@pytest.fixture
def mock_economic_impact():
    """ProgramEconomicImpact with known values."""
    return ProgramEconomicImpact(
        program_id="bia_tcr",
        total_obligation=500_000.0,
        estimated_impact_low=900_000.0,
        estimated_impact_high=1_200_000.0,
        jobs_estimate_low=4.0,
        jobs_estimate_high=7.5,
        fema_bcr=None,
        is_benchmark=False,
    )


@pytest.fixture
def mock_economic_summary(mock_economic_impact):
    """TribeEconomicSummary wrapping the mock impact."""
    return TribeEconomicSummary(
        tribe_id="epa_001",
        tribe_name="Alpha Tribe of Arizona",
        total_obligation=500_000.0,
        total_impact_low=900_000.0,
        total_impact_high=1_200_000.0,
        total_jobs_low=4.0,
        total_jobs_high=7.5,
        by_program=[mock_economic_impact],
    )


def _make_renderer():
    """Create a fresh Document + StyleManager + HotSheetRenderer."""
    doc = Document()
    sm = StyleManager(doc)
    renderer = HotSheetRenderer(doc, sm)
    return doc, sm, renderer


def _all_text(doc: Document) -> str:
    """Concatenate all paragraph text in the document."""
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Rendering tests
# ---------------------------------------------------------------------------


class TestRenderHotsheet:
    """Core rendering tests for HotSheetRenderer."""

    def test_render_hotsheet_creates_content(
        self, mock_context, mock_program_stable, mock_economic_impact, mock_structural_asks
    ):
        """Rendering one Hot Sheet adds paragraphs to the document."""
        doc, _, renderer = _make_renderer()
        renderer.render_hotsheet(
            mock_context, mock_program_stable, mock_economic_impact, mock_structural_asks
        )
        assert len(doc.paragraphs) > 5

    def test_render_all_hotsheets_count(
        self,
        mock_context,
        mock_program_stable,
        mock_program_at_risk,
        mock_economic_summary,
        mock_structural_asks,
    ):
        """render_all_hotsheets with 3 programs returns 3."""
        doc, _, renderer = _make_renderer()
        programs = [
            mock_program_stable,
            mock_program_at_risk,
            dict(mock_program_stable, id="prog_3", name="Third Prog"),
        ]
        # Add impacts for extra programs
        extra_impact = ProgramEconomicImpact(
            program_id="irs_elective_pay",
            total_obligation=200_000.0,
            estimated_impact_low=360_000.0,
            estimated_impact_high=480_000.0,
            jobs_estimate_low=1.6,
            jobs_estimate_high=3.0,
        )
        summary = TribeEconomicSummary(
            tribe_id="epa_001",
            tribe_name="Alpha Tribe of Arizona",
            by_program=[mock_economic_summary.by_program[0], extra_impact],
        )
        count = renderer.render_all_hotsheets(
            mock_context, programs, summary, mock_structural_asks
        )
        assert count == 3

    def test_render_all_hotsheets_page_breaks(
        self, mock_context, mock_program_stable, mock_economic_summary, mock_structural_asks
    ):
        """Page breaks between Hot Sheets but not after the last."""
        doc, _, renderer = _make_renderer()
        programs = [
            mock_program_stable,
            dict(mock_program_stable, id="prog_2", name="Second Prog"),
            dict(mock_program_stable, id="prog_3", name="Third Prog"),
        ]
        renderer.render_all_hotsheets(
            mock_context, programs, mock_economic_summary, mock_structural_asks
        )
        # Count page breaks in paragraphs (they appear as paragraph with page break before)
        page_break_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type='page']"):
                    page_break_count += 1
        # Should have 2 page breaks (between 1-2 and 2-3, but not after 3)
        assert page_break_count == 2


# ---------------------------------------------------------------------------
# Section-specific tests
# ---------------------------------------------------------------------------


class TestTitleAndStatus:
    """Tests for title and status badge rendering."""

    def test_title_and_status(self, mock_program_stable):
        """Program name and CI badge rendered."""
        doc, _, renderer = _make_renderer()
        renderer._add_title_and_status(mock_program_stable)
        text = _all_text(doc)
        assert "BIA Tribal Climate Resilience" in text
        # Badge circle should be present
        runs_text = []
        for p in doc.paragraphs:
            for r in p.runs:
                runs_text.append(r.text)
        # Unicode circle character
        assert any("\u25CF" in t for t in runs_text)


class TestAwardHistory:
    """Tests for award history table and first-time applicant framing."""

    def test_award_history_with_awards(self, mock_context, mock_program_stable):
        """Table created with correct row count when awards exist."""
        doc, _, renderer = _make_renderer()
        renderer._add_award_history(mock_context, mock_program_stable)
        text = _all_text(doc)
        assert "Award History" in text
        # Should have a table (1 header + 1 data row)
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 2  # header + 1 award
        assert "Total:" in text

    def test_award_history_zero_awards(self, mock_context_zero_awards, mock_program_stable):
        """First-Time Applicant Advantage text rendered for zero-award Tribes."""
        doc, _, renderer = _make_renderer()
        renderer._add_award_history(mock_context_zero_awards, mock_program_stable)
        text = _all_text(doc)
        assert "First-Time Applicant Advantage" in text
        assert "Beta Tribe of Montana" in text
        assert "first-time applicant" in text
        # No table should be created
        assert len(doc.tables) == 0


class TestHazardRelevance:
    """Tests for hazard relevance section."""

    def test_hazard_relevance_with_data(self, mock_context, mock_program_stable):
        """Top hazards listed when hazard data exists."""
        doc, _, renderer = _make_renderer()
        renderer._add_hazard_relevance(mock_context, mock_program_stable)
        text = _all_text(doc)
        assert "Local Hazard Relevance" in text
        assert "Wildfire" in text
        assert "92.0" in text

    def test_hazard_relevance_empty(self, mock_context_zero_awards, mock_program_stable):
        """Section omitted when no hazard data exists."""
        doc, _, renderer = _make_renderer()
        renderer._add_hazard_relevance(mock_context_zero_awards, mock_program_stable)
        text = _all_text(doc)
        assert "Local Hazard Relevance" not in text

    def test_hazard_relevance_usfs_wildfire(self, mock_context, mock_program_stable):
        """USFS wildfire data shown for wildfire-related programs."""
        doc, _, renderer = _make_renderer()
        renderer._add_hazard_relevance(mock_context, mock_program_stable)
        text = _all_text(doc)
        assert "USFS Wildfire Risk to Homes" in text
        assert "0.85" in text


class TestEconomicImpact:
    """Tests for economic impact section."""

    def test_economic_impact_with_data(
        self, mock_context, mock_program_stable, mock_economic_impact
    ):
        """Impact range rendered when economic data exists."""
        doc, _, renderer = _make_renderer()
        renderer._add_economic_impact(
            mock_economic_impact, mock_context, mock_program_stable
        )
        text = _all_text(doc)
        assert "District Economic Impact" in text
        assert "$900,000" in text
        assert "$1,200,000" in text
        assert "BEA RIMS II" in text

    def test_economic_impact_benchmark(self, mock_context, mock_program_stable):
        """Benchmark framing shown for benchmark-based impact."""
        doc, _, renderer = _make_renderer()
        benchmark_impact = ProgramEconomicImpact(
            program_id="bia_tcr",
            total_obligation=150_000.0,
            estimated_impact_low=270_000.0,
            estimated_impact_high=360_000.0,
            jobs_estimate_low=1.2,
            jobs_estimate_high=2.25,
            is_benchmark=True,
        )
        renderer._add_economic_impact(
            benchmark_impact, mock_context, mock_program_stable
        )
        text = _all_text(doc)
        assert "Based on program averages" in text

    def test_economic_impact_none_skipped(self, mock_context, mock_program_stable):
        """Section omitted when economic impact is None."""
        doc, _, renderer = _make_renderer()
        renderer._add_economic_impact(None, mock_context, mock_program_stable)
        text = _all_text(doc)
        assert "District Economic Impact" not in text

    def test_economic_impact_multi_district(self, mock_program_stable, mock_economic_impact):
        """Per-district breakdown shown for multi-district Tribes."""
        ctx = TribePacketContext(
            tribe_id="epa_003",
            tribe_name="Gamma Tribe",
            districts=[
                {"district": "OK-02", "overlap_pct": 60.0},
                {"district": "TX-01", "overlap_pct": 40.0},
            ],
        )
        doc, _, renderer = _make_renderer()
        renderer._add_economic_impact(mock_economic_impact, ctx, mock_program_stable)
        text = _all_text(doc)
        assert "Per-District Impact Allocation" in text
        assert "OK-02" in text
        assert "TX-01" in text

    def test_economic_impact_bcr_narrative(self, mock_context, mock_program_stable):
        """BCR narrative shown for mitigation programs."""
        doc, _, renderer = _make_renderer()
        mit_impact = ProgramEconomicImpact(
            program_id="fema_bric",
            total_obligation=500_000.0,
            estimated_impact_low=900_000.0,
            estimated_impact_high=1_200_000.0,
            jobs_estimate_low=4.0,
            jobs_estimate_high=7.5,
            fema_bcr=4.0,
        )
        renderer._add_economic_impact(
            mit_impact, mock_context, mock_program_stable
        )
        text = _all_text(doc)
        assert "MitSaves" in text or "$4" in text


# ---------------------------------------------------------------------------
# Advocacy language tests
# ---------------------------------------------------------------------------


class TestAdvocacyLanguage:
    """Tests for advocacy position section."""

    def test_advocacy_with_tightened_language(self, mock_program_stable):
        """tightened_language text appears in document."""
        doc, _, renderer = _make_renderer()
        renderer._add_advocacy_language(mock_program_stable)
        text = _all_text(doc)
        assert "Protect the $34.291M base allocation." in text

    def test_advocacy_fallback_no_tightened(self, mock_program_no_tightened):
        """description + advocacy_lever used when no tightened_language."""
        doc, _, renderer = _make_renderer()
        renderer._add_advocacy_language(mock_program_no_tightened)
        text = _all_text(doc)
        assert "Push for dedicated Tribal set-aside" in text
        assert "State and Tribal Assistance Grants" in text

    def test_framing_defensive_at_risk(self, mock_program_at_risk):
        """AT_RISK program gets defensive framing prefix."""
        doc, _, renderer = _make_renderer()
        renderer._add_advocacy_language(mock_program_at_risk)
        text = _all_text(doc)
        assert "Critical:" in text

    def test_framing_growth_stable(self, mock_program_stable):
        """STABLE program gets growth framing prefix."""
        doc, _, renderer = _make_renderer()
        renderer._add_advocacy_language(mock_program_stable)
        text = _all_text(doc)
        assert "Opportunity:" in text

    def test_framing_protective_uncertain(self, mock_program_no_tightened):
        """UNCERTAIN program gets protective framing prefix."""
        doc, _, renderer = _make_renderer()
        renderer._add_advocacy_language(mock_program_no_tightened)
        text = _all_text(doc)
        assert "Watch:" in text


# ---------------------------------------------------------------------------
# Key Ask tests
# ---------------------------------------------------------------------------


class TestKeyAsk:
    """Tests for Key Ask callout block."""

    def test_key_ask_three_lines(
        self, mock_context, mock_program_stable, mock_economic_impact
    ):
        """Key Ask renders ASK, WHY, and IMPACT lines."""
        doc, _, renderer = _make_renderer()
        renderer._add_key_ask(mock_program_stable, mock_context, mock_economic_impact)
        text = _all_text(doc)
        assert "Key Ask" in text
        assert "ASK:" in text
        assert "WHY:" in text
        assert "IMPACT:" in text

    def test_key_ask_no_advocacy_lever(self, mock_context):
        """Key Ask section skipped when program has no advocacy_lever."""
        doc, _, renderer = _make_renderer()
        program = {"id": "test", "name": "Test", "ci_status": "STABLE"}
        renderer._add_key_ask(program, mock_context, None)
        text = _all_text(doc)
        assert "Key Ask" not in text


# ---------------------------------------------------------------------------
# Structural asks tests
# ---------------------------------------------------------------------------


class TestStructuralAsks:
    """Tests for structural policy asks section."""

    def test_structural_asks_linked(
        self, mock_context, mock_program_stable, mock_structural_asks
    ):
        """Matching asks rendered for program."""
        doc, _, renderer = _make_renderer()
        renderer._add_structural_asks(
            mock_program_stable, mock_structural_asks, mock_context
        )
        text = _all_text(doc)
        assert "Structural Policy Asks" in text
        # bia_tcr is in ask_multi_year programs
        assert "Multi-Year Funding Stability" in text
        assert "[FY26]" in text
        assert "Congress" in text

    def test_structural_asks_none_matching(self, mock_context, mock_structural_asks):
        """Section omitted when no asks match the program."""
        doc, _, renderer = _make_renderer()
        program = {"id": "nonexistent_program", "name": "No Match"}
        renderer._add_structural_asks(program, mock_structural_asks, mock_context)
        text = _all_text(doc)
        assert "Structural Policy Asks" not in text


# ---------------------------------------------------------------------------
# Delegation tests
# ---------------------------------------------------------------------------


class TestDelegation:
    """Tests for delegation section and committee-match flags."""

    def test_delegation_rendered(self, mock_context, mock_program_stable):
        """Senator and representative names appear."""
        doc, _, renderer = _make_renderer()
        renderer._add_delegation_info(mock_context, mock_program_stable)
        text = _all_text(doc)
        assert "Your Delegation" in text
        assert "Sen. Jane Smith (D-AZ)" in text
        assert "Rep. Carol Lee (D-AZ-02)" in text

    def test_committee_match_flag(self, mock_context, mock_program_stable):
        """Committee match callout rendered when member sits on relevant committee."""
        doc, _, renderer = _make_renderer()
        renderer._add_delegation_info(mock_context, mock_program_stable)
        text = _all_text(doc)
        # Sen. Jane Smith sits on Indian Affairs -> match
        assert "direct influence on this program" in text
        assert "Indian Affairs" in text

    def test_delegation_empty(self, mock_program_stable):
        """Note shown when no delegation data available."""
        ctx = TribePacketContext(
            tribe_id="epa_099",
            tribe_name="No Delegation Tribe",
            senators=[],
            representatives=[],
        )
        doc, _, renderer = _make_renderer()
        renderer._add_delegation_info(ctx, mock_program_stable)
        text = _all_text(doc)
        assert "Congressional delegation data not available" in text


# ---------------------------------------------------------------------------
# Edge case tests
# ---------------------------------------------------------------------------


class TestEdgeCases:
    """Edge case and normalization tests."""

    def test_all_ci_statuses_render(
        self, mock_context, mock_structural_asks, mock_economic_impact
    ):
        """All 6 CI status values render without error."""
        all_statuses = list(CI_STATUS_COLORS.keys())
        assert len(all_statuses) == 6

        for status in all_statuses:
            doc, _, renderer = _make_renderer()
            program = {
                "id": "test_prog",
                "name": f"Test {status}",
                "agency": "TEST",
                "ci_status": status,
                "advocacy_lever": "Test lever",
                "description": "Test description.",
                "cfda": "99.999",
            }
            # Should not raise
            renderer.render_hotsheet(
                mock_context, program, mock_economic_impact, mock_structural_asks
            )
            text = _all_text(doc)
            assert f"Test {status}" in text

    def test_cfda_string_normalized(self, mock_context):
        """CFDA string field handled correctly in award matching."""
        doc, _, renderer = _make_renderer()
        program = {
            "id": "test_prog",
            "name": "String CFDA Prog",
            "cfda": "15.156",
        }
        # Should not crash
        renderer._add_award_history(mock_context, program)
        text = _all_text(doc)
        # context has award with cfda "15.156" and program_id "bia_tcr"
        # program_id doesn't match, but cfda matches
        assert "Award History" in text

    def test_cfda_list_normalized(self, mock_context):
        """CFDA list field handled correctly."""
        doc, _, renderer = _make_renderer()
        program = {
            "id": "fema_bric",
            "name": "FEMA BRIC",
            "cfda": ["97.047", "97.039"],
        }
        # Should not crash
        renderer._add_award_history(mock_context, program)
        text = _all_text(doc)
        assert "Award History" in text

    def test_cfda_null_handled(self, mock_context):
        """CFDA null field does not crash."""
        doc, _, renderer = _make_renderer()
        program = {
            "id": "no_cfda",
            "name": "No CFDA Prog",
            "cfda": None,
        }
        # Should not crash
        renderer._add_award_history(mock_context, program)
        text = _all_text(doc)
        assert "Award History" in text

    def test_empty_program_dict(self, mock_context, mock_structural_asks):
        """Rendering a nearly-empty program dict does not crash."""
        doc, _, renderer = _make_renderer()
        program = {"id": "empty", "name": "Empty Prog"}
        renderer.render_hotsheet(mock_context, program, None, mock_structural_asks)
        text = _all_text(doc)
        assert "Empty Prog" in text

    def test_methodology_footnote(self):
        """Methodology footnote includes BEA and BLS citations."""
        doc, _, renderer = _make_renderer()
        renderer._add_methodology_footnote()
        text = _all_text(doc)
        assert "BEA" in text or "Bureau of Economic Analysis" in text
        assert "BLS" in text or "Bureau of Labor Statistics" in text

    def test_framing_by_status_covers_all_statuses(self):
        """FRAMING_BY_STATUS has entries for all 6 CI statuses."""
        expected = {"FLAGGED", "AT_RISK", "UNCERTAIN", "STABLE_BUT_VULNERABLE", "STABLE", "SECURE"}
        assert set(FRAMING_BY_STATUS.keys()) == expected

    def test_hazard_profile_sources_key_path(self):
        """Hazard data accessible via sources key wrapper."""
        ctx = TribePacketContext(
            tribe_id="epa_010",
            tribe_name="Sources Key Tribe",
            hazard_profile={
                "sources": {
                    "fema_nri": {
                        "top_hazards": [
                            {"type": "Drought", "code": "DRGT", "risk_score": 55.0, "risk_rating": "Moderate"},
                        ],
                    },
                },
            },
        )
        doc, _, renderer = _make_renderer()
        program = {"id": "bia_tcr", "name": "Test", "cfda": "15.156"}
        renderer._add_hazard_relevance(ctx, program)
        text = _all_text(doc)
        assert "Drought" in text

    def test_render_hotsheet_save_docx(
        self, tmp_path, mock_context, mock_program_stable,
        mock_economic_impact, mock_structural_asks,
    ):
        """Complete Hot Sheet renders and saves to valid .docx file."""
        doc, _, renderer = _make_renderer()
        renderer.render_hotsheet(
            mock_context, mock_program_stable, mock_economic_impact, mock_structural_asks
        )
        output_path = tmp_path / "hotsheet_test.docx"
        doc.save(str(output_path))
        assert output_path.exists()
        assert output_path.stat().st_size > 0

        # Verify it reopens
        reopened = Document(str(output_path))
        assert len(reopened.paragraphs) > 5
