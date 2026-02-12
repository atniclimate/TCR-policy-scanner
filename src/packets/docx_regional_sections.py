"""Regional document section renderers for Doc C and Doc D.

Renders cover page, executive summary, hazard synthesis, award landscape,
delegation (with overlap analysis for Doc C), and appendix sections for
regional documents that aggregate data across all Tribes in a region.

Uses the same StyleManager and formatting patterns as Tribal section
renderers. Regional sections use RegionalContext instead of
TribePacketContext.

Air gap: no organizational names or tool attribution appear in any
rendered section content.
"""

from __future__ import annotations

import logging
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Pt

from src.config import FISCAL_YEAR_SHORT
from src.packets.confidence import section_confidence
from src.utils import format_dollars
from src.packets.docx_styles import (
    COLORS,
    StyleManager,
    apply_zebra_stripe,
    format_header_row,
)

if TYPE_CHECKING:
    from src.packets.doc_types import DocumentTypeConfig
    from src.packets.regional import RegionalContext

logger = logging.getLogger(__name__)


def render_regional_cover_page(
    doc: Document,
    ctx: RegionalContext,
    sm: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the cover page for a regional document.

    Shows region name, Tribe count, state coverage, and audience
    classification. Internal docs (Doc C) include CONFIDENTIAL banner.

    Args:
        doc: python-docx Document to render into.
        ctx: RegionalContext with aggregated region data.
        sm: StyleManager with registered custom styles.
        doc_type_config: DocumentTypeConfig for audience-aware rendering.
    """
    doc.add_paragraph("")

    # Confidentiality banner for internal docs
    if doc_type_config is not None and doc_type_config.confidential:
        banner = doc.add_paragraph(
            "CONFIDENTIAL -- FOR INTERNAL TRIBAL USE ONLY",
            style="HS Small",
        )
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    if doc_type_config is not None:
        title_text = doc_type_config.title_template
    else:
        title_text = f"{FISCAL_YEAR_SHORT} Regional Climate Resilience Overview"
    doc.add_paragraph(title_text, style="Heading 1")

    # Region name subtitle
    doc.add_paragraph(ctx.region_name, style="HS Title")

    # Region summary
    states_str = ", ".join(ctx.states) if ctx.states else "All states"
    doc.add_paragraph(
        f"{ctx.tribe_count} Tribal Nations across {states_str}",
        style="HS Body",
    )

    # Delegation summary
    doc.add_paragraph(
        f"Congressional representation: {ctx.total_senators} Senator(s), "
        f"{ctx.total_representatives} Representative(s)",
        style="HS Body",
    )

    # Generation date
    gen_date = ctx.generated_at[:10] if ctx.generated_at else "N/A"
    doc.add_paragraph(f"Generated: {gen_date}", style="HS Body")

    # Classification notice
    if doc_type_config is not None:
        classification = doc_type_config.footer_template
    else:
        classification = "For Congressional Office Use"
    notice = doc.add_paragraph(classification, style="HS Small")
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    logger.debug("Rendered regional cover page for %s", ctx.region_name)


def render_regional_executive_summary(
    doc: Document,
    ctx: RegionalContext,
    sm: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the regional executive summary section.

    For Doc C (internal): includes coverage gap analysis and delegation
    overlap as shared leverage opportunity. Frames coordination through data.
    For Doc D (congressional): evidence-only resource utilization framing.

    Args:
        doc: python-docx Document to render into.
        ctx: RegionalContext with aggregated region data.
        sm: StyleManager with registered custom styles.
        doc_type_config: DocumentTypeConfig for audience differentiation.
    """
    doc.add_paragraph("Regional Executive Summary", style="Heading 2")

    # Core statistics
    doc.add_paragraph(
        f"{ctx.region_name}: {ctx.tribe_count} Tribal Nations",
        style="HS Title",
    )

    # Hazard profile
    if ctx.top_shared_hazards:
        hazard_names = [
            h["hazard_type"] for h in ctx.top_shared_hazards[:3]
        ]
        doc.add_paragraph(
            f"Top Shared Hazards: {', '.join(hazard_names)}",
            style="HS Body",
        )
    if ctx.composite_risk_score > 0:
        doc.add_paragraph(
            f"Composite Risk Score (regional average): "
            f"{ctx.composite_risk_score:.1f}",
            style="HS Body",
        )

    # Award summary
    if ctx.total_awards > 0:
        doc.add_paragraph(
            f"Total Federal Climate Resilience Investment: "
            f"{format_dollars(ctx.total_awards)} across {ctx.award_coverage} "
            f"Tribal Nations",
            style="HS Body",
        )
    else:
        doc.add_paragraph(
            "No prior federal climate resilience awards recorded "
            "for Tribes in this region.",
            style="HS Body",
        )

    # Economic impact
    econ = ctx.aggregate_economic_impact
    if econ.get("total_high", 0) > 0:
        doc.add_paragraph(
            f"Aggregate Economic Impact: "
            f"{format_dollars(econ.get('total_low', 0))} to "
            f"{format_dollars(econ.get('total_high', 0))}",
            style="HS Body",
        )
        jobs_low = econ.get("total_jobs_low", 0)
        jobs_high = econ.get("total_jobs_high", 0)
        if jobs_high > 0:
            doc.add_paragraph(
                f"Estimated Jobs Supported: {jobs_low:,.0f} to "
                f"{jobs_high:,.0f}",
                style="HS Body",
            )

    # Audience-differentiated content
    is_internal = (
        doc_type_config is not None and doc_type_config.is_internal
    )

    if is_internal:
        # Doc C: coverage gap analysis + delegation overlap as leverage
        doc.add_paragraph(
            "Coverage and Coordination Analysis", style="HS Section"
        )

        # Coverage gaps
        gap_count = len(ctx.tribes_without_awards)
        if gap_count > 0:
            doc.add_paragraph(
                f"Coverage Gap: {gap_count} of {ctx.tribe_count} Tribal "
                f"Nations in this region have received zero federal climate "
                f"resilience awards. Coordinated regional approaches can "
                f"support first-time applicants through shared technical "
                f"assistance and complementary proposals.",
                style="HS Body",
            )

        # Delegation overlap as leverage
        if ctx.delegation_overlap:
            overlap_count = len(ctx.delegation_overlap)
            doc.add_paragraph(
                f"Shared Delegation: {overlap_count} congressional "
                f"member(s) serve multiple Tribal Nations in this region. "
                f"Coordinated asks from multiple Tribes to shared members "
                f"create a stronger signal than individual approaches.",
                style="HS Body",
            )
    else:
        # Doc D: resource utilization framing
        doc.add_paragraph(
            "Federal Investment Overview", style="HS Section"
        )
        doc.add_paragraph(
            f"This overview documents federal climate resilience program "
            f"utilization across {ctx.tribe_count} Tribal Nations in the "
            f"{ctx.short_name} region. Federal investments in Tribal "
            f"resilience infrastructure protect communities and reduce "
            f"future disaster response costs.",
            style="HS Body",
        )

    doc.add_page_break()
    logger.debug(
        "Rendered regional executive summary for %s", ctx.region_name
    )


def render_regional_hazard_synthesis(
    doc: Document,
    ctx: RegionalContext,
    sm: StyleManager,
) -> None:
    """Render the regional hazard synthesis section.

    Shows shared hazards across Tribes with Tribe count and average score.

    Args:
        doc: python-docx Document to render into.
        ctx: RegionalContext with aggregated region data.
        sm: StyleManager with registered custom styles.
    """
    doc.add_paragraph("Regional Hazard Synthesis", style="Heading 2")

    # Core frame from regional config
    if ctx.core_frame:
        doc.add_paragraph(
            f"Regional hazard context: {ctx.core_frame}",
            style="HS Body",
        )

    if not ctx.top_shared_hazards:
        doc.add_paragraph(
            "Hazard data not available for Tribes in this region.",
            style="HS Body",
        )
        doc.add_page_break()
        return

    # Shared hazards table
    doc.add_paragraph(
        f"Top shared hazards across {ctx.tribe_count} Tribal Nations:",
        style="HS Body",
    )
    headers = ["Hazard Type", "Tribes Affected", "Average Risk Score"]
    table = doc.add_table(
        rows=1 + len(ctx.top_shared_hazards), cols=3
    )
    format_header_row(table, headers)

    for i, hazard in enumerate(ctx.top_shared_hazards, start=1):
        row = table.rows[i]
        row.cells[0].text = hazard.get("hazard_type", "Unknown")
        row.cells[1].text = str(hazard.get("tribe_count", 0))
        avg_score = hazard.get("avg_score", 0.0)
        row.cells[2].text = f"{avg_score:.1f}"

    apply_zebra_stripe(table)

    # Coverage note
    if ctx.hazard_coverage < ctx.tribe_count:
        doc.add_paragraph(
            f"Note: {ctx.hazard_coverage} of {ctx.tribe_count} Tribal "
            f"Nations have scored hazard profiles. "
            f"Source: FEMA National Risk Index, county-level data "
            f"aggregated to Tribal geographic boundaries.",
            style="HS Small",
        )
    else:
        doc.add_paragraph(
            "Source: FEMA National Risk Index, county-level data "
            "aggregated to Tribal geographic boundaries.",
            style="HS Small",
        )

    doc.add_page_break()
    logger.debug(
        "Rendered regional hazard synthesis for %s", ctx.region_name
    )


def render_regional_award_landscape(
    doc: Document,
    ctx: RegionalContext,
    sm: StyleManager,
) -> None:
    """Render the regional award landscape section.

    Shows aggregate awards, coverage analysis, and investment gap.

    Args:
        doc: python-docx Document to render into.
        ctx: RegionalContext with aggregated region data.
        sm: StyleManager with registered custom styles.
    """
    doc.add_paragraph("Regional Award Landscape", style="Heading 2")

    # Aggregate summary
    doc.add_paragraph(
        f"Total Federal Climate Resilience Awards: "
        f"{format_dollars(ctx.total_awards)}",
        style="HS Body",
    )
    doc.add_paragraph(
        f"Tribal Nations with Awards: {ctx.award_coverage} of "
        f"{ctx.tribe_count} ({_pct(ctx.award_coverage, ctx.tribe_count)})",
        style="HS Body",
    )

    # Coverage gap
    gap = len(ctx.tribes_without_awards)
    if gap > 0:
        doc.add_paragraph(
            f"Investment Gap: {gap} Tribal Nation(s) in this region have "
            f"received zero federal climate resilience funding through "
            f"tracked programs.",
            style="HS Body",
        )

    doc.add_paragraph(
        "Source: USASpending.gov, tracked federal climate resilience "
        "programs.",
        style="HS Small",
    )

    doc.add_page_break()
    logger.debug(
        "Rendered regional award landscape for %s", ctx.region_name
    )


def render_regional_delegation(
    doc: Document,
    ctx: RegionalContext,
    sm: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the regional congressional delegation section.

    For Doc C (internal): includes delegation overlap analysis showing
    members serving multiple Tribes as shared leverage opportunity.
    For Doc D (congressional): clean delegation summary table.

    Args:
        doc: python-docx Document to render into.
        ctx: RegionalContext with aggregated region data.
        sm: StyleManager with registered custom styles.
        doc_type_config: DocumentTypeConfig for audience differentiation.
    """
    doc.add_paragraph("Regional Congressional Delegation", style="Heading 2")

    # Summary
    doc.add_paragraph(
        f"{ctx.total_senators} Senator(s) and "
        f"{ctx.total_representatives} Representative(s) "
        f"serve Tribal Nations in this region.",
        style="HS Body",
    )

    is_internal = (
        doc_type_config is not None and doc_type_config.is_internal
    )

    # Delegation overlap analysis (Doc C internal only)
    if is_internal and ctx.delegation_overlap:
        doc.add_paragraph(
            "Delegation Overlap Analysis", style="HS Section"
        )
        doc.add_paragraph(
            "The following congressional members serve multiple Tribal "
            "Nations in this region. Coordinated engagement with these "
            "shared delegation members amplifies regional voice.",
            style="HS Body",
        )

        headers = ["Member", "Role", "Tribes Served", "Key Committees"]
        table = doc.add_table(
            rows=1 + len(ctx.delegation_overlap), cols=4
        )
        format_header_row(table, headers)

        for i, member in enumerate(ctx.delegation_overlap, start=1):
            row = table.rows[i]
            row.cells[0].text = member.get("member_name", "Unknown")
            row.cells[1].text = member.get("role", "")
            row.cells[2].text = str(member.get("tribe_count", 0))
            committees = member.get("committees", [])
            row.cells[3].text = "; ".join(
                c for c in committees if c
            )[:100]  # Truncate long committee lists

        apply_zebra_stripe(table)

    elif not is_internal:
        # Doc D: clean summary only
        doc.add_paragraph(
            f"Congressional delegation data covers "
            f"{ctx.congressional_coverage} of {ctx.tribe_count} "
            f"Tribal Nations in this region.",
            style="HS Body",
        )
        doc.add_paragraph(
            "Source: Congress.gov, 119th Congress committee assignments.",
            style="HS Small",
        )

    doc.add_page_break()
    logger.debug(
        "Rendered regional delegation for %s", ctx.region_name
    )


def render_regional_appendix(
    doc: Document,
    ctx: RegionalContext,
    sm: StyleManager,
) -> None:
    """Render the regional appendix listing all Tribes in the region.

    Args:
        doc: python-docx Document to render into.
        ctx: RegionalContext with aggregated region data.
        sm: StyleManager with registered custom styles.
    """
    doc.add_paragraph(
        "Appendix: Tribal Nations in this Region", style="Heading 2"
    )

    if not ctx.tribes:
        doc.add_paragraph(
            "No Tribal Nations assigned to this region.",
            style="HS Body",
        )
        return

    # Tribe listing table
    headers = ["Tribe Name", "State(s)"]
    table = doc.add_table(rows=1 + len(ctx.tribes), cols=2)
    format_header_row(table, headers)

    for i, tribe in enumerate(ctx.tribes, start=1):
        row = table.rows[i]
        row.cells[0].text = tribe.get("tribe_name", "Unknown")
        row.cells[1].text = ", ".join(tribe.get("states", []))

    apply_zebra_stripe(table)

    doc.add_paragraph(
        f"Total: {ctx.tribe_count} Tribal Nations",
        style="HS Small",
    )

    logger.debug(
        "Rendered regional appendix for %s with %d Tribes",
        ctx.region_name,
        ctx.tribe_count,
    )


def render_regional_congressional_section(
    doc: Document,
    ctx: RegionalContext,
    sm: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the regional congressional legislation section for Doc C/D.

    Aggregates bills across all Tribes in the region, groups by bill_id
    so shared bills appear once with affected Tribe count, and shows the
    top 5 bills by combined relevance. Includes per-state delegation
    summary of members who sponsored relevant bills.

    Confidence badge (HIGH/MEDIUM/LOW) appears after the section heading.

    Ends with a page break.

    Args:
        doc: python-docx Document to render into.
        ctx: RegionalContext with aggregated region data (must have
            per-Tribe congressional_intel available through tribes list).
        sm: StyleManager with registered custom styles.
        doc_type_config: DocumentTypeConfig for audience differentiation.
    """
    heading_para = doc.add_paragraph(
        "Regional Legislation Overview", style="Heading 2"
    )

    # Confidence badge -- use earliest scan_date from region's intel
    scan_date = ""
    for tribe_info in ctx.tribes:
        tribe_intel = tribe_info.get("congressional_intel", {})
        td = tribe_intel.get("scan_date", "")
        if td and (not scan_date or td < scan_date):
            scan_date = td

    conf = section_confidence("congress_gov", scan_date)
    badge_run = heading_para.add_run(f" [{conf['level']}]")
    badge_run.font.size = Pt(8)
    badge_run.font.color.rgb = COLORS.muted

    # Aggregate bills across Tribes, deduplicating by bill_id
    bill_agg: dict[str, dict] = {}  # bill_id -> aggregated info
    for tribe_info in ctx.tribes:
        tribe_intel = tribe_info.get("congressional_intel", {})
        for bill in tribe_intel.get("bills", []):
            bid = bill.get("bill_id", "")
            if not bid:
                continue
            if bid not in bill_agg:
                bill_agg[bid] = {
                    "bill": bill,
                    "tribe_count": 0,
                    "combined_relevance": 0.0,
                }
            bill_agg[bid]["tribe_count"] += 1
            bill_agg[bid]["combined_relevance"] += bill.get(
                "relevance_score", 0
            )

    if not bill_agg:
        doc.add_paragraph(
            "No tracked legislation identified for Tribal Nations in "
            "this region during the current session.",
            style="HS Body",
        )
        doc.add_page_break()
        return

    # Sort by combined relevance, take top 5
    ranked = sorted(
        bill_agg.values(),
        key=lambda x: x["combined_relevance"],
        reverse=True,
    )[:5]

    # Render table
    headers = [
        "Bill",
        "Title",
        "Status",
        "Tribes Affected",
        "Relevance",
    ]
    table = doc.add_table(rows=1 + len(ranked), cols=5)
    format_header_row(table, headers)

    for i, entry in enumerate(ranked, start=1):
        bill = entry["bill"]
        row = table.rows[i]
        bill_type = bill.get("bill_type", "")
        bill_number = bill.get("bill_number", "")
        if bill_type and bill_number:
            row.cells[0].text = f"{bill_type} {bill_number}"
        else:
            row.cells[0].text = bill.get("bill_id", "")
        row.cells[1].text = (bill.get("title", "") or "")[:80]
        latest = bill.get("latest_action")
        if latest:
            row.cells[2].text = (latest.get("text", "") or "")[:60]
        else:
            row.cells[2].text = "N/A"
        row.cells[3].text = str(entry["tribe_count"])
        row.cells[4].text = f"{entry['combined_relevance']:.2f}"

    apply_zebra_stripe(table)

    # Per-state delegation summary: members who sponsored relevant bills
    is_internal = (
        doc_type_config is not None and doc_type_config.is_internal
    )
    sponsor_activity: dict[str, set[str]] = {}  # state -> set of names
    for entry in ranked:
        bill = entry["bill"]
        sponsor = bill.get("sponsor", {}) or {}
        sp_state = sponsor.get("state", "")
        sp_name = sponsor.get("name", "")
        if sp_state and sp_name:
            sponsor_activity.setdefault(sp_state, set()).add(sp_name)

    if sponsor_activity:
        doc.add_paragraph(
            "Delegation Sponsorship Activity", style="HS Section"
        )
        for state in sorted(sponsor_activity):
            if state in set(ctx.states):
                names = sorted(sponsor_activity[state])
                doc.add_paragraph(
                    f"{state}: {', '.join(names)}",
                    style="HS Body",
                )

    if is_internal:
        # Doc C: coordination note
        shared_count = sum(
            1 for e in ranked if e["tribe_count"] > 1
        )
        if shared_count > 0:
            doc.add_paragraph(
                f"{shared_count} of {len(ranked)} top bills affect "
                f"multiple Tribal Nations in this region. Coordinated "
                f"engagement on shared legislation amplifies regional "
                f"voice with congressional offices.",
                style="HS Body",
            )

    doc.add_page_break()
    logger.debug(
        "Rendered regional congressional section for %s (%d bills)",
        ctx.region_name,
        len(ranked),
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _pct(numerator: int, denominator: int) -> str:
    """Format a percentage string, handling zero denominator."""
    if denominator == 0:
        return "0%"
    return f"{numerator / denominator * 100:.0f}%"
