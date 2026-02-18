"""Document section renderers for Tribal advocacy packets.

Provides cover page, table of contents, executive summary, congressional
delegation, hazard summary, structural asks, messaging framework, and
appendix section renderers called by DocxEngine.generate() to assemble
the full document.

Each function renders directly into the provided python-docx Document,
adding paragraphs and page breaks. No new Word sections are created
(all content shares the same header/footer).

Section renderers accept an optional ``doc_type_config`` parameter
(DocumentTypeConfig) for audience-differentiated rendering. When provided,
sections adjust content, voice, and included subsections based on the
target audience (internal vs. congressional).

Air gap: no organizational names or tool attribution appear in any
rendered section content.
"""

from __future__ import annotations

import logging
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.config import FISCAL_YEAR_SHORT
from src.packets.context import TribePacketContext
from src.packets.confidence import confidence_level, section_confidence
from src.packets.docx_hotsheet import RELEVANT_COMMITTEE_KEYWORDS
from src.utils import format_dollars
from src.packets.docx_styles import (
    CI_STATUS_LABELS,
    COLORS,
    StyleManager,
    apply_zebra_stripe,
    format_header_row,
)
from src.packets.economic import TribeEconomicSummary

if TYPE_CHECKING:
    from src.packets.doc_types import DocumentTypeConfig

logger = logging.getLogger(__name__)


def render_cover_page(
    document: Document,
    context: TribePacketContext,
    style_manager: StyleManager,
    doc_type_config=None,
) -> None:
    """Render the cover page with Tribe identity and metadata.

    Adds title, Tribe name, delegation summary, geographic info,
    congressional session, generation date, and audience classification.
    Ends with a page break.

    When ``doc_type_config`` is provided, the title and classification
    notice are customized for the target audience.

    Args:
        document: python-docx Document to render into.
        context: TribePacketContext with all Tribe data.
        style_manager: StyleManager with registered custom styles.
        doc_type_config: Optional DocumentTypeConfig for audience-aware
            cover page rendering. When None, uses default fiscal year title.
    """
    # Top spacing
    document.add_paragraph("")

    # Confidentiality banner for internal docs (before title)
    if doc_type_config is not None and doc_type_config.confidential:
        banner_para = document.add_paragraph(
            "CONFIDENTIAL -- FOR INTERNAL TRIBAL USE ONLY",
            style="HS Small",
        )
        banner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Main title -- use doc_type_config title or default
    if doc_type_config is not None:
        title_text = doc_type_config.title_template
    else:
        title_text = f"{FISCAL_YEAR_SHORT} Climate Resilience Program Priorities"

    document.add_paragraph(title_text, style="Heading 1")

    # Tribe name subtitle
    document.add_paragraph(context.tribe_name, style="HS Title")

    # Delegation summary
    num_senators = len(context.senators) if context.senators else 0
    num_reps = len(context.representatives) if context.representatives else 0
    delegation_parts = []
    if num_senators:
        delegation_parts.append(
            f"{num_senators} Senator{'s' if num_senators != 1 else ''}"
        )
    if num_reps:
        delegation_parts.append(
            f"{num_reps} Representative{'s' if num_reps != 1 else ''}"
        )
    delegation_summary = (
        ", ".join(delegation_parts) if delegation_parts else "No delegation data"
    )

    document.add_paragraph(
        f"Prepared for the offices of: {delegation_summary}",
        style="HS Body",
    )

    # Geographic and session metadata
    states_str = ", ".join(context.states) if context.states else "N/A"
    document.add_paragraph(f"State(s): {states_str}", style="HS Body")

    ecoregions_str = ", ".join(context.ecoregions) if context.ecoregions else "N/A"
    document.add_paragraph(f"Ecoregion(s): {ecoregions_str}", style="HS Body")

    document.add_paragraph(
        f"Congressional Session: {context.congress_session}",
        style="HS Body",
    )

    # Generation date (date portion only)
    gen_date = context.generated_at[:10] if context.generated_at else "N/A"
    document.add_paragraph(f"Generated: {gen_date}", style="HS Body")

    # Audience classification notice (centered, small)
    if doc_type_config is not None:
        classification_text = doc_type_config.footer_template
    else:
        classification_text = "For Congressional Office Use"

    conf_para = document.add_paragraph(classification_text, style="HS Small")
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page break after cover
    document.add_page_break()

    logger.debug("Rendered cover page for %s", context.tribe_name)


def render_table_of_contents(
    document: Document,
    programs: list[dict],
    context: TribePacketContext,
    style_manager: StyleManager,
) -> None:
    """Render a table of contents listing program Hot Sheets.

    Lists each program with a numbered entry and inline CI status badge.
    Ends with a page break.

    Args:
        document: python-docx Document to render into.
        programs: List of relevant program dicts (already sorted by relevance).
        context: TribePacketContext for the Tribe name reference.
        style_manager: StyleManager with registered custom styles.
    """
    document.add_paragraph("Program Hot Sheets", style="Heading 2")

    for i, program in enumerate(programs, 1):
        name = program.get("name", "Unknown Program")
        ci_status = program.get("ci_status", "")

        entry_para = document.add_paragraph(style="HS Body")
        entry_para.add_run(f"{i}. {name}")

        if ci_status:
            # Add space then inline badge
            entry_para.add_run("  ")
            label = CI_STATUS_LABELS.get(ci_status, ci_status)
            badge_run = entry_para.add_run(f"[{label}]")
            badge_run.font.size = Pt(8)
            badge_run.font.color.rgb = COLORS.muted

    # Summary note
    document.add_paragraph(
        f"This packet contains {len(programs)} program analyses tailored "
        f"to {context.tribe_name}'s climate risk profile and geographic context.",
        style="HS Body",
    )

    # Page break after TOC
    document.add_page_break()

    logger.debug(
        "Rendered table of contents with %d programs for %s",
        len(programs),
        context.tribe_name,
    )


def render_executive_summary(
    document: Document,
    context: TribePacketContext,
    relevant_programs: list[dict],
    economic_summary: TribeEconomicSummary,
    style_manager: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the executive summary section with Tribe overview.

    Includes Tribe identity, hazard snapshot, funding summary, top programs
    with CI status, delegation snapshot, and structural asks count.

    When ``doc_type_config`` is provided, adjusts narrative voice:
      - Internal (assertive): includes strategic framing paragraph
      - Congressional (objective): evidence-first opening, hazard profile
        summary, award history summary

    Ends with a page break.

    Args:
        document: python-docx Document to render into.
        context: TribePacketContext with all Tribe data.
        relevant_programs: List of relevant program dicts.
        economic_summary: Aggregated economic impact calculations.
        style_manager: StyleManager with registered custom styles.
        doc_type_config: Optional DocumentTypeConfig for audience-aware
            rendering. When None, uses default (backward compatible).
    """
    document.add_paragraph("Executive Summary", style="Heading 2")

    # Tribe identity line
    states_str = ", ".join(context.states) if context.states else "N/A"
    document.add_paragraph(
        f"{context.tribe_name} -- {states_str}", style="HS Title"
    )

    # Audience-specific opening
    is_internal = doc_type_config is not None and doc_type_config.is_internal

    if is_internal:
        # Assertive voice: strategic framing paragraph
        document.add_paragraph(
            f"This strategy document outlines {context.tribe_name}'s federal "
            f"climate resilience funding landscape, identifies strategic "
            f"leverage points across key programs, and recommends sequenced "
            f"advocacy actions for the current congressional session.",
            style="HS Body",
        )

    # Hazard snapshot: top 3 hazards
    hazard_profile = context.hazard_profile or {}
    fema_nri = hazard_profile.get("fema_nri") or hazard_profile.get(
        "sources", {}
    ).get("fema_nri", {})
    top_hazards = fema_nri.get("top_hazards", []) if fema_nri else []

    if top_hazards:
        hazard_names = [h.get("type", "Unknown") for h in top_hazards[:3]]
        document.add_paragraph(
            f"Top Climate Hazards: {', '.join(hazard_names)}", style="HS Body"
        )
    else:
        document.add_paragraph("Hazard profile data pending.", style="HS Body")

    # Funding summary from awards
    total_obligation = 0.0
    for award in context.awards:
        obligation = award.get("obligation", 0.0)
        if isinstance(obligation, (int, float)):
            total_obligation += obligation

    if total_obligation > 0:
        document.add_paragraph(
            f"Total Federal Climate Resilience Funding: {format_dollars(total_obligation)}",
            style="HS Body",
        )
    else:
        document.add_paragraph(
            "No prior federal climate resilience awards on record. "
            "First-time applicant positioning available.",
            style="HS Body",
        )

    # Congressional-facing: add award history summary line
    if doc_type_config is not None and doc_type_config.is_congressional:
        award_count = len(context.awards)
        if award_count > 0:
            document.add_paragraph(
                f"{award_count} federal climate resilience award(s) on record.",
                style="HS Body",
            )

    # Top programs (first 3) with CI status badge
    for program in relevant_programs[:3]:
        name = program.get("name", "Unknown Program")
        ci_status = program.get("ci_status", "")

        prog_para = document.add_paragraph(style="HS Body")
        prog_para.add_run(f"  {name}")

        if ci_status:
            prog_para.add_run("  ")
            label = CI_STATUS_LABELS.get(ci_status, ci_status)
            badge_run = prog_para.add_run(f"[{label}]")
            badge_run.font.size = Pt(8)
            badge_run.font.color.rgb = COLORS.muted

    # Delegation snapshot
    num_senators = len(context.senators) if context.senators else 0
    num_reps = len(context.representatives) if context.representatives else 0
    num_districts = len(context.districts) if context.districts else 0
    document.add_paragraph(
        f"{num_senators} Senator(s) and {num_reps} Representative(s) "
        f"across {num_districts} district(s)",
        style="HS Body",
    )

    # Structural asks count (always 5 per project design)
    document.add_paragraph(
        "5 structural policy asks identified", style="HS Body"
    )

    # Page break
    document.add_page_break()

    logger.info("Rendered executive summary for %s", context.tribe_name)


def render_delegation_section(
    document: Document,
    context: TribePacketContext,
    style_manager: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the congressional delegation section with tabular member listing.

    Lists senators and representatives with their relevant committee
    assignments filtered by RELEVANT_COMMITTEE_KEYWORDS.

    When ``doc_type_config`` is provided:
      - Internal: adds STRATEGIC NOTES subsection after the delegation table
        with per-member approach guidance
      - Congressional: delegation table with committee assignments only,
        no strategic notes

    Ends with a page break.

    Args:
        document: python-docx Document to render into.
        context: TribePacketContext with all Tribe data.
        style_manager: StyleManager with registered custom styles.
        doc_type_config: Optional DocumentTypeConfig for audience-aware
            rendering. When None, uses default (backward compatible).
    """
    document.add_paragraph("Congressional Delegation", style="Heading 2")

    senators = context.senators or []
    representatives = context.representatives or []

    if not senators and not representatives:
        document.add_paragraph(
            "Congressional delegation data not available.", style="HS Body"
        )
        document.add_page_break()
        return

    # Build all member rows
    all_members = []
    for senator in senators:
        name = senator.get("formatted_name", "Unknown")
        state = senator.get("state", "")
        bioguide = senator.get("bioguide_id", "")
        # Filter committees by relevant keywords
        committees = senator.get("committees", [])
        matching = _filter_relevant_committees(committees)
        all_members.append({
            "name": name,
            "role": "Senator",
            "state": state,
            "district": "At-Large",
            "committees": "; ".join(matching),
            "bioguide_id": bioguide,
        })

    for rep in representatives:
        name = rep.get("formatted_name", "Unknown")
        state = rep.get("state", "")
        district = rep.get("district", "")
        bioguide = rep.get("bioguide_id", "")
        committees = rep.get("committees", [])
        matching = _filter_relevant_committees(committees)
        all_members.append({
            "name": name,
            "role": "Representative",
            "state": state,
            "district": str(district) if district else "",
            "committees": "; ".join(matching),
            "bioguide_id": bioguide,
        })

    # Create table
    headers = ["Name", "Role", "State", "District", "Relevant Committees"]
    table = document.add_table(rows=1 + len(all_members), cols=5)
    format_header_row(table, headers)

    for row_idx, member in enumerate(all_members, start=1):
        row = table.rows[row_idx]
        row.cells[0].text = member["name"]
        row.cells[1].text = member["role"]
        row.cells[2].text = member["state"]
        row.cells[3].text = member["district"]
        row.cells[4].text = member["committees"]

    apply_zebra_stripe(table)

    # Summary: count of members with relevant committee assignments
    members_with_committees = sum(
        1 for m in all_members if m["committees"]
    )
    document.add_paragraph(
        f"{members_with_committees} members with relevant committee assignments",
        style="HS Body",
    )

    # Congressional intel: committee assignments and bill sponsorship activity
    delegation_activity = context.congressional_intel.get(
        "delegation_activity", {}
    )
    if delegation_activity:
        document.add_paragraph(
            "Delegation Legislative Activity", style="HS Section"
        )
        for member in all_members:
            member_key = member.get("bioguide_id", member["name"])
            activity = delegation_activity.get(member_key, {})
            sponsored = activity.get("sponsored", [])
            cosponsored = activity.get("cosponsored", [])
            relevant_committees = activity.get("committees", [])

            activity_parts: list[str] = []
            if sponsored:
                bill_labels = [b.get("bill_id", "") for b in sponsored[:3]]
                labels_str = ", ".join(b for b in bill_labels if b)
                if labels_str:
                    activity_parts.append(f"Sponsored: {labels_str}")
            if cosponsored:
                activity_parts.append(
                    f"{len(cosponsored)} bill(s) cosponsored"
                )
            if relevant_committees:
                activity_parts.append(
                    f"Committees: {'; '.join(relevant_committees[:3])}"
                )
            if activity_parts:
                document.add_paragraph(
                    f"{member['name']}: {' | '.join(activity_parts)}",
                    style="HS Body",
                )

    # Internal-only: STRATEGIC NOTES subsection
    is_internal = (
        doc_type_config is not None and doc_type_config.include_member_approach
    )
    if is_internal:
        document.add_paragraph("Strategic Notes", style="HS Section")
        for member in all_members:
            if member["committees"]:
                document.add_paragraph(
                    f"{member['name']}: Sits on {member['committees']}. "
                    f"Direct influence on relevant programs.",
                    style="HS Body",
                )

    # Page break
    document.add_page_break()

    total = len(all_members)
    logger.info(
        "Rendered delegation section with %d members for %s",
        total,
        context.tribe_name,
    )


def render_hazard_summary(
    document: Document,
    context: TribePacketContext,
    style_manager: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the climate hazard profile summary section.

    Displays FEMA NRI composite ratings, top 5 hazards in a table, and
    USFS wildfire risk data if available.

    When ``doc_type_config`` is provided:
      - Internal: includes hazard-to-program alignment narrative
      - Congressional: presents hazard data factually with NRI scores
        and source citation

    Both versions use real hazard data when available.

    Ends with a page break.

    Args:
        document: python-docx Document to render into.
        context: TribePacketContext with all Tribe data.
        style_manager: StyleManager with registered custom styles.
        doc_type_config: Optional DocumentTypeConfig for audience-aware
            rendering. When None, uses default (backward compatible).
    """
    document.add_paragraph("Climate Hazard Profile", style="Heading 2")

    # Extract hazard data -- same nesting pattern as docx_hotsheet
    hazard_profile = context.hazard_profile or {}
    fema_nri = hazard_profile.get("fema_nri") or hazard_profile.get(
        "sources", {}
    ).get("fema_nri", {})
    usfs_data = hazard_profile.get("usfs_wildfire") or hazard_profile.get(
        "sources", {}
    ).get("usfs_wildfire", {})

    if not fema_nri and not usfs_data:
        document.add_paragraph(
            "Hazard profile data not available for this Tribe.",
            style="HS Body",
        )
        document.add_page_break()
        return

    # Congressional: source citation up front
    is_congressional = (
        doc_type_config is not None and doc_type_config.is_congressional
    )
    if is_congressional:
        document.add_paragraph(
            "Source: FEMA National Risk Index, county-level data aggregated "
            "to Tribal geographic boundaries.",
            style="HS Small",
        )

    # FEMA NRI composite section
    composite = fema_nri.get("composite", {}) if fema_nri else {}
    if composite:
        risk_rating = composite.get("risk_rating", "N/A")
        risk_score = composite.get("risk_score")
        sovi_rating = composite.get("sovi_rating", "N/A")
        sovi_score = composite.get("sovi_score")

        if risk_score is not None:
            document.add_paragraph(
                f"Overall Risk Rating: {risk_rating} (Score: {risk_score:.1f})",
                style="HS Body",
            )
        else:
            document.add_paragraph(
                f"Overall Risk Rating: {risk_rating}", style="HS Body"
            )

        if sovi_score is not None:
            document.add_paragraph(
                f"Social Vulnerability: {sovi_rating} (Score: {sovi_score:.1f})",
                style="HS Body",
            )
        else:
            document.add_paragraph(
                f"Social Vulnerability: {sovi_rating}", style="HS Body"
            )

    # Top 5 hazards table
    top_hazards = fema_nri.get("top_hazards", []) if fema_nri else []
    if top_hazards:
        headers = ["Hazard Type", "Risk Score", "Risk Rating", "Expected Annual Loss"]
        display_hazards = top_hazards[:5]
        table = document.add_table(rows=1 + len(display_hazards), cols=4)
        format_header_row(table, headers)

        for row_idx, hazard in enumerate(display_hazards, start=1):
            row = table.rows[row_idx]
            row.cells[0].text = hazard.get("type", "Unknown")
            risk_score = hazard.get("risk_score")
            row.cells[1].text = f"{risk_score:.1f}" if risk_score is not None else "N/A"
            row.cells[2].text = hazard.get("risk_rating", "N/A")
            eal = hazard.get("eal_total")
            row.cells[3].text = format_dollars(eal) if eal is not None else "N/A"

        apply_zebra_stripe(table)

    # Internal-only: hazard-to-program alignment narrative
    is_internal = doc_type_config is not None and doc_type_config.is_internal
    if is_internal and top_hazards:
        document.add_paragraph(
            "Hazard-to-Program Alignment", style="HS Section"
        )
        hazard_names = [h.get("type", "Unknown") for h in top_hazards[:3]]
        document.add_paragraph(
            f"{context.tribe_name}'s top hazard exposures "
            f"({', '.join(hazard_names)}) align directly with federal "
            f"programs designed to address these risks. Strategic advocacy "
            f"should emphasize this alignment when engaging congressional "
            f"offices.",
            style="HS Body",
        )

    # USFS wildfire section
    if usfs_data:
        document.add_paragraph(
            "USFS Wildfire Risk Assessment", style="HS Section"
        )
        risk_to_homes = usfs_data.get("risk_to_homes")
        likelihood = usfs_data.get("likelihood")
        parts = []
        if risk_to_homes is not None:
            parts.append(f"Risk to Homes: {risk_to_homes:.2f}")
        if likelihood is not None:
            parts.append(f"Fire Likelihood: {likelihood:.2f}")
        if parts:
            document.add_paragraph(" | ".join(parts), style="HS Body")

    # Page break
    document.add_page_break()

    logger.info("Rendered hazard summary for %s", context.tribe_name)


def render_structural_asks_section(
    document: Document,
    structural_asks: list[dict],
    context: TribePacketContext,
    style_manager: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the standalone structural policy asks section.

    Lists all structural asks with descriptions, targets, urgency badges,
    programs advanced, and evidence lines connecting to Tribe data.

    When ``doc_type_config`` is provided:
      - Strategic framing (internal): assertive, action-oriented intro
      - Policy recommendation framing (congressional): measured,
        evidence-based intro

    Ends with a page break.

    Args:
        document: python-docx Document to render into.
        structural_asks: List of structural ask dicts.
        context: TribePacketContext with all Tribe data.
        style_manager: StyleManager with registered custom styles.
        doc_type_config: Optional DocumentTypeConfig for audience-aware
            rendering. When None, uses default (backward compatible).
    """
    document.add_paragraph("Structural Policy Asks", style="Heading 2")

    if not structural_asks:
        document.add_paragraph(
            "No structural policy asks defined.", style="HS Body"
        )
        document.add_page_break()
        return

    # Audience-calibrated intro
    is_strategic = (
        doc_type_config is not None
        and doc_type_config.structural_asks_framing == "strategic"
    )
    if is_strategic:
        document.add_paragraph(
            f"The following structural policy changes are priority advocacy "
            f"targets for {context.tribe_name}. Each ask has been identified "
            f"as high-impact based on current program vulnerabilities and "
            f"congressional dynamics.",
            style="HS Body",
        )
    else:
        document.add_paragraph(
            f"The following structural policy changes would advance climate "
            f"resilience funding for {context.tribe_name} and similarly "
            f"situated Tribes.",
            style="HS Body",
        )

    for ask in structural_asks:
        ask_name = ask.get("name", "")
        ask_urgency = ask.get("urgency", "")
        ask_desc = ask.get("description", "")
        ask_target = ask.get("target", "")
        ask_programs = ask.get("programs", [])

        # Ask name (bold) with urgency badge
        heading_parts = [ask_name]
        if ask_urgency:
            heading_parts.append(f"[{ask_urgency}]")
        heading_para = document.add_paragraph(style="HS Body")
        heading_run = heading_para.add_run(" ".join(heading_parts))
        heading_run.bold = True

        # Description
        if ask_desc:
            document.add_paragraph(ask_desc, style="HS Body")

        # Target
        if ask_target:
            document.add_paragraph(
                f"Target: {ask_target}", style="HS Body"
            )

        # Programs advanced
        if ask_programs:
            program_names = ", ".join(ask_programs)
            document.add_paragraph(
                f"Programs advanced: {program_names}", style="HS Body"
            )

        # Evidence line from Tribe data
        evidence = _build_ask_evidence(ask, context)
        if evidence:
            para = document.add_paragraph(style="HS Small")
            para.add_run(f"Evidence: {evidence}")

    # Page break
    document.add_page_break()

    logger.info(
        "Rendered %d structural asks for %s",
        len(structural_asks),
        context.tribe_name,
    )


def render_bill_intelligence_section(
    document: Document,
    context: TribePacketContext,
    style_manager: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the congressional bill intelligence section.

    Displays tracked legislation relevant to the Tribe's programs with
    audience-differentiated content:

      - Doc A (internal): Full briefing for top 2-3 bills with talking
        points, timing, and strategy notes. Remaining bills get compact
        cards. This is the advocacy preparation version.
      - Doc B (congressional): Facts-only bill listing with zero strategy
        language. Filtered to bills with relevance_score >= 0.5.
      - Default (no dtc): Renders compact cards for all bills.

    Confidence badge (HIGH/MEDIUM/LOW) appears after the section heading
    based on data freshness from congressional_intel scan_date.

    Ends with a page break.

    Args:
        document: python-docx Document to render into.
        context: TribePacketContext with congressional_intel data.
        style_manager: StyleManager with registered custom styles.
        doc_type_config: Optional DocumentTypeConfig for audience-aware
            rendering.
    """
    is_internal = doc_type_config is not None and doc_type_config.is_internal
    is_congressional = (
        doc_type_config is not None and doc_type_config.is_congressional
    )

    # Section heading with audience differentiation
    if is_internal:
        heading_text = "Legislation Affecting Your Programs"
    elif is_congressional:
        heading_text = "Relevant Legislation"
    else:
        heading_text = "Relevant Legislation"

    heading_para = document.add_paragraph(heading_text, style="Heading 2")

    # Confidence badge after heading
    intel = context.congressional_intel or {}
    scan_date = intel.get("scan_date") or intel.get("metadata", {}).get(
        "built_at", ""
    )
    conf = section_confidence("congress_gov", scan_date)
    badge_text = f" [{conf['level']}]"
    badge_run = heading_para.add_run(badge_text)
    badge_run.font.size = Pt(8)
    badge_run.font.color.rgb = COLORS.muted

    # Get bills
    bills = intel.get("bills", [])
    if not bills:
        document.add_paragraph(
            "No tracked legislation identified in the current session.",
            style="HS Body",
        )
        document.add_page_break()
        logger.debug(
            "Bill intelligence: no bills for %s", context.tribe_name
        )
        return

    # Sort by relevance_score descending
    sorted_bills = sorted(
        bills,
        key=lambda b: b.get("relevance_score", 0),
        reverse=True,
    )

    if is_congressional:
        # Doc B: facts only, filtered to relevance >= 0.5
        filtered = [
            b for b in sorted_bills if b.get("relevance_score", 0) >= 0.5
        ]
        if not filtered:
            document.add_paragraph(
                "No legislation with direct program relevance identified "
                "in the current session.",
                style="HS Body",
            )
        else:
            for bill in filtered:
                _render_bill_facts_only(document, bill)
    elif is_internal:
        # Doc A: top 2-3 get full briefing, rest get compact cards
        full_count = min(3, len(sorted_bills))
        for bill in sorted_bills[:full_count]:
            _render_bill_full_briefing(document, bill, context)

        remaining = sorted_bills[full_count:]
        if remaining:
            document.add_paragraph(
                "Additional Tracked Legislation", style="HS Section"
            )
            for bill in remaining:
                _render_bill_compact_card(document, bill)
    else:
        # Default (no doc_type_config): compact cards for all
        for bill in sorted_bills:
            _render_bill_compact_card(document, bill)

    document.add_page_break()

    logger.info(
        "Rendered bill intelligence (%d bills) for %s",
        len(sorted_bills),
        context.tribe_name,
    )


# ---------------------------------------------------------------------------
# Bill rendering helpers (internal to this module)
# ---------------------------------------------------------------------------


def _render_bill_full_briefing(
    document: Document,
    bill: dict,
    context: TribePacketContext,
) -> None:
    """Render a full briefing card for a high-relevance bill (Doc A).

    Includes bill number/title, status, sponsor, cosponsors, committees,
    affected programs, talking points, and timing/urgency note.

    Args:
        document: python-docx Document to render into.
        bill: Bill dict from congressional_intel.
        context: TribePacketContext for program/hazard cross-references.
    """
    bill_id = bill.get("bill_id", "")
    title = bill.get("title", "Unknown")
    bill_type = bill.get("bill_type", "")
    bill_number = bill.get("bill_number", "")

    # Build display label (e.g. "HR 1234" or use bill_id)
    if bill_type and bill_number:
        display_label = f"{bill_type} {bill_number}"
    else:
        display_label = bill_id

    # Subheading: bold bill number + title
    sub_para = document.add_paragraph(style="HS Body")
    sub_run = sub_para.add_run(f"{display_label}: {title}")
    sub_run.bold = True

    # Status line
    latest_action = bill.get("latest_action")
    if latest_action:
        action_date = latest_action.get("date", "")
        action_text = latest_action.get("text", "")
        if action_date:
            action_date = action_date[:10]  # date portion only
        document.add_paragraph(
            f"Status: {action_text} ({action_date})" if action_date
            else f"Status: {action_text}",
            style="HS Body",
        )

    # Sponsor + cosponsor count
    sponsor = bill.get("sponsor")
    cosponsor_count = bill.get("cosponsor_count", 0)
    if sponsor:
        sponsor_name = sponsor.get("name", "Unknown")
        sponsor_party = sponsor.get("party", "")
        sponsor_state = sponsor.get("state", "")
        sponsor_parts = [sponsor_name]
        if sponsor_party:
            sponsor_parts.append(f"({sponsor_party})")
        if sponsor_state:
            sponsor_parts.append(f"- {sponsor_state}")
        document.add_paragraph(
            f"Sponsor: {' '.join(sponsor_parts)} | "
            f"{cosponsor_count} cosponsor(s)",
            style="HS Body",
        )

    # Committee referral
    committees = bill.get("committees", [])
    if committees:
        committee_names = [
            c.get("name", c.get("committee_name", "")) for c in committees
        ]
        committee_names = [n for n in committee_names if n]
        if committee_names:
            document.add_paragraph(
                f"Committees: {'; '.join(committee_names)}",
                style="HS Body",
            )

    # Affected programs
    matched = bill.get("matched_programs", [])
    if matched:
        document.add_paragraph(
            f"Affected Programs: {', '.join(matched)}",
            style="HS Body",
        )

    # Talking points (Doc A only -- derived from subjects + program relevance)
    subjects = bill.get("subjects", [])
    _render_talking_points(document, bill, subjects, context)

    # Timing/urgency note
    _render_timing_note(document, latest_action)


def _render_bill_facts_only(document: Document, bill: dict) -> None:
    """Render a facts-only bill entry for Doc B (congressional).

    Shows bill number, title, status, sponsor, and affected programs.
    Contains ZERO strategy, talking points, timing, or urgency language.

    Args:
        document: python-docx Document to render into.
        bill: Bill dict from congressional_intel.
    """
    bill_type = bill.get("bill_type", "")
    bill_number = bill.get("bill_number", "")
    title = bill.get("title", "Unknown")
    bill_id = bill.get("bill_id", "")

    if bill_type and bill_number:
        display_label = f"{bill_type} {bill_number}"
    else:
        display_label = bill_id

    # Bill number + title
    name_para = document.add_paragraph(style="HS Body")
    name_run = name_para.add_run(f"{display_label}: {title}")
    name_run.bold = True

    # Status
    latest_action = bill.get("latest_action")
    if latest_action:
        action_date = latest_action.get("date", "")
        action_text = latest_action.get("text", "")
        if action_date:
            action_date = action_date[:10]
        status_text = (
            f"Status: {action_text} ({action_date})" if action_date
            else f"Status: {action_text}"
        )
        document.add_paragraph(status_text, style="HS Body")

    # Sponsor name
    sponsor = bill.get("sponsor")
    if sponsor:
        document.add_paragraph(
            f"Sponsor: {sponsor.get('name', 'Unknown')}",
            style="HS Body",
        )

    # Affected programs
    matched = bill.get("matched_programs", [])
    if matched:
        document.add_paragraph(
            f"Affected Programs: {', '.join(matched)}",
            style="HS Body",
        )


def _render_bill_compact_card(document: Document, bill: dict) -> None:
    """Render a compact single-card bill entry.

    Shows bill number + title (bold), status + affected programs on one
    line, and a one-sentence impact summary.

    Args:
        document: python-docx Document to render into.
        bill: Bill dict from congressional_intel.
    """
    bill_type = bill.get("bill_type", "")
    bill_number = bill.get("bill_number", "")
    title = bill.get("title", "Unknown")
    bill_id = bill.get("bill_id", "")

    if bill_type and bill_number:
        display_label = f"{bill_type} {bill_number}"
    else:
        display_label = bill_id

    # Bill number + title (bold)
    name_para = document.add_paragraph(style="HS Body")
    name_run = name_para.add_run(f"{display_label}: {title}")
    name_run.bold = True

    # Status + affected programs (single line)
    parts = []
    latest_action = bill.get("latest_action")
    if latest_action:
        action_text = latest_action.get("text", "")
        if action_text:
            parts.append(action_text)
    matched = bill.get("matched_programs", [])
    if matched:
        parts.append(f"Programs: {', '.join(matched)}")
    if parts:
        document.add_paragraph(" | ".join(parts), style="HS Small")

    # Impact summary from subjects
    subjects = bill.get("subjects", [])
    if subjects:
        document.add_paragraph(
            f"Subject areas: {', '.join(subjects[:5])}",
            style="HS Small",
        )


def _render_talking_points(
    document: Document,
    bill: dict,
    subjects: list[str],
    context: TribePacketContext,
) -> None:
    """Render 2-3 talking points derived from bill subjects and Tribe context.

    Doc A internal only -- these talking points help Tribal leadership
    prepare for congressional engagement around this legislation.

    Args:
        document: python-docx Document to render into.
        bill: Bill dict with matched_programs and subjects.
        subjects: Bill subject area list.
        context: TribePacketContext for hazard/program cross-references.
    """
    points: list[str] = []

    # Point 1: program alignment
    matched = bill.get("matched_programs", [])
    if matched:
        points.append(
            f"This bill directly affects {', '.join(matched[:3])}, "
            f"programs that {context.tribe_name} relies on for climate "
            f"resilience funding."
        )

    # Point 2: hazard relevance
    hazard_profile = context.hazard_profile or {}
    fema_nri = hazard_profile.get("fema_nri") or hazard_profile.get(
        "sources", {}
    ).get("fema_nri", {})
    top_hazards = fema_nri.get("top_hazards", []) if fema_nri else []
    if top_hazards:
        top_hazard = top_hazards[0].get("type", "")
        if top_hazard:
            points.append(
                f"{context.tribe_name}'s top hazard exposure "
                f"({top_hazard}) underscores the importance of "
                f"sustained federal investment in these programs."
            )

    # Point 3: subject area context
    if subjects:
        points.append(
            f"Key subject areas ({', '.join(subjects[:3])}) align with "
            f"the Tribe's priority program areas."
        )

    if points:
        document.add_paragraph("Talking Points", style="HS Section")
        for point in points[:3]:
            document.add_paragraph(point, style="HS Body")


def _render_timing_note(
    document: Document,
    latest_action: dict | None,
    doc_type_config: "DocumentTypeConfig | None" = None,
) -> None:
    """Render a timing/urgency note based on latest action recency.

    **Doc A (internal) ONLY.** Contains strategy-adjacent language
    ('Active movement', 'Monitoring') that must never appear in
    congressional documents. When ``doc_type_config`` indicates a
    congressional audience, this function returns immediately without
    rendering anything.

    Categorizes bills by latest action date into immediate (< 30 days),
    active (30-90 days), or monitoring (> 90 days).

    Args:
        document: python-docx Document to render into.
        latest_action: Latest action dict with date and text fields.
        doc_type_config: Optional DocumentTypeConfig. If congressional,
            timing note is suppressed entirely.
    """
    # Explicit audience guard: timing notes are Doc A only
    if doc_type_config is not None and doc_type_config.is_congressional:
        return
    if not latest_action:
        return

    action_date = latest_action.get("date", "")
    if not action_date:
        return

    try:
        from datetime import datetime, timezone

        action_dt = datetime.fromisoformat(action_date)
        if action_dt.tzinfo is None:
            action_dt = action_dt.replace(tzinfo=timezone.utc)
        now = datetime.now(timezone.utc)
        days_ago = (now - action_dt).days

        if days_ago < 30:
            note = "Active movement -- recent action within 30 days."
        elif days_ago < 90:
            note = "Moderate activity -- last action within 90 days."
        else:
            note = "Monitoring -- no recent movement."

        timing_para = document.add_paragraph(style="HS Small")
        timing_run = timing_para.add_run(f"Timing: {note}")
        timing_run.italic = True
    except (ValueError, TypeError):
        pass


# ---------------------------------------------------------------------------
# Internal helpers for new section renderers
# ---------------------------------------------------------------------------


def _filter_relevant_committees(committees: list[dict]) -> list[str]:
    """Filter committee list to those matching RELEVANT_COMMITTEE_KEYWORDS.

    Args:
        committees: List of committee dicts with ``committee_name`` key.

    Returns:
        List of matching committee names.
    """
    matching = []
    for committee in committees:
        committee_name = committee.get("committee_name", "")
        for keyword in RELEVANT_COMMITTEE_KEYWORDS:
            if keyword.lower() in committee_name.lower():
                matching.append(committee_name)
                break
    return matching


def _build_ask_evidence(ask: dict, context: TribePacketContext) -> str:
    """Build a brief evidence line connecting an ask to Tribe data.

    Checks for matching awards and top hazards to generate evidence text.

    Args:
        ask: Structural ask dict with ``programs`` list.
        context: TribePacketContext with awards and hazard data.

    Returns:
        Evidence string, or empty string if no evidence found.
    """
    parts: list[str] = []

    # Check if any ask programs match context awards
    ask_programs = set(ask.get("programs", []))
    for award in context.awards:
        if award.get("program_id") in ask_programs:
            parts.append("Active awards in affected programs")
            break

    # Check hazard profile for top hazard type
    hazard_profile = context.hazard_profile or {}
    fema_nri = hazard_profile.get("fema_nri") or hazard_profile.get(
        "sources", {}
    ).get("fema_nri", {})
    top_hazards = fema_nri.get("top_hazards", []) if fema_nri else []
    if top_hazards:
        top_type = top_hazards[0].get("type", "")
        if top_type:
            parts.append(f"{top_type} risk profile")

    return "; ".join(parts) if parts else ""


def render_messaging_framework(
    document: Document,
    context: TribePacketContext,
    style_manager: StyleManager,
) -> None:
    """Render the messaging framework section (Doc A only).

    Provides audience-calibrated talking points for Tribal leadership
    when engaging congressional offices, committee staff, and media.
    This section contains strategic framing that should never appear
    in congressional-facing documents.

    Ends with a page break.

    Args:
        document: python-docx Document to render into.
        context: TribePacketContext with all Tribe data.
        style_manager: StyleManager with registered custom styles.
    """
    document.add_paragraph("Messaging Framework", style="Heading 2")

    document.add_paragraph(
        f"Audience-calibrated talking points for {context.tribe_name}'s "
        f"advocacy engagements during the current congressional session.",
        style="HS Body",
    )

    # Core message
    document.add_paragraph("Core Message", style="HS Section")
    states_str = ", ".join(context.states) if context.states else "our region"

    # Extract top hazard for contextual messaging
    hazard_profile = context.hazard_profile or {}
    fema_nri = hazard_profile.get("fema_nri") or hazard_profile.get(
        "sources", {}
    ).get("fema_nri", {})
    top_hazards = fema_nri.get("top_hazards", []) if fema_nri else []
    if top_hazards:
        top_hazard = top_hazards[0].get("type", "extreme weather")
    else:
        top_hazard = "extreme weather"

    document.add_paragraph(
        f"{context.tribe_name} faces significant {top_hazard.lower()} risk. "
        f"Congress has funded programs that address this risk, but "
        f"administrative delivery gaps and structural barriers prevent "
        f"these resources from reaching Tribal communities in {states_str}.",
        style="HS Body",
    )

    # Congressional framing
    document.add_paragraph(
        "Congressional Office Framing", style="HS Section"
    )
    document.add_paragraph(
        "Lead with bipartisan infrastructure investment and return on "
        "investment. Emphasize that Congress made the funding decisions; "
        "the ask is for oversight and delivery accountability.",
        style="HS Body",
    )

    # Committee-specific guidance
    document.add_paragraph(
        "Committee-Specific Guidance", style="HS Section"
    )
    document.add_paragraph(
        "Appropriations: Frame around funding continuity and program "
        "effectiveness metrics. Indian Affairs: Frame around trust "
        "responsibility, self-determination, and direct service delivery. "
        "Natural Resources / Energy: Frame around infrastructure "
        "resilience and economic development.",
        style="HS Body",
    )

    document.add_page_break()

    logger.info(
        "Rendered messaging framework for %s", context.tribe_name
    )


def render_appendix(
    document: Document,
    omitted_programs: list[dict],
    context: TribePacketContext,
    style_manager: StyleManager,
    doc_type_config: DocumentTypeConfig | None = None,
) -> None:
    """Render the appendix listing omitted (lower-priority) programs.

    If no programs were omitted, adds a brief note and returns.
    Otherwise lists each omitted program with description, access/funding
    type, and CI status.

    Args:
        document: python-docx Document to render into.
        omitted_programs: List of program dicts not included in Hot Sheets.
        context: TribePacketContext for Tribe name reference.
        style_manager: StyleManager with registered custom styles.
        doc_type_config: Optional DocumentTypeConfig for audience-aware
            rendering. Currently unused but accepted for signature
            consistency with other section renderers.
    """
    document.add_paragraph("Appendix: Additional Programs", style="Heading 2")

    if not omitted_programs:
        document.add_paragraph(
            "All tracked programs are included in this packet.",
            style="HS Body",
        )
        logger.debug("Appendix: no omitted programs for %s", context.tribe_name)
        return

    # Intro paragraph
    document.add_paragraph(
        f"The following programs were assessed as lower priority for "
        f"{context.tribe_name} based on hazard profile and geographic "
        f"relevance. They may still be relevant for specific project needs.",
        style="HS Body",
    )

    for program in omitted_programs:
        name = program.get("name", "Unknown Program")
        ci_status = program.get("ci_status", "")
        description = program.get("description", "")
        access_type = program.get("access_type", "")
        funding_type = program.get("funding_type", "")
        agency = program.get("agency", "the administering agency")

        # Program name (bold) + CI status inline
        name_para = document.add_paragraph(style="HS Body")
        name_run = name_para.add_run(name)
        name_run.bold = True

        if ci_status:
            label = CI_STATUS_LABELS.get(ci_status, ci_status)
            name_para.add_run(f"  [{label}]")

        # Description
        if description:
            document.add_paragraph(description, style="HS Body")

        # Access type and funding type
        meta_parts = []
        if access_type:
            meta_parts.append(
                f"Access: {access_type.replace('_', ' ').title()}"
            )
        if funding_type:
            meta_parts.append(f"Funding: {funding_type}")
        if meta_parts:
            document.add_paragraph(" | ".join(meta_parts), style="HS Small")

        # Contact info
        document.add_paragraph(
            f"Learn more: Contact your program specialist or visit "
            f"{agency} for current application information.",
            style="HS Small",
        )

    logger.debug(
        "Rendered appendix with %d omitted programs for %s",
        len(omitted_programs),
        context.tribe_name,
    )


def render_change_tracking(
    document: Document,
    changes: list[dict],
    previous_date: str,
    style_manager: StyleManager,
) -> None:
    """Render 'Since Last Packet' section showing changes between generations.

    Groups changes by type with human-readable descriptions. If the changes
    list is empty, the section is NOT rendered (omitted entirely). The caller
    is responsible for adding a page break before this section if needed.

    Args:
        document: python-docx Document to render into.
        changes: List of change dicts with ``type`` and ``description`` keys.
        previous_date: ISO date string of the previous packet generation.
        style_manager: StyleManager with registered custom styles.
    """
    if not changes:
        return

    date_display = previous_date[:10] if previous_date else "unknown"
    # Use add_paragraph with Heading 2 style (not add_heading) to ensure
    # the StyleManager-configured Arial font is applied consistently.
    # See ACCURACY-006: add_heading bypasses style overrides.
    document.add_paragraph(
        f"Since Last Packet (generated {date_display})", style="Heading 2"
    )

    type_labels = {
        "ci_status_change": "Confidence Index Changes",
        "new_award": "New Awards",
        "award_total_change": "Funding Changes",
        "advocacy_goal_shift": "Advocacy Position Shifts",
        "new_threat": "New Hazard Threats",
    }

    # Group changes by type
    grouped: dict[str, list[dict]] = {}
    for change in changes:
        ctype = change.get("type", "other")
        if ctype not in grouped:
            grouped[ctype] = []
        grouped[ctype].append(change)

    for ctype, group in grouped.items():
        label = type_labels.get(ctype, ctype.replace("_", " ").title())
        p = document.add_paragraph()
        run = p.add_run(label)
        run.bold = True

        for change in group:
            description = change.get("description", str(change))
            document.add_paragraph(description, style="List Bullet")

    summary = document.add_paragraph()
    run = summary.add_run(
        f"{len(changes)} change(s) detected since last generation."
    )
    run.italic = True

    logger.debug(
        "Rendered change tracking section with %d changes", len(changes)
    )
