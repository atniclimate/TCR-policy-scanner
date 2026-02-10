"""Hot Sheet renderer for Tribe-specific advocacy packets.

Builds the per-program Hot Sheet pages: status badge, program description,
award history, hazard relevance, economic impact, advocacy language,
Key Ask callout, structural asks, and delegation info.

Each Hot Sheet is a self-contained 1-2 page section ordered by relevance
score.  Empty sections are omitted (hidden), not shown with N/A.
"""

import logging
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.packets.context import TribePacketContext
from src.packets.docx_styles import (
    CI_STATUS_COLORS,
    CI_STATUS_LABELS,
    COLORS,
    StyleManager,
    add_status_badge,
    apply_zebra_stripe,
    format_header_row,
    set_cell_shading,
)
from src.packets.economic import (
    METHODOLOGY_CITATION,
    EconomicImpactCalculator,
    ProgramEconomicImpact,
    TribeEconomicSummary,
)

logger = logging.getLogger("tcr_scanner.packets.docx_hotsheet")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

FRAMING_BY_STATUS: dict[str, dict[str, Any]] = {
    "FLAGGED": {
        "tone": "defensive",
        "verbs": ["protect", "defend", "restore", "prevent loss of"],
        "prefix": "Critical:",
    },
    "AT_RISK": {
        "tone": "defensive",
        "verbs": ["protect", "defend", "restore", "prevent loss of"],
        "prefix": "Critical:",
    },
    "UNCERTAIN": {
        "tone": "protective",
        "verbs": ["safeguard", "ensure continuity of", "prepare", "maintain"],
        "prefix": "Watch:",
    },
    "STABLE_BUT_VULNERABLE": {
        "tone": "protective",
        "verbs": ["safeguard", "ensure continuity of", "prepare", "maintain"],
        "prefix": "Watch:",
    },
    "STABLE": {
        "tone": "growth",
        "verbs": ["expand", "build on", "strengthen", "invest in"],
        "prefix": "Opportunity:",
    },
    "SECURE": {
        "tone": "growth",
        "verbs": ["expand", "build on", "strengthen", "invest in"],
        "prefix": "Opportunity:",
    },
}
"""Framing language indexed by CI status value."""

RELEVANT_COMMITTEE_KEYWORDS: list[str] = [
    "Appropriations",
    "Indian Affairs",
    "Energy",
    "Natural Resources",
    "Environment",
    "Commerce",
    "Transportation",
    "Agriculture",
]
"""Substrings used to match delegation committees to program relevance."""

# Wildfire-related program IDs (for USFS wildfire data display)
_WILDFIRE_PROGRAM_IDS: set[str] = {
    "usda_wildfire",
    "fema_bric",
    "fema_tribal_mitigation",
    "bia_tcr",
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _normalize_cfda(cfda: Any) -> list[str]:
    """Normalize a CFDA value (str, list, or None) to a list of strings."""
    if isinstance(cfda, str):
        return [cfda]
    if isinstance(cfda, list):
        return [c for c in cfda if isinstance(c, str)]
    return []


def _format_dollars(amount: float) -> str:
    """Format a dollar amount with commas and no decimals."""
    return f"${amount:,.0f}"


def _add_paragraph_shading(paragraph, color_hex: str) -> None:
    """Apply background shading to a paragraph via XML manipulation."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


# ---------------------------------------------------------------------------
# HotSheetRenderer
# ---------------------------------------------------------------------------

class HotSheetRenderer:
    """Renders per-program Hot Sheet sections into a DOCX document.

    Each Hot Sheet contains up to 10 sub-sections (empty sections omitted):
    title+status, program description, award history, hazard relevance,
    economic impact, advocacy language, key ask callout, structural asks,
    delegation info, and methodology footnote.

    Usage::

        doc = Document()
        sm = StyleManager(doc)
        renderer = HotSheetRenderer(doc, sm)
        renderer.render_hotsheet(context, program, economic_impact, structural_asks)
    """

    def __init__(self, document: Document, style_manager: StyleManager) -> None:
        """Initialize the renderer with document and style manager.

        Args:
            document: python-docx Document to render into.
            style_manager: StyleManager with all custom styles registered.
        """
        self.document = document
        self.style_manager = style_manager
        self._calc = EconomicImpactCalculator()

    def render_all_hotsheets(
        self,
        context: TribePacketContext,
        programs: list[dict],
        economic_summary: TribeEconomicSummary,
        structural_asks: list[dict],
    ) -> int:
        """Render Hot Sheets for all programs into the document.

        Programs should already be filtered and sorted by relevance score.
        A page break is added between Hot Sheets (but not after the last one).

        Args:
            context: TribePacketContext with all Tribe data.
            programs: List of program dicts sorted by relevance.
            economic_summary: Aggregated economic impact for the Tribe.
            structural_asks: List of structural ask dicts.

        Returns:
            Number of Hot Sheets rendered.
        """
        count = 0
        # Build lookup from economic summary
        impact_by_program: dict[str, ProgramEconomicImpact] = {}
        for impact in economic_summary.by_program:
            impact_by_program[impact.program_id] = impact

        for i, program in enumerate(programs):
            if i > 0:
                self.document.add_page_break()

            program_id = program.get("id", "")
            economic_impact = impact_by_program.get(program_id)

            self.render_hotsheet(
                context=context,
                program=program,
                economic_impact=economic_impact,
                structural_asks=structural_asks,
            )
            count += 1
            logger.debug(
                "Rendered Hot Sheet %d/%d: %s",
                count,
                len(programs),
                program.get("name", program_id),
            )

        logger.info(
            "Rendered %d Hot Sheet(s) for %s", count, context.tribe_name
        )
        return count

    def render_hotsheet(
        self,
        context: TribePacketContext,
        program: dict,
        economic_impact: ProgramEconomicImpact | None,
        structural_asks: list[dict],
    ) -> None:
        """Render a single Hot Sheet section into the document.

        Sub-sections are rendered in fixed order; sections with no data are
        omitted entirely (never shown as blank or N/A).

        Args:
            context: TribePacketContext with all Tribe data.
            program: Program metadata dict.
            economic_impact: Economic impact for this program (or None).
            structural_asks: List of all structural ask dicts.
        """
        self._add_title_and_status(program)
        self._add_program_description(program)
        self._add_award_history(context, program)
        self._add_hazard_relevance(context, program)
        self._add_economic_impact(economic_impact, context, program)
        self._add_advocacy_language(program)
        self._add_key_ask(program, context, economic_impact)
        self._add_structural_asks(program, structural_asks, context)
        self._add_delegation_info(context, program)
        self._add_methodology_footnote()

    # ------------------------------------------------------------------
    # Private section methods
    # ------------------------------------------------------------------

    def _add_title_and_status(self, program: dict) -> None:
        """Add program title, metadata line, and CI status badge."""
        name = program.get("name", "Unknown Program")
        self.document.add_paragraph(name, style="HS Title")

        # Metadata line: agency, access type, funding type
        meta_parts: list[str] = []
        if program.get("agency"):
            meta_parts.append(program["agency"])
        if program.get("access_type"):
            meta_parts.append(program["access_type"].replace("_", " ").title())
        if program.get("funding_type"):
            meta_parts.append(program["funding_type"])
        if meta_parts:
            self.document.add_paragraph(
                " | ".join(meta_parts), style="HS Subtitle"
            )

        # Status badge
        ci_status = program.get("ci_status", "")
        if ci_status:
            badge_para = self.document.add_paragraph(style="HS Body")
            add_status_badge(badge_para, ci_status)

        # CI determination text
        determination = program.get("ci_determination", "")
        if determination:
            self.document.add_paragraph(determination, style="HS Body")

    def _add_program_description(self, program: dict) -> None:
        """Add program overview section."""
        description = program.get("description", "")
        federal_home = program.get("federal_home", "")
        if not description and not federal_home:
            return

        self.document.add_paragraph("Program Overview", style="HS Section")
        if description:
            self.document.add_paragraph(description, style="HS Body")
        if federal_home:
            self.document.add_paragraph(
                f"Federal Home: {federal_home}", style="HS Body"
            )

    def _add_award_history(
        self, context: TribePacketContext, program: dict
    ) -> None:
        """Add award history table or first-time applicant framing."""
        program_id = program.get("id", "")
        cfda_list = _normalize_cfda(program.get("cfda"))

        # Find matching awards
        matching_awards: list[dict] = []
        for award in context.awards:
            if award.get("program_id") == program_id:
                matching_awards.append(award)
            elif award.get("cfda") and award.get("cfda") in cfda_list:
                matching_awards.append(award)

        self.document.add_paragraph("Award History", style="HS Section")

        if matching_awards:
            # Create table with award data
            headers = ["Program", "CFDA", "Obligation", "Period"]
            table = self.document.add_table(
                rows=1 + len(matching_awards), cols=4
            )
            format_header_row(table, headers)

            total_obligation = 0.0
            for row_idx, award in enumerate(matching_awards, start=1):
                row = table.rows[row_idx]
                row.cells[0].text = award.get("program_id", program_id)
                row.cells[1].text = str(award.get("cfda", ""))
                obligation = award.get("obligation", 0.0)
                if isinstance(obligation, (int, float)):
                    total_obligation += obligation
                    row.cells[2].text = _format_dollars(obligation)
                else:
                    row.cells[2].text = str(obligation)
                start = award.get("start_date", "")
                end = award.get("end_date", "")
                if start or end:
                    row.cells[3].text = f"{start} - {end}"

            apply_zebra_stripe(table)

            # Total row
            total_para = self.document.add_paragraph(style="HS Body")
            total_run = total_para.add_run(
                f"Total: {_format_dollars(total_obligation)}"
            )
            total_run.bold = True
        else:
            # First-Time Applicant Advantage framing
            program_name = program.get("name", "this program")
            tribe_name = context.tribe_name

            bold_para = self.document.add_paragraph(style="HS Body")
            bold_run = bold_para.add_run("First-Time Applicant Advantage")
            bold_run.bold = True

            self.document.add_paragraph(
                f"{tribe_name} has not yet received funding through "
                f"{program_name}. This positions the Tribe as a first-time "
                f"applicant -- a category often given priority consideration "
                f"in competitive programs.",
                style="HS Body",
            )

    def _add_hazard_relevance(
        self, context: TribePacketContext, program: dict
    ) -> None:
        """Add local hazard relevance section."""
        # Extract hazard data -- handle both nested and flat formats
        hazard_profile = context.hazard_profile or {}
        fema_nri = hazard_profile.get("fema_nri") or hazard_profile.get(
            "sources", {}
        ).get("fema_nri", {})
        top_hazards = fema_nri.get("top_hazards", []) if fema_nri else []
        usfs_data = hazard_profile.get("usfs_wildfire") or hazard_profile.get(
            "sources", {}
        ).get("usfs_wildfire", {})

        if not top_hazards and not usfs_data:
            # No hazard data at all -- skip section
            return

        self.document.add_paragraph(
            "Local Hazard Relevance", style="HS Section"
        )

        if not top_hazards:
            # Only have USFS data or minimal hazard info
            self.document.add_paragraph(
                "Hazard relevance determined by ecoregion classification and "
                "available wildfire risk assessments.",
                style="HS Body",
            )
        else:
            # Show top 3 hazards
            for hazard in top_hazards[:3]:
                haz_type = hazard.get("type", "Unknown")
                risk_score = hazard.get("risk_score")
                risk_rating = hazard.get("risk_rating", "")

                line_parts = [haz_type]
                if risk_score is not None:
                    line_parts.append(f"Risk Score: {risk_score:.1f}")
                if risk_rating:
                    line_parts.append(f"({risk_rating})")

                self.document.add_paragraph(
                    " -- ".join(line_parts), style="HS Body"
                )

        # USFS wildfire data for wildfire-related programs
        program_id = program.get("id", "")
        if program_id in _WILDFIRE_PROGRAM_IDS and usfs_data:
            risk_to_homes = usfs_data.get("risk_to_homes")
            likelihood = usfs_data.get("likelihood")
            parts: list[str] = []
            if risk_to_homes is not None:
                parts.append(
                    f"USFS Wildfire Risk to Homes: {risk_to_homes:.2f}"
                )
            if likelihood is not None:
                parts.append(f"Fire Likelihood: {likelihood:.2f}")
            if parts:
                self.document.add_paragraph(
                    " | ".join(parts), style="HS Body"
                )

    def _add_economic_impact(
        self,
        economic_impact: ProgramEconomicImpact | None,
        context: TribePacketContext,
        program: dict,
    ) -> None:
        """Add district economic impact section."""
        if economic_impact is None:
            return

        self.document.add_paragraph(
            "District Economic Impact", style="HS Section"
        )

        program_name = program.get("name", "this program")
        narrative = self._calc.format_impact_narrative(
            economic_impact, context.tribe_name, program_name
        )

        if economic_impact.is_benchmark:
            self.document.add_paragraph(
                f"Based on program averages for {program_name}:",
                style="HS Body",
            )

        self.document.add_paragraph(narrative, style="HS Body")

        # Multi-district breakdown
        if len(context.districts) > 1:
            self.document.add_paragraph(
                "Per-District Impact Allocation:", style="HS Body"
            )
            for district_info in context.districts:
                district_key = district_info.get("district", "unknown")
                overlap = district_info.get("overlap_pct", 0.0)
                if overlap and overlap > 0:
                    frac = overlap / 100.0
                    low = economic_impact.estimated_impact_low * frac
                    high = economic_impact.estimated_impact_high * frac
                    self.document.add_paragraph(
                        f"  {district_key} ({overlap:.0f}% overlap): "
                        f"{_format_dollars(low)} - {_format_dollars(high)}",
                        style="HS Body",
                    )

        # BCR narrative for mitigation programs
        bcr_text = self._calc.format_bcr_narrative(economic_impact)
        if bcr_text:
            self.document.add_paragraph(bcr_text, style="HS Body")

    def _add_advocacy_language(self, program: dict) -> None:
        """Add advocacy position section with framing by CI status."""
        ci_status = program.get("ci_status", "STABLE")
        framing = FRAMING_BY_STATUS.get(
            ci_status,
            FRAMING_BY_STATUS["STABLE"],
        )
        prefix = framing["prefix"]

        tightened = program.get("tightened_language")
        description = program.get("description", "")
        advocacy_lever = program.get("advocacy_lever", "")

        if not tightened and not description and not advocacy_lever:
            return

        self.document.add_paragraph("Advocacy Position", style="HS Section")

        if tightened:
            self.document.add_paragraph(
                f"{prefix} {tightened}", style="HS Body"
            )
        else:
            # Fallback: combine description + advocacy_lever
            fallback_text = f"{prefix} {advocacy_lever}. {description}".strip()
            if fallback_text.endswith("."):
                pass  # Already ends with period
            else:
                fallback_text += "."
            self.document.add_paragraph(fallback_text, style="HS Body")

        # Optional proposed fix
        proposed_fix = program.get("proposed_fix")
        if proposed_fix:
            self.document.add_paragraph(
                f"Proposed Fix: {proposed_fix}", style="HS Body"
            )

        # Optional specialist focus
        specialist_focus = program.get("specialist_focus")
        if specialist_focus:
            para = self.document.add_paragraph(style="HS Body")
            run = para.add_run(f"Specialist Focus: {specialist_focus}")
            run.bold = True

    def _add_key_ask(
        self,
        program: dict,
        context: TribePacketContext,
        economic_impact: ProgramEconomicImpact | None,
    ) -> None:
        """Add Key Ask callout block (ASK / WHY / IMPACT)."""
        advocacy_lever = program.get("advocacy_lever", "")
        if not advocacy_lever:
            return

        self.document.add_paragraph("Key Ask", style="HS Section")

        # ASK line
        ask_para = self.document.add_paragraph(style="HS Callout")
        ask_para.add_run(f"ASK: {advocacy_lever}")
        _add_paragraph_shading(ask_para, COLORS.callout_bg)

        # WHY line -- connect to hazard context
        hazard_context = self._get_hazard_context_brief(context)
        why_text = (
            f"{context.tribe_name}'s {hazard_context} makes this program "
            f"critical for community resilience"
        )
        why_para = self.document.add_paragraph(style="HS Callout Detail")
        why_para.add_run(f"WHY: {why_text}")
        _add_paragraph_shading(why_para, COLORS.callout_bg)

        # IMPACT line
        if economic_impact is not None:
            program_name = program.get("name", "this program")
            impact_text = self._calc.format_impact_narrative(
                economic_impact, context.tribe_name, program_name
            )
        else:
            impact_text = (
                "Federal investment in this program supports Tribal "
                "community resilience and regional economic development"
            )
        impact_para = self.document.add_paragraph(style="HS Callout Detail")
        impact_para.add_run(f"IMPACT: {impact_text}")
        _add_paragraph_shading(impact_para, COLORS.callout_bg)

    def _add_structural_asks(
        self,
        program: dict,
        structural_asks: list[dict],
        context: TribePacketContext,
    ) -> None:
        """Add structural policy asks that include this program."""
        program_id = program.get("id", "")
        matching_asks = [
            ask
            for ask in structural_asks
            if program_id in ask.get("programs", [])
        ]

        if not matching_asks:
            return

        self.document.add_paragraph(
            "Structural Policy Asks", style="HS Section"
        )

        for ask in matching_asks:
            ask_name = ask.get("name", "")
            ask_desc = ask.get("description", "")
            ask_target = ask.get("target", "")
            ask_urgency = ask.get("urgency", "")

            # Ask heading with urgency
            heading_parts = [ask_name]
            if ask_urgency:
                heading_parts.append(f"[{ask_urgency}]")
            heading_para = self.document.add_paragraph(style="HS Body")
            heading_run = heading_para.add_run(" ".join(heading_parts))
            heading_run.bold = True

            # Description and target
            if ask_desc:
                self.document.add_paragraph(ask_desc, style="HS Body")
            if ask_target:
                self.document.add_paragraph(
                    f"Target: {ask_target}", style="HS Body"
                )

            # Brief evidence line from hazard/award data
            evidence = self._get_evidence_for_ask(ask, context, program)
            if evidence:
                para = self.document.add_paragraph(style="HS Small")
                para.add_run(f"Evidence: {evidence}")

    def _add_delegation_info(
        self, context: TribePacketContext, program: dict
    ) -> None:
        """Add congressional delegation section with committee-match flags."""
        senators = context.senators or []
        representatives = context.representatives or []

        if not senators and not representatives:
            self.document.add_paragraph("Your Delegation", style="HS Section")
            self.document.add_paragraph(
                "Congressional delegation data not available.",
                style="HS Body",
            )
            return

        self.document.add_paragraph("Your Delegation", style="HS Section")

        all_members = []
        for s in senators:
            all_members.append(s)
        for r in representatives:
            all_members.append(r)

        for member in all_members:
            name = member.get("formatted_name", "Unknown Member")
            self.document.add_paragraph(name, style="HS Body")

            # Committee-match flag
            committees = member.get("committees", [])
            for committee in committees:
                committee_name = committee.get("committee_name", "")
                for keyword in RELEVANT_COMMITTEE_KEYWORDS:
                    if keyword.lower() in committee_name.lower():
                        callout_para = self.document.add_paragraph(
                            style="HS Body"
                        )
                        callout_run = callout_para.add_run(
                            f"{name} sits on {committee_name} "
                            f"-- direct influence on this program"
                        )
                        callout_run.bold = True
                        break  # One match per committee is enough

    def _add_methodology_footnote(self) -> None:
        """Add methodology citation as footnote with horizontal rule."""
        # Thin horizontal rule
        rule_para = self.document.add_paragraph()
        rule_run = rule_para.add_run(
            "\u2500" * 40
        )  # Box-drawing horizontal line
        rule_run.font.size = Pt(6)
        rule_run.font.color.rgb = COLORS.muted

        # Citation text
        self.document.add_paragraph(METHODOLOGY_CITATION, style="HS Small")

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _get_hazard_context_brief(self, context: TribePacketContext) -> str:
        """Build a brief hazard context string for Key Ask WHY line."""
        hazard_profile = context.hazard_profile or {}
        fema_nri = hazard_profile.get("fema_nri") or hazard_profile.get(
            "sources", {}
        ).get("fema_nri", {})
        top_hazards = fema_nri.get("top_hazards", []) if fema_nri else []

        if not top_hazards:
            return "climate risk exposure"

        hazard_names = [
            h.get("type", "unknown") for h in top_hazards[:3]
        ]
        if len(hazard_names) == 1:
            return f"{hazard_names[0]} risk"
        return f"{', '.join(hazard_names[:-1])} and {hazard_names[-1]} risk"

    def _get_evidence_for_ask(
        self,
        ask: dict,
        context: TribePacketContext,
        program: dict,
    ) -> str:
        """Build a brief evidence line linking ask to Tribe data."""
        parts: list[str] = []

        # Check for award evidence
        program_id = program.get("id", "")
        cfda_list = _normalize_cfda(program.get("cfda"))
        has_awards = any(
            a.get("program_id") == program_id or a.get("cfda") in cfda_list
            for a in context.awards
        )
        if has_awards:
            parts.append(f"Active awards under {program.get('name', program_id)}")

        # Check for hazard relevance
        hazard_profile = context.hazard_profile or {}
        fema_nri = hazard_profile.get("fema_nri") or hazard_profile.get(
            "sources", {}
        ).get("fema_nri", {})
        top_hazards = fema_nri.get("top_hazards", []) if fema_nri else []
        if top_hazards:
            top_type = top_hazards[0].get("type", "")
            if top_type:
                parts.append(f"{top_type} risk profile")

        return "; ".join(parts) if parts else ""
