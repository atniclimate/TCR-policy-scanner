"""DOCX styling foundation for Tribal advocacy packets.

Provides StyleManager (programmatic custom styles), helper functions
(cell shading, zebra striping, status badges, header formatting),
and color constants for CI status values.

All styles use the Arial font family for professional, accessible output.
"""

import logging

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Color constants
# ---------------------------------------------------------------------------

CI_STATUS_COLORS: dict[str, RGBColor] = {
    "FLAGGED": RGBColor(0xDC, 0x26, 0x26),
    "AT_RISK": RGBColor(0xDC, 0x26, 0x26),
    "UNCERTAIN": RGBColor(0xB4, 0x53, 0x09),
    "STABLE_BUT_VULNERABLE": RGBColor(0xB4, 0x53, 0x09),
    "STABLE": RGBColor(0x16, 0xA3, 0x4A),
    "SECURE": RGBColor(0x16, 0xA3, 0x4A),
}

CI_STATUS_LABELS: dict[str, str] = {
    "FLAGGED": "Flagged -- Requires Specialist Attention",
    "AT_RISK": "At Risk",
    "UNCERTAIN": "Uncertain",
    "STABLE_BUT_VULNERABLE": "Stable but Vulnerable",
    "STABLE": "Stable",
    "SECURE": "Secure",
}


class _ColorNamespace:
    """Palette constants for document styling."""

    primary_dark = RGBColor(0x1A, 0x23, 0x7E)
    heading = RGBColor(0x37, 0x41, 0x51)
    body_text = RGBColor(0x1F, 0x29, 0x37)
    muted = RGBColor(0x4B, 0x55, 0x63)
    table_header_bg = "1A237E"
    table_stripe_even = "F2F2F2"
    table_stripe_odd = "FFFFFF"
    callout_bg = "EEF2FF"


COLORS = _ColorNamespace()

# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------


def set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading color on a table cell.

    Creates a FRESH OxmlElement for each call to avoid element reuse bugs
    where only the last cell gets styled.

    Args:
        cell: A python-docx table cell object.
        color_hex: Six-character hex color string (e.g. ``"F2F2F2"``).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def apply_zebra_stripe(table, header_rows: int = 1) -> None:
    """Apply alternating row colors to a table for readability.

    Skips the specified number of header rows, then alternates between
    ``COLORS.table_stripe_even`` and ``COLORS.table_stripe_odd``.

    Args:
        table: A python-docx Table object.
        header_rows: Number of header rows to skip (default 1).
    """
    for i, row in enumerate(table.rows):
        if i < header_rows:
            continue
        data_index = i - header_rows
        color = COLORS.table_stripe_even if data_index % 2 == 0 else COLORS.table_stripe_odd
        for cell in row.cells:
            set_cell_shading(cell, color)


def add_status_badge(paragraph, ci_status: str) -> None:
    """Add a colored Unicode circle and styled status label to a paragraph.

    Renders a colored circle (U+25CF) followed by the human-readable
    status label from ``CI_STATUS_LABELS``.

    Args:
        paragraph: A python-docx Paragraph object.
        ci_status: One of the six CI status strings (e.g. ``"FLAGGED"``).
    """
    color = CI_STATUS_COLORS.get(ci_status, RGBColor(0x6B, 0x72, 0x80))
    label = CI_STATUS_LABELS.get(ci_status, ci_status)

    # First run: colored circle
    circle_run = paragraph.add_run("\u25CF ")
    circle_run.font.color.rgb = color
    circle_run.font.size = Pt(14)
    circle_run.font.name = "Arial"

    # Second run: bold colored status label
    label_run = paragraph.add_run(label)
    label_run.font.color.rgb = color
    label_run.font.size = Pt(11)
    label_run.font.bold = True
    label_run.font.name = "Arial"


def format_header_row(table, headers: list[str]) -> None:
    """Style the first row of a table as a dark header with white text.

    Applies dark background (``COLORS.table_header_bg``) and white bold
    Arial 9pt text to each cell in the first row. Also sets the
    ``w:tblHeader`` property for Section 508 screen reader accessibility.

    Args:
        table: A python-docx Table object.
        headers: List of header text strings (must match column count).
    """
    header_row = table.rows[0]
    for i, cell in enumerate(header_row.cells):
        set_cell_shading(cell, COLORS.table_header_bg)
        if i < len(headers):
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(headers[i])
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Arial"

    # Section 508: mark header row for screen readers
    trPr = header_row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


# ---------------------------------------------------------------------------
# StyleManager
# ---------------------------------------------------------------------------


class StyleManager:
    """Creates and manages all custom paragraph styles for advocacy packets.

    Instantiate with a python-docx Document to register all custom styles.
    Handles duplicate style names gracefully (skips if already present).

    Usage::

        doc = Document()
        sm = StyleManager(doc)
        # All custom styles now available on doc.styles
    """

    def __init__(self, document: Document) -> None:
        self.doc = document
        self._create_styles()

    def _create_styles(self) -> None:
        """Create all custom paragraph styles, skipping any that already exist."""
        existing_names = {s.name for s in self.doc.styles}

        # Custom paragraph style definitions: (name, size_pt, bold, color, extras)
        custom_styles = [
            {
                "name": "HS Title",
                "size": Pt(18),
                "bold": True,
                "color": COLORS.primary_dark,
                "space_after": Pt(6),
            },
            {
                "name": "HS Subtitle",
                "size": Pt(11),
                "bold": False,
                "color": COLORS.heading,
                "space_after": Pt(4),
            },
            {
                "name": "HS Section",
                "size": Pt(12),
                "bold": True,
                "color": COLORS.heading,
                "space_before": Pt(12),
                "space_after": Pt(4),
            },
            {
                "name": "HS Body",
                "size": Pt(10),
                "bold": False,
                "color": COLORS.body_text,
                "space_after": Pt(4),
            },
            {
                "name": "HS Small",
                "size": Pt(8),
                "bold": False,
                "color": COLORS.muted,
                "space_after": Pt(2),
            },
            {
                "name": "HS Callout",
                "size": Pt(10),
                "bold": True,
                "color": COLORS.body_text,
                "left_indent": Inches(0.3),
                "space_before": Pt(8),
            },
            {
                "name": "HS Callout Detail",
                "size": Pt(10),
                "bold": False,
                "color": COLORS.body_text,
                "left_indent": Inches(0.3),
                "space_after": Pt(2),
            },
            {
                "name": "HS Table Header",
                "size": Pt(9),
                "bold": True,
                "color": RGBColor(0xFF, 0xFF, 0xFF),
                "space_after": Pt(0),
            },
            {
                "name": "HS Table Body",
                "size": Pt(9),
                "bold": False,
                "color": COLORS.body_text,
                "space_after": Pt(0),
            },
        ]

        for spec in custom_styles:
            name = spec["name"]
            if name in existing_names:
                logger.debug("Style '%s' already exists, skipping", name)
                continue

            style = self.doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Arial"
            style.font.size = spec["size"]
            style.font.bold = spec.get("bold", False)
            style.font.color.rgb = spec["color"]

            fmt = style.paragraph_format
            if "space_after" in spec:
                fmt.space_after = spec["space_after"]
            if "space_before" in spec:
                fmt.space_before = spec["space_before"]
            if "left_indent" in spec:
                fmt.left_indent = spec["left_indent"]

            existing_names.add(name)

        # Override built-in heading styles to use Arial
        heading_overrides = [
            {"name": "Heading 1", "size": Pt(16), "bold": True, "color": COLORS.primary_dark},
            {"name": "Heading 2", "size": Pt(14), "bold": True, "color": COLORS.heading},
            {"name": "Heading 3", "size": Pt(12), "bold": True, "color": COLORS.heading},
        ]

        for spec in heading_overrides:
            try:
                style = self.doc.styles[spec["name"]]
                style.font.name = "Arial"
                style.font.size = spec["size"]
                style.font.bold = spec["bold"]
                style.font.color.rgb = spec["color"]
            except KeyError:
                logger.warning("Built-in style '%s' not found", spec["name"])
