"""Hot Sheet template builder for Tribal advocacy packets.

Generates a reusable .docx template pre-baked with all HS-* styles,
page layout, header/footer, and design primitives (2-column grid,
info boxes, horizontal rules). Rendering code opens the template
and fills in content -- never redefines styles at runtime.

Usage::

    from src.packets.docx_template import build_template
    path = build_template()  # -> templates/hot_sheet_template.docx

Or use the class directly::

    builder = TemplateBuilder()
    doc = builder.build()
    builder.save("output/my_template.docx")
"""

import logging
import os
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

from src.config import FISCAL_YEAR_SHORT
from src.packets.docx_styles import COLORS, StyleManager, set_cell_shading
from src.paths import PROJECT_ROOT

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN_TOP = Inches(1)
MARGIN_BOTTOM = Inches(1)
MARGIN_LEFT = Inches(1)
MARGIN_RIGHT = Inches(1)
GUTTER = Inches(0.25)
LEFT_RATIO = 0.62

FONT_FAMILY = "Arial"

BOX_SCHEMES: dict[str, dict[str, str]] = {
    "neutral": {"bg": "F3F4F6", "border": "D1D5DB"},
    "alert": {"bg": "FEF2F2", "border": "F87171"},
    "success": {"bg": "F0FDF4", "border": "4ADE80"},
    "callout": {"bg": "EEF2FF", "border": "818CF8"},
}
"""Color schemes for information boxes keyed by semantic type."""

# Extended styles that supplement the existing StyleManager definitions
_EXTENDED_STYLES: list[dict] = [
    {
        "name": "HS Column Primary",
        "size": Pt(10),
        "bold": False,
        "color": COLORS.body_text,
        "space_after": Pt(4),
    },
    {
        "name": "HS Column Secondary",
        "size": Pt(9),
        "bold": False,
        "color": COLORS.body_text,
        "space_after": Pt(3),
    },
    {
        "name": "HS Box Title",
        "size": Pt(10),
        "bold": True,
        "color": COLORS.heading,
        "space_after": Pt(2),
    },
    {
        "name": "HS Box Body",
        "size": Pt(9),
        "bold": False,
        "color": COLORS.body_text,
        "space_after": Pt(2),
    },
    {
        "name": "HS Ask Label",
        "size": Pt(10),
        "bold": True,
        "color": RGBColor(0x1A, 0x23, 0x7E),
        "space_after": Pt(1),
    },
    {
        "name": "HS Ask Detail",
        "size": Pt(9),
        "bold": False,
        "color": COLORS.body_text,
        "space_after": Pt(2),
    },
    {
        "name": "HS Footer",
        "size": Pt(7),
        "bold": False,
        "color": COLORS.muted,
        "space_after": Pt(0),
    },
    {
        "name": "HS Caption",
        "size": Pt(8),
        "bold": False,
        "color": COLORS.muted,
        "space_after": Pt(2),
    },
    {
        "name": "HS Page Number",
        "size": Pt(8),
        "bold": False,
        "color": COLORS.muted,
    },
]


# ---------------------------------------------------------------------------
# Low-level OxmlElement helpers
# ---------------------------------------------------------------------------


def _remove_table_borders(table) -> None:
    """Remove all borders from a table to create an invisible layout grid.

    Sets all six border edges (top, left, bottom, right, insideH, insideV)
    to ``val="none"`` so the table renders as a pure layout container.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)


def _set_cell_margins(cell, top: int = 100, bottom: int = 100,
                       start: int = 150, end: int = 150) -> None:
    """Set internal padding on a table cell.

    Values are in twentieths of a point (dxa).  Default values give a
    comfortable 5pt top/bottom and ~7.5pt left/right padding.

    Args:
        cell: A python-docx table cell object.
        top: Top margin in dxa.
        bottom: Bottom margin in dxa.
        start: Left margin in dxa.
        end: Right margin in dxa.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("start", start), ("end", end)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_border(cell, color_hex: str, size: int = 4) -> None:
    """Apply a thin solid border on all four sides of a table cell.

    Creates FRESH OxmlElement objects per call to avoid element-move bugs.

    Args:
        cell: A python-docx table cell object.
        color_hex: Six-character hex color string (e.g. ``"D1D5DB"``).
        size: Border width in eighths of a point (4 = 0.5pt).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# Public layout helpers
# ---------------------------------------------------------------------------


def create_two_column_grid(document: Document,
                           left_ratio: float = LEFT_RATIO):
    """Create a borderless 2-column asymmetric table (default 62/38 split).

    Uses invisible borders for a clean layout grid.  Column widths are
    computed from the first section's page dimensions minus margins.

    Args:
        document: The python-docx Document to add the grid to.
        left_ratio: Fraction of usable width for the left column (0.0-1.0).

    Returns:
        The created Table object with two columns.
    """
    section = document.sections[-1]
    usable = section.page_width - section.left_margin - section.right_margin
    gutter_emu = Emu(GUTTER)
    left_width = int((usable - gutter_emu) * left_ratio)
    right_width = usable - left_width - gutter_emu

    table = document.add_table(rows=1, cols=2)
    table.columns[0].width = left_width
    table.columns[1].width = right_width
    table.autofit = False

    _remove_table_borders(table)

    # Set gutter via right padding on left cell and left padding on right cell
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    _set_cell_margins(left_cell, top=0, bottom=0, start=0,
                       end=int(gutter_emu / 2 / 635))  # Emu to dxa
    _set_cell_margins(right_cell, top=0, bottom=0,
                       start=int(gutter_emu / 2 / 635), end=0)

    return table


def create_info_box(document: Document,
                    box_type: str = "neutral") -> tuple:
    """Create a styled single-cell table for information boxes.

    Returns a ``(table, cell)`` tuple so the caller can add content
    paragraphs to the cell.

    Args:
        document: The python-docx Document.
        box_type: One of ``"neutral"``, ``"alert"``, ``"success"``,
            ``"callout"``.  Determines background and border colors.

    Returns:
        Tuple of (Table, Cell) for the information box.
    """
    scheme = BOX_SCHEMES.get(box_type, BOX_SCHEMES["neutral"])

    table = document.add_table(rows=1, cols=1)
    table.autofit = False

    cell = table.rows[0].cells[0]
    set_cell_shading(cell, scheme["bg"])
    _set_cell_border(cell, scheme["border"], size=4)
    _set_cell_margins(cell, top=100, bottom=100, start=150, end=150)

    # Remove outer table borders so only cell borders show
    _remove_table_borders(table)

    return table, cell


def add_horizontal_rule(document: Document, color: str = "D1D5DB",
                        weight: int = 4, space: int = 6):
    """Add a horizontal rule via paragraph bottom border.

    Uses OOXML paragraph border (``w:pBdr``) instead of Unicode
    box-drawing characters for consistent rendering across Word and
    LibreOffice.

    Args:
        document: The python-docx Document.
        color: Six-character hex color string.
        weight: Border width in eighths of a point.
        space: Space between text and border in points.

    Returns:
        The Paragraph containing the rule.
    """
    para = document.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(weight))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Set font size to 2pt so the rule line is tight
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    run.font.size = Pt(2)

    return para


def mark_header_row(table) -> None:
    """Mark the first row of a table as a header for Section 508 accessibility.

    Sets the ``w:tblHeader`` property so screen readers recognize the row
    as a repeating table header.

    Args:
        table: A python-docx Table object.
    """
    for row in table.rows[:1]:
        trPr = row._tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)


def add_page_number_field(paragraph) -> None:
    """Insert a PAGE field code for automatic page numbering.

    Uses the OOXML field-code triplet (fldChar begin, instrText, fldChar end).

    Args:
        paragraph: A python-docx Paragraph object.
    """
    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r.append(fld_sep)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_end)


# ---------------------------------------------------------------------------
# TemplateBuilder
# ---------------------------------------------------------------------------


class TemplateBuilder:
    """Builds a reusable .docx template with pre-registered styles and layout.

    The template includes all HS-* paragraph styles, page dimensions,
    header/footer, and the font definitions needed for Hot Sheet rendering.
    Opening the saved template at runtime avoids re-creating styles on
    every document and saves ~0.2s per document in batch generation.

    Usage::

        builder = TemplateBuilder()
        doc = builder.build()
        path = builder.save("templates/hot_sheet.docx")
    """

    def __init__(self) -> None:
        self._document: Document | None = None

    def build(self) -> Document:
        """Create a fresh Document with all styles, layout, and metadata.

        Returns:
            A python-docx Document ready to be saved as a template.
        """
        document = Document()

        # 1. Register base styles via existing StyleManager
        StyleManager(document)

        # 2. Register extended styles for template-based workflow
        self._register_extended_styles(document)

        # 3. Configure page layout
        self._configure_page_layout(document)

        # 4. Set up header and footer
        self._setup_header_footer(document)

        # 5. Set default paragraph format
        self._set_default_paragraph_format(document)

        self._document = document
        logger.info("Built Hot Sheet template with %d custom styles",
                     len(_EXTENDED_STYLES) + 7)
        return document

    def save(self, output_path: Path | str) -> Path:
        """Save the template document to disk using atomic write.

        Args:
            output_path: Destination file path.

        Returns:
            Path to the saved .docx file.

        Raises:
            RuntimeError: If ``build()`` has not been called.
        """
        if self._document is None:
            raise RuntimeError("Call build() before save()")

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Atomic write: tmp file -> os.replace
        tmp_fd = tempfile.NamedTemporaryFile(
            dir=str(output_path.parent), suffix=".docx", delete=False
        )
        tmp_name = tmp_fd.name
        tmp_fd.close()
        try:
            self._document.save(tmp_name)
            os.replace(tmp_name, str(output_path))
        except Exception:
            try:
                os.unlink(tmp_name)
            except OSError:
                pass
            raise

        logger.info("Saved Hot Sheet template: %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _register_extended_styles(document: Document) -> None:
        """Register extended HS-* styles beyond the base StyleManager set."""
        from docx.enum.style import WD_STYLE_TYPE

        existing_names = {s.name for s in document.styles}

        for spec in _EXTENDED_STYLES:
            name = spec["name"]
            if name in existing_names:
                logger.debug("Extended style '%s' already exists, skipping",
                             name)
                continue

            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = FONT_FAMILY
            style.font.size = spec["size"]
            style.font.bold = spec.get("bold", False)
            style.font.color.rgb = spec["color"]

            fmt = style.paragraph_format
            if "space_after" in spec:
                fmt.space_after = spec["space_after"]
            if "space_before" in spec:
                fmt.space_before = spec["space_before"]
            if "left_indent" in spec:
                fmt.left_indent = spec["left_indent"]

            # Body text defaults based on style purpose
            centered_styles = {"HS Footer", "HS Page Number"}
            if name in centered_styles:
                fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Widow control and line spacing for multi-line body styles only
            single_line_styles = {"HS Footer", "HS Page Number", "HS Box Title"}
            if name not in single_line_styles:
                fmt.widow_control = True
                fmt.line_spacing = 1.15

            existing_names.add(name)

    @staticmethod
    def _configure_page_layout(document: Document) -> None:
        """Set page dimensions and margins on all sections."""
        for section in document.sections:
            section.page_width = PAGE_WIDTH
            section.page_height = PAGE_HEIGHT
            section.top_margin = MARGIN_TOP
            section.bottom_margin = MARGIN_BOTTOM
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT

    @staticmethod
    def _setup_header_footer(document: Document) -> None:
        """Configure header and footer with descriptive metadata and page numbers.

        Air gap compliant: no organizational names or tool attribution.
        """
        for section in document.sections:
            # --- Header ---
            header = section.header
            header.is_linked_to_previous = False
            header_para = (header.paragraphs[0] if header.paragraphs
                           else header.add_paragraph())
            header_run = header_para.add_run(
                f"{FISCAL_YEAR_SHORT} Climate Resilience Program Priorities"
            )
            header_run.font.size = Pt(8)
            header_run.font.color.rgb = COLORS.muted
            header_run.font.name = FONT_FAMILY

            # --- Footer ---
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = (footer.paragraphs[0] if footer.paragraphs
                           else footer.add_paragraph())
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # "For Congressional Office Use | Page X"
            conf_run = footer_para.add_run(
                "For Congressional Office Use | Page "
            )
            conf_run.font.size = Pt(7)
            conf_run.font.color.rgb = COLORS.muted
            conf_run.font.name = FONT_FAMILY

            add_page_number_field(footer_para)

            # Style the page number field
            for run in footer_para.runs:
                run.font.size = Pt(7)
                run.font.color.rgb = COLORS.muted
                run.font.name = FONT_FAMILY

    @staticmethod
    def _set_default_paragraph_format(document: Document) -> None:
        """Set Normal style defaults: Arial 10pt, rag-right, 1.15 spacing."""
        normal = document.styles["Normal"]
        normal.font.name = FONT_FAMILY
        normal.font.size = Pt(10)
        normal.font.color.rgb = COLORS.body_text
        fmt = normal.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.line_spacing = 1.15
        fmt.widow_control = True
        fmt.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Convenience function
# ---------------------------------------------------------------------------


def build_template(output_path: Path | str | None = None) -> Path:
    """Build and save the canonical Hot Sheet template.

    If ``output_path`` is None, saves to
    ``PROJECT_ROOT / "templates" / "hot_sheet_template.docx"``.

    Args:
        output_path: Optional destination path.

    Returns:
        Path to the saved template file.
    """
    if output_path is None:
        output_path = PROJECT_ROOT / "templates" / "hot_sheet_template.docx"

    builder = TemplateBuilder()
    builder.build()
    return builder.save(output_path)
