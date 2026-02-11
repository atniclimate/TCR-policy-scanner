"""Strategic overview document generator (Document 2).

Generates a single shared DOCX -- "{FY} Federal Funding Overview & Strategy" --
from existing v1.0 pipeline data: program_inventory.json, policy_tracking.json,
graph_schema.json, and optionally LATEST-MONITOR-DATA.json.

This document is generated once (not per-Tribe) and provides the cross-cutting
policy landscape used by all Tribal advocacy offices.
"""

import json
import logging
import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from src.config import FISCAL_YEAR_SHORT
from src.paths import (
    GRAPH_SCHEMA_PATH,
    LATEST_MONITOR_DATA_PATH,
    POLICY_TRACKING_PATH,
    PROJECT_ROOT,
)
from src.packets.docx_hotsheet import FRAMING_BY_STATUS
from src.packets.docx_styles import (
    CI_STATUS_LABELS,
    COLORS,
    StyleManager,
    add_status_badge,
    apply_zebra_stripe,
    format_header_row,
)
from src.packets.ecoregion import EcoregionMapper

logger = logging.getLogger(__name__)

_MAX_DATA_FILE_BYTES: int = 10 * 1024 * 1024  # 10 MB

# Advocacy goal groupings derived from CI status
_GOAL_BY_STATUS: dict[str, str] = {
    "FLAGGED": "DEFEND",
    "AT_RISK": "DEFEND",
    "UNCERTAIN": "PROTECT",
    "STABLE_BUT_VULNERABLE": "PROTECT",
    "STABLE": "EXPAND",
    "SECURE": "EXPAND",
}

# FEMA program IDs for the dedicated FEMA analysis section
_FEMA_PROGRAM_IDS: set[str] = {
    "fema_bric", "fema_hmgp", "fema_fma", "fema_pdm",
    "fema_tribal_mitigation",
}

# Subheadings for each advocacy goal group
_GOAL_HEADINGS: dict[str, str] = {
    "DEFEND": "DEFEND: Critical Programs Requiring Immediate Action",
    "PROTECT": "PROTECT: Programs Requiring Vigilance",
    "EXPAND": "EXPAND: Programs with Growth Opportunity",
}


class StrategicOverviewGenerator:
    """Generates the shared strategic overview DOCX (Document 2).

    Uses existing v1.0 data: program_inventory.json, policy_tracking.json,
    graph_schema.json, and optionally LATEST-MONITOR-DATA.json.

    Usage::

        generator = StrategicOverviewGenerator(config, programs)
        path = generator.generate(output_dir)
    """

    def __init__(self, config: dict, programs: dict[str, dict]) -> None:
        """Initialize the generator with config and program inventory.

        Args:
            config: Application configuration dict (with ``packets`` section).
            programs: Dict of program dicts keyed by program ID.
        """
        self.config = config
        self.programs = programs
        self.ecoregion = EcoregionMapper(config)
        self.policy_tracking = self._load_json(str(POLICY_TRACKING_PATH))
        self.graph_schema = self._load_json(str(GRAPH_SCHEMA_PATH))
        self.monitor_data = self._load_json(str(LATEST_MONITOR_DATA_PATH))

    def _load_json(self, path_str: str) -> dict:
        """Load a JSON file with size-limit guard and graceful fallback.

        Args:
            path_str: Relative or absolute path to the JSON file.

        Returns:
            Parsed dict, or empty dict on missing/corrupt/oversized files.
        """
        path = Path(path_str)
        if not path.is_absolute():
            path = PROJECT_ROOT / path

        if not path.exists():
            logger.debug("Data file not found: %s", path)
            return {}

        try:
            file_size = path.stat().st_size
        except OSError:
            return {}

        if file_size > _MAX_DATA_FILE_BYTES:
            logger.warning(
                "Data file %s exceeds size limit (%d bytes > %d), skipping",
                path, file_size, _MAX_DATA_FILE_BYTES,
            )
            return {}

        try:
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            return data if isinstance(data, dict) else {}
        except (json.JSONDecodeError, OSError) as exc:
            logger.warning("Failed to load %s: %s", path, exc)
            return {}

    def generate(self, output_dir: Path) -> Path:
        """Generate the strategic overview DOCX and save to output_dir.

        Creates the document with 7 sections, applies styling, and saves
        using an atomic write pattern (tmp file + os.replace).

        Args:
            output_dir: Directory to write STRATEGIC-OVERVIEW.docx.

        Returns:
            Path to the generated STRATEGIC-OVERVIEW.docx file.
        """
        document = Document()
        style_manager = StyleManager(document)

        # Configure page margins (1 inch all sides)
        for section in document.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        self._setup_header_footer(document)

        # Render 7 sections
        self._render_cover(document, style_manager)
        self._render_appropriations(document, style_manager)
        self._render_ecoregion_priorities(document, style_manager)
        self._render_science_threats(document, style_manager)
        self._render_fema_analysis(document, style_manager)
        self._render_cross_cutting_framework(document, style_manager)
        self._render_messaging_guidance(document, style_manager)

        # Atomic save
        output_path = output_dir / "STRATEGIC-OVERVIEW.docx"
        output_dir.mkdir(parents=True, exist_ok=True)

        tmp_fd = tempfile.NamedTemporaryFile(
            dir=str(output_dir), suffix=".docx", delete=False
        )
        tmp_name = tmp_fd.name
        tmp_fd.close()
        try:
            document.save(tmp_name)
            os.replace(tmp_name, str(output_path))
        except Exception:
            try:
                os.unlink(tmp_name)
            except OSError:
                pass
            raise

        logger.info("Saved strategic overview: %s", output_path)
        return output_path

    def _setup_header_footer(self, document: Document) -> None:
        """Configure header and footer for the strategic overview.

        Header: f"{FISCAL_YEAR_SHORT} Federal Funding Overview" (Arial 8pt, muted)
        Footer: "CONFIDENTIAL | Generated by TCR Policy Scanner" (Arial 7pt, muted, centered)
        """
        for section in document.sections:
            # Header
            header = section.header
            header.is_linked_to_previous = False
            header_para = (
                header.paragraphs[0]
                if header.paragraphs
                else header.add_paragraph()
            )
            header_run = header_para.add_run(
                f"{FISCAL_YEAR_SHORT} Federal Funding Overview"
            )
            header_run.font.size = Pt(8)
            header_run.font.color.rgb = COLORS.muted
            header_run.font.name = "Arial"

            # Footer
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = (
                footer.paragraphs[0]
                if footer.paragraphs
                else footer.add_paragraph()
            )
            footer_run = footer_para.add_run(
                "CONFIDENTIAL | Generated by TCR Policy Scanner"
            )
            footer_run.font.size = Pt(7)
            footer_run.font.color.rgb = COLORS.muted
            footer_run.font.name = "Arial"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------------
    # Section renderers
    # ------------------------------------------------------------------

    def _render_cover(self, document: Document, style_manager: StyleManager) -> None:
        """Render the cover page (Section 1)."""
        # Spacing
        document.add_paragraph("", style="HS Body")

        # Main title
        document.add_heading(f"{FISCAL_YEAR_SHORT} Federal Funding Overview & Strategy", level=1)

        # Subtitle
        document.add_paragraph(
            "Tribal Climate Resilience Programs", style="HS Title"
        )

        # Description
        document.add_paragraph(
            "A comprehensive analysis of 16 federal programs across 7 ecoregions",
            style="HS Body",
        )

        # Date
        now = datetime.now(timezone.utc)
        document.add_paragraph(
            f"Generated: {now.strftime('%B %d, %Y')}", style="HS Body"
        )

        # Confidentiality
        conf_para = document.add_paragraph(
            "CONFIDENTIAL -- For Advocacy Office Use Only", style="HS Small"
        )
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_page_break()

    def _render_appropriations(
        self, document: Document, style_manager: StyleManager
    ) -> None:
        """Render the Appropriations Landscape section (Section 2)."""
        document.add_heading("Appropriations Landscape", level=2)

        document.add_paragraph(
            "Current CI status and funding positions for all 16 tracked "
            "federal climate resilience programs.",
            style="HS Body",
        )

        # Build policy tracking lookup
        pt_programs = self.policy_tracking.get("positions", [])
        pt_by_id: dict[str, dict] = {}
        for pos in pt_programs:
            pid = pos.get("program_id", "")
            if pid:
                pt_by_id[pid] = pos

        # Sort programs by name
        sorted_programs = sorted(
            self.programs.values(), key=lambda p: p.get("name", "")
        )

        # Create table: Program | Agency | CI Status | Funding Type | Authorization
        headers = ["Program", "Agency", "CI Status", "Funding Type", "Authorization"]
        table = document.add_table(
            rows=1 + len(sorted_programs), cols=len(headers)
        )

        format_header_row(table, headers)

        flagged_or_risk = 0
        stable_or_secure = 0

        for row_idx, program in enumerate(sorted_programs, start=1):
            row = table.rows[row_idx]
            prog_id = program.get("id", "")
            ci_status = program.get("ci_status", "")

            row.cells[0].text = program.get("name", "")
            row.cells[1].text = program.get("agency", "")
            row.cells[2].text = CI_STATUS_LABELS.get(ci_status, ci_status)
            row.cells[3].text = program.get("funding_type", "")

            # Authorization from policy_tracking
            auth_status = ""
            pt_entry = pt_by_id.get(prog_id, {})
            ext = pt_entry.get("extensible_fields", {})
            auth_status = ext.get("authorization_status", "")
            if not auth_status:
                # Fallback: use reasoning as brief summary
                auth_status = pt_entry.get("reasoning", "")[:60]
                if len(pt_entry.get("reasoning", "")) > 60:
                    auth_status += "..."
            row.cells[4].text = auth_status

            # Count statuses
            if ci_status in ("FLAGGED", "AT_RISK"):
                flagged_or_risk += 1
            elif ci_status in ("STABLE", "SECURE"):
                stable_or_secure += 1

        apply_zebra_stripe(table)

        # Summary
        document.add_paragraph(
            f"{flagged_or_risk} programs flagged or at risk, "
            f"{stable_or_secure} stable or secure",
            style="HS Body",
        )

        document.add_page_break()

    def _render_ecoregion_priorities(
        self, document: Document, style_manager: StyleManager
    ) -> None:
        """Render the Ecoregion Strategic Priorities section (Section 3)."""
        document.add_heading("Ecoregion Strategic Priorities", level=2)

        document.add_paragraph(
            "Priority programs by ecoregion, based on geographic climate "
            "risk profiles.",
            style="HS Body",
        )

        eco_data = self.ecoregion.get_all_ecoregions()
        ecoregions = eco_data.get("ecoregions", {})

        for eco_id in sorted(ecoregions.keys()):
            eco_def = ecoregions[eco_id]
            eco_name = eco_def.get("name", eco_id)

            document.add_heading(eco_name, level=3)

            priority_ids = self.ecoregion.get_priority_programs(eco_id)

            if not priority_ids:
                document.add_paragraph(
                    "No specific program priorities defined for this ecoregion.",
                    style="HS Body",
                )
                continue

            for pid in priority_ids:
                prog = self.programs.get(pid)
                if prog:
                    ci_status = prog.get("ci_status", "")
                    label = CI_STATUS_LABELS.get(ci_status, ci_status)
                    para = document.add_paragraph(style="HS Body")
                    run = para.add_run(prog.get("name", pid))
                    run.bold = True
                    para.add_run(f" -- {label}")
                else:
                    document.add_paragraph(
                        f"{pid} (not in current inventory)", style="HS Body"
                    )

        document.add_page_break()

    def _render_science_threats(
        self, document: Document, style_manager: StyleManager
    ) -> None:
        """Render the Science Infrastructure Threats section (Section 4)."""
        document.add_heading("Science Infrastructure Threats", level=2)

        if not self.monitor_data:
            document.add_paragraph(
                "Monitor data not available. Run the scan pipeline to "
                "generate threat assessments.",
                style="HS Body",
            )
            document.add_page_break()
            return

        # Try to extract alerts from monitor data
        alerts = self.monitor_data.get("alerts", [])
        findings = self.monitor_data.get("findings", [])
        threats = self.monitor_data.get("threats", [])
        monitors = self.monitor_data.get("monitors", [])

        items = alerts or findings or threats or monitors

        if not items:
            document.add_paragraph(
                "Monitor data format not recognized. See "
                "outputs/LATEST-MONITOR-DATA.json for raw data.",
                style="HS Body",
            )
            document.add_page_break()
            return

        for item in items:
            title = item.get("title", item.get("name", "Untitled"))
            description = item.get("description", item.get("detail", ""))
            severity = item.get("severity", item.get("level", ""))

            # Title bold
            title_para = document.add_paragraph(style="HS Body")
            title_run = title_para.add_run(title)
            title_run.bold = True

            if severity:
                title_para.add_run(f"  [{severity.upper()}]")

            if description:
                document.add_paragraph(description, style="HS Body")

        document.add_page_break()

    def _render_fema_analysis(
        self, document: Document, style_manager: StyleManager
    ) -> None:
        """Render the FEMA Programs Analysis section (Section 5)."""
        document.add_heading("FEMA Programs Analysis", level=2)

        document.add_paragraph(
            "Dedicated analysis of FEMA climate resilience programs critical "
            "for Tribal hazard mitigation.",
            style="HS Body",
        )

        # Filter for FEMA programs
        fema_programs = [
            p for p in self.programs.values()
            if p.get("id", "") in _FEMA_PROGRAM_IDS
        ]
        fema_programs.sort(key=lambda p: p.get("name", ""))

        if not fema_programs:
            document.add_paragraph(
                "No FEMA programs found in the current inventory.",
                style="HS Body",
            )
            document.add_page_break()
            return

        for program in fema_programs:
            # Program name bold
            name_para = document.add_paragraph(style="HS Body")
            name_run = name_para.add_run(program.get("name", ""))
            name_run.bold = True

            # Status badge
            ci_status = program.get("ci_status", "")
            if ci_status:
                badge_para = document.add_paragraph(style="HS Body")
                add_status_badge(badge_para, ci_status)

            # Description
            description = program.get("description", "")
            if description:
                document.add_paragraph(description, style="HS Body")

            # Advocacy lever
            advocacy_lever = program.get("advocacy_lever", "")
            if advocacy_lever:
                lever_para = document.add_paragraph(style="HS Body")
                lever_run = lever_para.add_run("Advocacy Lever: ")
                lever_run.bold = True
                lever_para.add_run(advocacy_lever)

        # BCR methodology note
        document.add_paragraph(
            "FEMA programs use a 4:1 benefit-cost ratio standard for "
            "mitigation project evaluation.",
            style="HS Small",
        )

        document.add_page_break()

    def _render_cross_cutting_framework(
        self, document: Document, style_manager: StyleManager
    ) -> None:
        """Render the Cross-Cutting Policy Framework section (Section 6)."""
        document.add_heading("Cross-Cutting Policy Framework", level=2)

        structural_asks = self.graph_schema.get("structural_asks", [])

        if not structural_asks:
            document.add_paragraph(
                "No structural asks defined in policy graph.",
                style="HS Body",
            )
            document.add_page_break()
            return

        document.add_paragraph(
            "Five structural policy changes that would advance climate "
            "resilience funding across multiple programs.",
            style="HS Body",
        )

        for ask in structural_asks:
            ask_name = ask.get("name", "")
            ask_desc = ask.get("description", "")
            ask_target = ask.get("target", "")
            ask_urgency = ask.get("urgency", "")

            # Name bold + urgency
            name_para = document.add_paragraph(style="HS Body")
            name_run = name_para.add_run(ask_name)
            name_run.bold = True
            if ask_urgency:
                name_para.add_run(f"  [{ask_urgency}]")

            # Description
            if ask_desc:
                document.add_paragraph(ask_desc, style="HS Body")

            # Target
            if ask_target:
                document.add_paragraph(
                    f"Target: {ask_target}", style="HS Body"
                )

            # Programs advanced
            ask_programs = ask.get("programs", [])
            if ask_programs:
                prog_names = []
                for pid in ask_programs:
                    prog = self.programs.get(pid)
                    prog_names.append(
                        prog.get("name", pid) if prog else pid
                    )
                document.add_paragraph(
                    f"Programs advanced: {', '.join(prog_names)}",
                    style="HS Body",
                )

        document.add_page_break()

    def _render_messaging_guidance(
        self, document: Document, style_manager: StyleManager
    ) -> None:
        """Render the Messaging Guidance section (Section 7)."""
        document.add_heading("Messaging Guidance", level=2)

        document.add_paragraph(
            "Advocacy framing organized by strategic goal. Each program's "
            "messaging is calibrated to its current CI status.",
            style="HS Body",
        )

        # Group programs by goal
        goal_groups: dict[str, list[dict]] = {
            "DEFEND": [],
            "PROTECT": [],
            "EXPAND": [],
        }

        for program in self.programs.values():
            ci_status = program.get("ci_status", "STABLE")
            goal = _GOAL_BY_STATUS.get(ci_status, "EXPAND")
            goal_groups[goal].append(program)

        # Sort programs within each group by name
        for goal in goal_groups:
            goal_groups[goal].sort(key=lambda p: p.get("name", ""))

        # Render in order: DEFEND, PROTECT, EXPAND
        for goal in ("DEFEND", "PROTECT", "EXPAND"):
            programs_in_group = goal_groups[goal]
            heading_text = _GOAL_HEADINGS.get(goal, goal)

            document.add_heading(heading_text, level=3)

            if not programs_in_group:
                document.add_paragraph(
                    "No programs in this category.", style="HS Body"
                )
                continue

            for program in programs_in_group:
                ci_status = program.get("ci_status", "STABLE")

                # Program name bold
                name_para = document.add_paragraph(style="HS Body")
                name_run = name_para.add_run(program.get("name", ""))
                name_run.bold = True

                # Tightened language
                tightened = program.get("tightened_language", "")
                if tightened:
                    document.add_paragraph(tightened, style="HS Body")

                # Advocacy lever
                advocacy_lever = program.get("advocacy_lever", "")
                if advocacy_lever:
                    lever_para = document.add_paragraph(style="HS Body")
                    lever_run = lever_para.add_run("Lever: ")
                    lever_run.bold = True
                    lever_para.add_run(advocacy_lever)

                # Framing verbs
                framing = FRAMING_BY_STATUS.get(
                    ci_status, FRAMING_BY_STATUS.get("STABLE", {})
                )
                verbs = framing.get("verbs", [])
                if verbs:
                    document.add_paragraph(
                        f"Framing: {', '.join(verbs)}", style="HS Body"
                    )

        # Final note
        document.add_paragraph(
            "All messaging is template-based from program inventory data. "
            "No AI-generated text.",
            style="HS Small",
        )
