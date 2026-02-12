"""DocxEngine -- assembles complete per-Tribe advocacy packet DOCX documents.

Creates a python-docx Document with programmatic styles, delegates section
rendering to specialized renderers, and saves to the output directory.

Supports template-based workflow: if a pre-built template exists at
``templates/hot_sheet_template.docx``, opens it (with all HS-* styles
pre-baked) instead of creating styles from scratch on every document.
This saves ~0.2s per document in batch generation.

Air gap: no organizational names or tool attribution appear in any
generated document surface (headers, footers, cover pages, body text).
"""

from __future__ import annotations

import logging
import os
import tempfile
from pathlib import Path

from src.paths import PACKETS_OUTPUT_DIR, PROJECT_ROOT

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from src.packets.context import TribePacketContext
from src.packets.doc_types import DocumentTypeConfig
from src.packets.docx_hotsheet import HotSheetRenderer
from src.packets.docx_sections import (
    render_appendix,
    render_change_tracking,
    render_cover_page,
    render_delegation_section,
    render_executive_summary,
    render_hazard_summary,
    render_messaging_framework,
    render_structural_asks_section,
    render_table_of_contents,
)
from src.packets.docx_styles import COLORS, StyleManager
from src.packets.economic import TribeEconomicSummary

logger = logging.getLogger(__name__)

DEFAULT_TEMPLATE_PATH = PROJECT_ROOT / "templates" / "hot_sheet_template.docx"


class DocxEngine:
    """Assembles and saves styled DOCX advocacy packets.

    Creates a python-docx Document with all custom styles applied via
    StyleManager, configures page layout, and saves to the output
    directory using an atomic write pattern.

    Supports audience-differentiated generation via DocumentTypeConfig.
    When ``doc_type_config`` is provided, headers, footers, and cover
    pages are tailored to the target audience (internal vs. congressional).

    Usage::

        engine = DocxEngine(config, programs, doc_type_config=DOC_A)
        doc, sm = engine.create_document()
        # ... add content via renderers ...
        path = engine.save(doc, "epa_123")
    """

    def __init__(
        self,
        config: dict,
        programs: dict[str, dict],
        template_path: Path | str | None = None,
        doc_type_config: DocumentTypeConfig | None = None,
    ) -> None:
        """Initialize the engine with config and program inventory.

        Args:
            config: Application configuration dict (with ``packets`` section).
            programs: Dict of program dicts keyed by program ID.
            template_path: Optional path to a pre-built .docx template.
                If ``None``, uses ``DEFAULT_TEMPLATE_PATH`` when it exists,
                otherwise falls back to creating styles from scratch.
                Pass ``False`` to explicitly disable template loading.
            doc_type_config: Optional document type configuration for
                audience-differentiated generation. Defaults to None
                (backward compatible -- uses DOC_B-style headers/footers).
        """
        self.config = config
        self.programs = programs
        self.doc_type_config = doc_type_config
        raw_dir = config.get("packets", {}).get("output_dir")
        self.output_dir = Path(raw_dir) if raw_dir else PACKETS_OUTPUT_DIR
        if not self.output_dir.is_absolute():
            self.output_dir = PROJECT_ROOT / self.output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Resolve template path
        if template_path is False:
            self._template_path: Path | None = None
        elif template_path is not None:
            self._template_path = Path(template_path)
        elif DEFAULT_TEMPLATE_PATH.exists():
            self._template_path = DEFAULT_TEMPLATE_PATH
        else:
            self._template_path = None

        logger.debug(
            "DocxEngine initialized, output_dir=%s, template=%s, doc_type=%s",
            self.output_dir,
            self._template_path,
            doc_type_config.doc_type if doc_type_config else "default",
        )

    def create_document(self) -> tuple[Document, StyleManager]:
        """Create a Document with all custom styles and page layout.

        If a pre-built template is available, opens it (styles already
        registered, margins set, header/footer configured).  Otherwise
        creates a blank document and applies styles programmatically.

        Note: When a template is used, its header/footer are replaced
        with air-gap-compliant defaults. The actual doc-type-specific
        header/footer is applied in ``generate()`` where context
        (tribe_name, region_name) is available.

        Returns:
            Tuple of (Document, StyleManager) ready for content rendering.
        """
        if self._template_path and self._template_path.exists():
            document = Document(str(self._template_path))
            # StyleManager is idempotent -- skips already-registered styles
            style_manager = StyleManager(document)
            # Replace template header/footer with air-gap-clean defaults
            self._setup_header_footer(document)
            logger.debug(
                "Opened document from template: %s", self._template_path
            )
        else:
            document = Document()
            style_manager = StyleManager(document)

            # Configure page margins (1 inch all sides)
            for section in document.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            self._setup_header_footer(document)
            logger.debug("Created new document with styles and page layout")

        return document, style_manager

    def _setup_header_footer(
        self,
        document: Document,
        header_text: str | None = None,
        footer_text: str | None = None,
        confidential: bool = False,
    ) -> None:
        """Configure document header and footer text.

        Air-gap compliant: no organizational names or tool attribution.

        When called without arguments (initial document setup), uses
        generic descriptive text. When called with specific text
        (from generate()), uses doc-type-specific content.

        Args:
            document: The python-docx Document to configure.
            header_text: Header text to display. Defaults to a generic
                FY26 program overview header.
            footer_text: Footer text prefix. Defaults to congressional
                office use text.
            confidential: If True, adds CONFIDENTIAL marking.
        """
        if header_text is None:
            header_text = "FY26 Climate Resilience Program Priorities"
        if footer_text is None:
            footer_text = "For Congressional Office Use"

        for section in document.sections:
            # Header
            header = section.header
            header.is_linked_to_previous = False
            header_para = (
                header.paragraphs[0]
                if header.paragraphs
                else header.add_paragraph()
            )
            # Clear existing header content (removes template air gap violations)
            header_para.clear()
            header_run = header_para.add_run(header_text)
            header_run.font.size = Pt(8)
            header_run.font.color.rgb = COLORS.muted
            header_run.font.name = "Arial"

            # Footer
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = (
                footer.paragraphs[0]
                if footer.paragraphs
                else footer.add_paragraph()
            )
            # Clear existing footer content (removes template air gap violations)
            footer_para.clear()
            footer_prefix = (
                f"CONFIDENTIAL | {footer_text}"
                if confidential
                else footer_text
            )
            footer_run = footer_para.add_run(footer_prefix)
            footer_run.font.size = Pt(7)
            footer_run.font.color.rgb = COLORS.muted
            footer_run.font.name = "Arial"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save(self, document: Document, tribe_id: str) -> Path:
        """Save the document to disk using an atomic write pattern.

        Writes to a temporary file first, then atomically replaces the
        target path. This prevents partial/corrupt files on crash.

        Warning:
            Not safe for concurrent execution by multiple processes
            targeting the same tribe_id. If parallel packet generation
            is needed, use external locking or a task queue with
            single-worker concurrency per tribe_id.

        Args:
            document: The completed python-docx Document.
            tribe_id: Tribe identifier used as filename stem.

        Returns:
            Path to the saved .docx file.

        Raises:
            ValueError: If tribe_id contains path traversal sequences.
        """
        # Sanitize tribe_id to prevent path traversal
        safe_id = Path(tribe_id).name
        if not safe_id or safe_id in (".", ".."):
            raise ValueError(f"Invalid tribe_id: {tribe_id!r}")

        output_path = self.output_dir / f"{safe_id}.docx"
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Atomic write: save to tmp, then replace.
        # Close handle before save to avoid Windows file-locking conflicts.
        tmp_fd = tempfile.NamedTemporaryFile(
            dir=str(self.output_dir), suffix=".docx", delete=False
        )
        tmp_name = tmp_fd.name
        tmp_fd.close()
        try:
            document.save(tmp_name)
            os.replace(tmp_name, str(output_path))
        except Exception:
            try:
                os.unlink(tmp_name)
            except OSError:
                pass
            raise

        logger.info("Saved advocacy packet: %s", output_path)
        return output_path

    def generate(
        self,
        context: TribePacketContext,
        relevant_programs: list[dict],
        economic_summary: TribeEconomicSummary,
        structural_asks: list[dict],
        omitted_programs: list[dict] | None = None,
        changes: list[dict] | None = None,
        previous_date: str | None = None,
        doc_type_config: DocumentTypeConfig | None = None,
    ) -> Path:
        """Generate a complete 9-section advocacy packet DOCX for a single Tribe.

        Assembles the full document in order: cover page, table of contents,
        executive summary, congressional delegation, Hot Sheet pages for each
        relevant program, hazard profile summary, structural policy asks,
        change tracking (if changes exist), and appendix of omitted programs.

        Args:
            context: TribePacketContext with all Tribe data.
            relevant_programs: List of program dicts relevant to this Tribe.
            economic_summary: Aggregated economic impact calculations.
            structural_asks: List of structural ask dicts from graph schema.
            omitted_programs: Optional list of programs not included in
                Hot Sheets (rendered in appendix).
            changes: Optional list of change dicts for change tracking
                (OPS-03). Empty or None means no "Since Last Packet" section.
            previous_date: Optional ISO date string of previous packet
                generation, displayed in the change tracking header.
            doc_type_config: Optional document type configuration that
                overrides self.doc_type_config for this generation.

        Returns:
            Path to the saved .docx file.
        """
        # Resolve doc_type_config: parameter > instance > None
        dtc = doc_type_config or self.doc_type_config

        document, style_manager = self.create_document()

        # Apply doc-type-specific headers/footers now that context is available
        if dtc is not None:
            header_text = dtc.format_header(
                tribe_name=context.tribe_name,
                region_name=None,  # Regional docs will pass region_name
            )
            self._setup_header_footer(
                document,
                header_text=header_text,
                footer_text=dtc.footer_template,
                confidential=dtc.confidential,
            )

        # 1. Cover page
        render_cover_page(
            document, context, style_manager, doc_type_config=dtc
        )

        # 2. Table of contents
        render_table_of_contents(
            document, relevant_programs, context, style_manager
        )

        # 3. Executive summary
        render_executive_summary(
            document, context, relevant_programs, economic_summary,
            style_manager, doc_type_config=dtc,
        )

        # 4. Congressional delegation
        render_delegation_section(
            document, context, style_manager, doc_type_config=dtc
        )

        # 5. Hot Sheets for all relevant programs (8-12 per Tribe)
        renderer = HotSheetRenderer(document, style_manager)
        renderer.render_all_hotsheets(
            context, relevant_programs, economic_summary, structural_asks,
            doc_type_config=dtc,
        )

        # 6. Hazard profile summary
        document.add_page_break()
        render_hazard_summary(
            document, context, style_manager, doc_type_config=dtc
        )

        # 7. Structural asks standalone
        render_structural_asks_section(
            document, structural_asks, context, style_manager,
            doc_type_config=dtc,
        )

        # 8. Messaging framework (Doc A only -- internal with messaging flag)
        if dtc is not None and dtc.include_messaging_framework:
            render_messaging_framework(document, context, style_manager)

        # 9. Change tracking (OPS-03, internal only, if changes exist)
        if changes and (dtc is None or dtc.is_internal):
            document.add_page_break()
            render_change_tracking(
                document, changes, previous_date, style_manager
            )

        # 10. Appendix (omitted programs or "all included" note)
        render_appendix(
            document, omitted_programs or [], context, style_manager,
            doc_type_config=dtc,
        )

        # Save with doc-type-aware filename
        tribe_id = getattr(context, "tribe_id", "unknown")
        if dtc is not None:
            filename_stem = dtc.format_filename(tribe_id).replace(".docx", "")
        else:
            filename_stem = tribe_id
        output_path = self.save(document, filename_stem)

        logger.info(
            "Generated complete packet for %s (doc_type=%s): "
            "%d Hot Sheets, %d structural asks, %d omitted, %d changes",
            context.tribe_name,
            dtc.doc_type if dtc else "default",
            len(relevant_programs),
            len(structural_asks),
            len(omitted_programs or []),
            len(changes or []),
        )
        return output_path
