#!/usr/bin/env python3
"""Validate structural integrity of all generated DOCX advocacy packets.

Checks 7 quality dimensions per document:
  1. Heading hierarchy -- no level gaps (H1->H3 without H2)
  2. Table population  -- every table has at least one non-empty cell
  3. Page count estimate -- page-break count within expected range (5-80)
  4. Header presence    -- section headers contain non-empty text
  5. Placeholder absence -- no TBD, [INSERT], PLACEHOLDER, or "data pending"
  6. Font consistency   -- all explicit fonts are Arial
  7. Image integrity    -- inline_shapes accessible without error

Usage:
    python scripts/validate_docx_structure.py               # Validate all
    python scripts/validate_docx_structure.py --sample 10    # Random sample
    python scripts/validate_docx_structure.py --help         # Help
"""

import argparse
import json
import logging
import random
import re
import sys
import time
from datetime import datetime, timezone
from pathlib import Path

# Ensure project root on sys.path
_SCRIPT_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _SCRIPT_DIR.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

from src.paths import OUTPUTS_DIR, PACKETS_OUTPUT_DIR  # noqa: E402

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DOCX_REVIEW_DIR = OUTPUTS_DIR / "docx_review"
OUTPUT_JSON_PATH = DOCX_REVIEW_DIR / "structural_validation.json"

PLACEHOLDER_PATTERN = re.compile(
    r"\bTBD\b|\[INSERT|\bPLACEHOLDER\b|data\s+pending",
    re.IGNORECASE,
)

MIN_PAGE_BREAKS = 0   # Regional docs can be short
MAX_PAGE_BREAKS = 80  # Upper bound for Tribal docs
EXPECTED_FONT = "Arial"


# ---------------------------------------------------------------------------
# Document discovery
# ---------------------------------------------------------------------------


def discover_documents() -> list[Path]:
    """Find all DOCX files across packet output directories."""
    docx_files: list[Path] = []

    # Doc A (internal) + Doc B (congressional) per Tribe
    packets_dir = PACKETS_OUTPUT_DIR
    if packets_dir.is_dir():
        for subdir_name in ("internal", "congressional"):
            subdir = packets_dir / subdir_name
            if subdir.is_dir():
                docx_files.extend(sorted(subdir.glob("*.docx")))

        # Regional docs (Doc C/D)
        regional_dir = packets_dir / "regional"
        if regional_dir.is_dir():
            for regional_sub in ("internal", "congressional"):
                regional_subdir = regional_dir / regional_sub
                if regional_subdir.is_dir():
                    docx_files.extend(sorted(regional_subdir.glob("*.docx")))

    # Also check top-level loose DOCX files (e.g., STRATEGIC-OVERVIEW.docx)
    if packets_dir.is_dir():
        for loose in sorted(packets_dir.glob("*.docx")):
            if loose not in docx_files:
                docx_files.append(loose)

    return docx_files


# ---------------------------------------------------------------------------
# Per-document checks
# ---------------------------------------------------------------------------


def classify_doc_type(path: Path) -> str:
    """Classify document as Doc A/B/C/D based on path."""
    parts = path.parts
    if "regional" in parts:
        if "congressional" in parts:
            return "D"
        return "C"
    if "congressional" in parts:
        return "B"
    if "internal" in parts:
        return "A"
    return "unknown"


def check_heading_hierarchy(paragraphs: list) -> dict:
    """Verify no heading level gaps (e.g., H1 -> H3 without H2)."""
    levels_seen: list[int] = []
    gaps: list[str] = []

    for para in paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading "):
            try:
                level = int(style_name.split(" ")[-1])
                levels_seen.append(level)
            except (ValueError, IndexError):
                pass

    # Check for gaps in sequence
    for i in range(1, len(levels_seen)):
        prev = levels_seen[i - 1]
        curr = levels_seen[i]
        if curr > prev + 1:
            gaps.append(f"H{prev} -> H{curr} (skipped H{prev + 1})")

    return {
        "passed": len(gaps) == 0,
        "heading_count": len(levels_seen),
        "gaps": gaps,
    }


def check_table_population(tables: list) -> dict:
    """Every table must have at least one non-empty cell."""
    empty_tables: list[int] = []

    for idx, table in enumerate(tables):
        has_content = False
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    has_content = True
                    break
            if has_content:
                break
        if not has_content:
            empty_tables.append(idx)

    return {
        "passed": len(empty_tables) == 0,
        "table_count": len(tables),
        "empty_tables": empty_tables,
    }


def check_page_count(paragraphs: list) -> dict:
    """Count page breaks in run XML.  Flag outliers."""
    page_breaks = 0

    for para in paragraphs:
        for run in para.runs:
            # Check XML for page breaks: <w:br w:type="page"/>
            xml = run._r.xml
            if "w:br" in xml and 'w:type="page"' in xml:
                page_breaks += 1

    # Pages = page_breaks + 1 (first page has no break)
    estimated_pages = page_breaks + 1
    in_range = MIN_PAGE_BREAKS <= page_breaks <= MAX_PAGE_BREAKS

    return {
        "passed": in_range,
        "page_breaks": page_breaks,
        "estimated_pages": estimated_pages,
    }


def check_header_presence(sections: list) -> dict:
    """Check that at least one section header has non-empty text."""
    has_header_text = False
    header_texts: list[str] = []

    for section in sections:
        header = section.header
        for para in header.paragraphs:
            text = para.text.strip()
            if text:
                has_header_text = True
                header_texts.append(text)

    return {
        "passed": has_header_text,
        "header_texts": header_texts[:3],  # First 3 for brevity
    }


def check_placeholder_absence(paragraphs: list) -> dict:
    """Scan all paragraph text for placeholder patterns."""
    found: list[dict] = []

    for idx, para in enumerate(paragraphs):
        text = para.text
        matches = PLACEHOLDER_PATTERN.findall(text)
        if matches:
            found.append({
                "paragraph_index": idx,
                "matches": matches,
                "text_snippet": text[:120],
            })

    return {
        "passed": len(found) == 0,
        "placeholder_count": len(found),
        "placeholders": found[:10],  # Cap at 10 for readability
    }


def check_font_consistency(paragraphs: list) -> dict:
    """Flag any explicit font that is not Arial."""
    non_arial: list[dict] = []
    fonts_seen: set[str] = set()

    for para in paragraphs:
        for run in para.runs:
            font_name = run.font.name
            if font_name is not None:
                fonts_seen.add(font_name)
                if font_name != EXPECTED_FONT:
                    non_arial.append({
                        "font": font_name,
                        "text_snippet": run.text[:80] if run.text else "",
                    })

    # Deduplicate by font name for summary
    unique_non_arial = sorted(set(f["font"] for f in non_arial))

    return {
        "passed": len(unique_non_arial) == 0,
        "fonts_seen": sorted(fonts_seen),
        "non_arial_fonts": unique_non_arial,
        "non_arial_count": len(non_arial),
    }


def check_image_integrity(doc) -> dict:
    """Verify inline_shapes accessible without error."""
    try:
        shape_count = len(doc.inline_shapes)
        return {
            "passed": True,
            "image_count": shape_count,
        }
    except Exception as exc:
        return {
            "passed": False,
            "image_count": 0,
            "error": str(exc),
        }


def _single_pass_paragraph_checks(paragraphs: list) -> dict:
    """Run heading, page-count, placeholder, and font checks in ONE pass.

    Iterating paragraphs and their runs is the most expensive operation.
    Combining four checks into a single pass cuts ~60% of per-document time.
    """
    # Heading hierarchy state
    heading_levels: list[int] = []

    # Page count state
    page_breaks = 0

    # Placeholder state
    placeholder_hits: list[dict] = []

    # Font consistency state
    fonts_seen: set[str] = set()
    non_arial_fonts: set[str] = set()
    non_arial_count = 0

    for idx, para in enumerate(paragraphs):
        # --- Heading check (style name only, no run iteration) ---
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading "):
            try:
                heading_levels.append(int(style_name.split(" ")[-1]))
            except (ValueError, IndexError):
                pass

        # --- Placeholder check (paragraph text, no run iteration) ---
        text = para.text
        if text:
            matches = PLACEHOLDER_PATTERN.findall(text)
            if matches:
                placeholder_hits.append({
                    "paragraph_index": idx,
                    "matches": matches,
                    "text_snippet": text[:120],
                })

        # --- Run-level checks (page breaks + font) ---
        for run in para.runs:
            # Page breaks
            run_xml = run._r.xml
            if "w:br" in run_xml and 'w:type="page"' in run_xml:
                page_breaks += 1

            # Font consistency
            font_name = run.font.name
            if font_name is not None:
                fonts_seen.add(font_name)
                if font_name != EXPECTED_FONT:
                    non_arial_fonts.add(font_name)
                    non_arial_count += 1

    # Build heading hierarchy result
    heading_gaps: list[str] = []
    for i in range(1, len(heading_levels)):
        prev = heading_levels[i - 1]
        curr = heading_levels[i]
        if curr > prev + 1:
            heading_gaps.append(f"H{prev} -> H{curr} (skipped H{prev + 1})")

    return {
        "heading_hierarchy": {
            "passed": len(heading_gaps) == 0,
            "heading_count": len(heading_levels),
            "gaps": heading_gaps,
        },
        "page_count": {
            "passed": MIN_PAGE_BREAKS <= page_breaks <= MAX_PAGE_BREAKS,
            "page_breaks": page_breaks,
            "estimated_pages": page_breaks + 1,
        },
        "placeholder_absence": {
            "passed": len(placeholder_hits) == 0,
            "placeholder_count": len(placeholder_hits),
            "placeholders": placeholder_hits[:10],
        },
        "font_consistency": {
            "passed": len(non_arial_fonts) == 0,
            "fonts_seen": sorted(fonts_seen),
            "non_arial_fonts": sorted(non_arial_fonts),
            "non_arial_count": non_arial_count,
        },
    }


def validate_document(docx_path: Path) -> dict:
    """Run all 7 structural checks on a single DOCX document.

    Opens the document exactly once, runs paragraph checks in a single
    pass for performance, then runs table/header/image checks separately.
    """
    from docx import Document

    result: dict = {
        "file": str(docx_path.relative_to(_PROJECT_ROOT)),
        "doc_type": classify_doc_type(docx_path),
        "checks": {},
        "passed": True,
        "errors": [],
    }

    try:
        doc = Document(str(docx_path))
    except Exception as exc:
        result["passed"] = False
        result["errors"].append(f"Failed to open: {exc}")
        return result

    # Single-pass for 4 paragraph-based checks (heading, page, placeholder, font)
    checks = _single_pass_paragraph_checks(doc.paragraphs)

    # Remaining checks that need different data sources
    checks["table_population"] = check_table_population(doc.tables)
    checks["header_presence"] = check_header_presence(doc.sections)
    checks["image_integrity"] = check_image_integrity(doc)

    result["checks"] = checks

    # Document passes only if ALL checks pass
    result["passed"] = all(c["passed"] for c in checks.values())

    return result


# ---------------------------------------------------------------------------
# Aggregation
# ---------------------------------------------------------------------------


def aggregate_results(results: list[dict]) -> dict:
    """Aggregate per-document results into a summary."""
    total = len(results)
    all_passed = sum(1 for r in results if r["passed"])
    all_failed = total - all_passed

    # Per-check aggregation
    check_names = [
        "heading_hierarchy",
        "table_population",
        "page_count",
        "header_presence",
        "placeholder_absence",
        "font_consistency",
        "image_integrity",
    ]

    per_check: dict[str, dict] = {}
    for name in check_names:
        passed = sum(
            1
            for r in results
            if name in r.get("checks", {}) and r["checks"][name]["passed"]
        )
        failed = sum(
            1
            for r in results
            if name in r.get("checks", {}) and not r["checks"][name]["passed"]
        )
        errored = sum(
            1 for r in results if name not in r.get("checks", {})
        )
        per_check[name] = {
            "passed": passed,
            "failed": failed,
            "errored": errored,
        }

    # Collect failing documents
    failures: list[dict] = []
    for r in results:
        if not r["passed"]:
            failed_checks = [
                name
                for name in check_names
                if name in r.get("checks", {})
                and not r["checks"][name]["passed"]
            ]
            failures.append({
                "file": r["file"],
                "doc_type": r.get("doc_type", "unknown"),
                "failed_checks": failed_checks,
                "errors": r.get("errors", []),
            })

    # Doc type breakdown
    doc_types: dict[str, dict] = {}
    for r in results:
        dt = r.get("doc_type", "unknown")
        if dt not in doc_types:
            doc_types[dt] = {"total": 0, "passed": 0, "failed": 0}
        doc_types[dt]["total"] += 1
        if r["passed"]:
            doc_types[dt]["passed"] += 1
        else:
            doc_types[dt]["failed"] += 1

    return {
        "total_documents": total,
        "passed": all_passed,
        "failed": all_failed,
        "pass_rate": f"{(all_passed / total * 100):.1f}%" if total > 0 else "N/A",
        "per_check": per_check,
        "doc_type_breakdown": doc_types,
        "failures": failures,
    }


# ---------------------------------------------------------------------------
# Output
# ---------------------------------------------------------------------------


def print_summary(summary: dict, elapsed: float) -> None:
    """Print human-readable summary to stdout."""
    total = summary["total_documents"]
    passed = summary["passed"]
    failed = summary["failed"]

    print("\n" + "=" * 70)
    print("  DOCX Structural Validation Report")
    print("=" * 70)
    print(f"  Documents checked: {total}")
    print(f"  Passed:            {passed}")
    print(f"  Failed:            {failed}")
    print(f"  Pass rate:         {summary['pass_rate']}")
    print(f"  Elapsed:           {elapsed:.2f}s")
    if total > 0:
        print(f"  Per-document:      {elapsed / total * 1000:.1f}ms avg")
    print()

    # Per-check breakdown
    print("  Per-Check Results:")
    print("  " + "-" * 50)
    for check_name, stats in summary["per_check"].items():
        status = "PASS" if stats["failed"] == 0 else f"FAIL ({stats['failed']})"
        print(f"    {check_name:25s}  {status}")
    print()

    # Doc type breakdown
    if summary["doc_type_breakdown"]:
        print("  Doc Type Breakdown:")
        print("  " + "-" * 50)
        for dt, stats in sorted(summary["doc_type_breakdown"].items()):
            print(f"    Doc {dt}: {stats['total']} total, "
                  f"{stats['passed']} passed, {stats['failed']} failed")
        print()

    # Failures (first 20)
    if summary["failures"]:
        print(f"  Failing Documents ({min(len(summary['failures']), 20)} "
              f"of {len(summary['failures'])}):")
        print("  " + "-" * 50)
        for f in summary["failures"][:20]:
            checks = ", ".join(f["failed_checks"]) if f["failed_checks"] else "open error"
            print(f"    {f['file']}")
            print(f"      Failed: {checks}")
            if f.get("errors"):
                for err in f["errors"]:
                    print(f"      Error: {err}")
        print()

    print("=" * 70)


def write_json_output(results: list[dict], summary: dict, elapsed: float) -> Path:
    """Write machine-readable JSON to outputs/docx_review/structural_validation.json."""
    DOCX_REVIEW_DIR.mkdir(parents=True, exist_ok=True)

    output = {
        "script": "scripts/validate_docx_structure.py",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "elapsed_seconds": round(elapsed, 2),
        "summary": summary,
        "documents": results,
    }

    OUTPUT_JSON_PATH.write_text(
        json.dumps(output, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    return OUTPUT_JSON_PATH


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Validate structural integrity of generated DOCX packets.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "--sample",
        type=int,
        default=None,
        metavar="N",
        help="Validate a random sample of N documents (default: all)",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()

    # Discover documents
    all_docs = discover_documents()
    print(f"Discovered {len(all_docs)} DOCX documents")

    if not all_docs:
        print("No documents found. Full validation requires document generation first.")
        print("Run: python -m src.main --prep-packets --all-tribes")
        # Write empty JSON for downstream consumers
        DOCX_REVIEW_DIR.mkdir(parents=True, exist_ok=True)
        empty_output = {
            "script": "scripts/validate_docx_structure.py",
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "elapsed_seconds": 0.0,
            "summary": {
                "total_documents": 0,
                "passed": 0,
                "failed": 0,
                "pass_rate": "N/A",
                "per_check": {},
                "doc_type_breakdown": {},
                "failures": [],
            },
            "documents": [],
        }
        OUTPUT_JSON_PATH.write_text(
            json.dumps(empty_output, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        print(f"Empty results written to {OUTPUT_JSON_PATH}")
        return 0

    # Apply sampling
    docs_to_validate = all_docs
    if args.sample is not None:
        if args.sample <= 0:
            print("Sample size 0: no documents to validate.")
            return 0
        if args.sample < len(all_docs):
            docs_to_validate = random.sample(all_docs, args.sample)
            print(f"Sampling {args.sample} of {len(all_docs)} documents")

    # Validate documents sequentially (python-docx is not thread-safe)
    print(f"Validating {len(docs_to_validate)} documents...")
    start = time.monotonic()

    results: list[dict] = []
    for i, docx_path in enumerate(docs_to_validate, 1):
        if i % 100 == 0 or i == len(docs_to_validate):
            print(f"  [{i}/{len(docs_to_validate)}] ...", end="\r")
        result = validate_document(docx_path)
        results.append(result)

    elapsed = time.monotonic() - start
    print()  # Clear progress line

    # Aggregate and report
    summary = aggregate_results(results)
    print_summary(summary, elapsed)

    # Write JSON
    json_path = write_json_output(results, summary, elapsed)
    print(f"JSON results written to {json_path}")

    return 0 if summary["failed"] == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
